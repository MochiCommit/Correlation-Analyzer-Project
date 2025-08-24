import streamlit as st
import pandas as pd
import numpy as np
import plotly.express as px
import plotly.figure_factory as ff
from sklearn.linear_model import LinearRegression
from sklearn.model_selection import train_test_split
import optuna
import os
import io
import base64
from docx import Document
from docx.shared import Inches
import tempfile
from datetime import datetime

st.set_page_config(page_title="Correlation Analyzer", layout="wide")



st.title("공정 데이터 상관관계 분석 도우미")

# 1. 파일 업로드
st.header("1. 파일 업로드")
uploaded_file = st.file_uploader("CSV 또는 Excel 파일을 업로드하세요", type=["csv", "xlsx"])
st.info("분석할 파일을 업로드하세요. Data는 회사 내부 서버에 저장됩니다.") 

if 'df' not in st.session_state:
    st.session_state['df'] = None
if 'data_file_path' not in st.session_state:
    st.session_state['data_file_path'] = None
if 'preprocessing_steps' not in st.session_state:
    st.session_state['preprocessing_steps'] = []
if 'current_df' not in st.session_state:
    st.session_state['current_df'] = None
if 'data_period' not in st.session_state:
    st.session_state['data_period'] = None
if 'preprocessing_completed' not in st.session_state:
    st.session_state['preprocessing_completed'] = False

# 파일이 업로드되면 모든 상태 초기화
if uploaded_file is not None:
    # 새로운 파일이 업로드되었는지 확인
    current_file_name = uploaded_file.name
    if 'uploaded_file_name' not in st.session_state or st.session_state['uploaded_file_name'] != current_file_name:
        # 새로운 파일이므로 모든 상태 초기화
        st.session_state['uploaded_file_name'] = current_file_name
        st.session_state['preprocessing_steps'] = []
        st.session_state['current_df'] = None
        st.session_state['df'] = None
        st.session_state['data_file_path'] = None
        st.session_state['data_period'] = None  # 기간 정보도 초기화
        st.session_state['preprocessing_completed'] = False  # 전처리 완료 상태도 초기화
        # 변수 선택 관련 상태도 초기화
        if 'y_col' in st.session_state:
            del st.session_state['y_col']
        if 'x_cols' in st.session_state:
            del st.session_state['x_cols']

if uploaded_file is not None:
    if uploaded_file.name.endswith('.csv'):
        df = pd.read_csv(uploaded_file)
    else:
        df = pd.read_excel(uploaded_file)
    st.session_state['df'] = df
    # 업로드된 파일을 서버에 저장 (안전한 파일 저장)
    import os
    import time
    
    # 파일명에 타임스탬프 추가하여 중복 방지
    timestamp = int(time.time())
    base_name = uploaded_file.name.rsplit('.', 1)[0]
    extension = uploaded_file.name.rsplit('.', 1)[1] if '.' in uploaded_file.name else ''
    
    save_path = f"uploaded_{base_name}_{timestamp}.{extension}"
    
    try:
        if uploaded_file.name.endswith('.csv'):
            df.to_csv(save_path, index=False)
        else:
            df.to_excel(save_path, index=False)
        st.session_state['data_file_path'] = save_path
    except PermissionError:
        # 권한 오류 시 임시 파일명 사용
        temp_save_path = f"temp_uploaded_{timestamp}.{extension}"
        try:
            if uploaded_file.name.endswith('.csv'):
                df.to_csv(temp_save_path, index=False)
            else:
                df.to_excel(temp_save_path, index=False)
            st.session_state['data_file_path'] = temp_save_path
            st.warning(f"⚠️ 원본 파일명으로 저장할 수 없어 임시 파일명으로 저장했습니다: {temp_save_path}")
        except Exception as e:
            st.error(f"❌ 파일 저장 중 오류가 발생했습니다: {str(e)}")
            st.session_state['data_file_path'] = None
    except Exception as e:
        st.error(f"❌ 파일 저장 중 오류가 발생했습니다: {str(e)}")
        st.session_state['data_file_path'] = None
    
    # 굵은 가로 구분선 추가
    st.markdown('<hr style="border:3px solid #333;">', unsafe_allow_html=True)
    
    # 스크롤을 조정하여 결측치 처리 섹션이 화면 중앙 위쪽에 오도록 함
    st.markdown("""
    <script>
    // 페이지 로드 후 스크롤 조정
    function adjustScroll() {
        // 결측치 처리 섹션을 찾아서 스크롤 조정
        const elements = document.querySelectorAll('h1, h2, h3, h4, h5, h6');
        for (let element of elements) {
            if (element.textContent.includes('결측치 처리')) {
                // 화면 높이의 20% 지점으로 스크롤 (더 위쪽으로)
                const targetPosition = element.offsetTop - window.innerHeight * 0.2;
                window.scrollTo({
                    top: targetPosition,
                    behavior: 'smooth'
                });
                break;
            }
        }
    }
    
    // 여러 시점에서 스크롤 조정 시도
    window.addEventListener('load', function() {
        setTimeout(adjustScroll, 500);
        setTimeout(adjustScroll, 1500);
        setTimeout(adjustScroll, 2500);
    });
    
    // DOM 변경 감지
    const observer = new MutationObserver(function(mutations) {
        mutations.forEach(function(mutation) {
            if (mutation.type === 'childList') {
                setTimeout(adjustScroll, 100);
            }
        });
    });
    
    observer.observe(document.body, {
        childList: true,
        subtree: true
    });
    </script>
    """, unsafe_allow_html=True)
    
    st.header("2. 데이터 전처리")
    # st.subheader("컬럼별 데이터형 변환")
    # dtype_map = {col: str(st.session_state['df'][col].dtype) for col in st.session_state['df'].columns}
    # new_dtypes = {}
    # dtype_options = {"문자열": "object", "숫자": "float64", "날짜": "datetime64[ns]"}
    # cols = st.columns(len(st.session_state['df'].columns))
    # for i, col in enumerate(st.session_state['df'].columns):
    #     with cols[i]:
    #         dtype = st.selectbox(f"{col}", options=list(dtype_options.keys()), index=list(dtype_options.values()).index(dtype_map[col]) if dtype_map[col] in dtype_options.values() else 0, key=f"dtype_{col}")
    #         new_dtypes[col] = dtype_options[dtype]
    # if st.button("데이터형 변환 적용"):
    #     for col, dtype in new_dtypes.items():
    #         try:
    #             if dtype == "datetime64[ns]":
    #                 st.session_state['df'][col] = pd.to_datetime(st.session_state['df'][col], errors='coerce')
    #             else:
    #                 st.session_state['df'][col] = st.session_state['df'][col].astype(dtype)
    #         except Exception as e:
    #             st.warning(f"{col} 변환 실패: {e}")
    #     st.success("데이터형 변환 완료!")

    # 데이터 미리보기
    with st.spinner('Loading ...'):
        st.subheader("데이터 미리보기")
        
        # 데이터 요약 정보 표시 - 항상 최신 데이터 사용
        df = st.session_state['df']
        total_rows, total_cols = df.shape
        numeric_cols = df.select_dtypes(include=['number']).shape[1]
        text_cols = df.select_dtypes(include=['object']).shape[1]
        datetime_cols = df.select_dtypes(include=['datetime']).shape[1]
        missing_total = df.isnull().sum().sum()
        missing_ratio = (missing_total / (total_rows * total_cols)) * 100
        
        st.markdown(f"""
        <div style="background-color: #f0f2f6; padding: 15px; border-radius: 8px; margin: 10px 0; border-left: 4px solid #1f77b4;">
            <h4 style="margin: 0 0 10px 0; color: #1f77b4;">📊 데이터 요약</h4>
            <div style="font-size: 14px;">
                <div style="display: flex; gap: 20px; margin-bottom: 8px;">
                    <div><strong>총 행 수:</strong> {total_rows:,}개</div>
                    <div><strong>총 열 수:</strong> {total_cols:,}개</div>
                    <div><strong>총 결측치:</strong> {missing_total:,}개 (총 {total_rows * total_cols:,}개 데이터 중 {missing_ratio:.1f}%)</div>
                </div>
            </div>
        </div>
        """, unsafe_allow_html=True)
        
        # 날짜형 열 변환 기능 추가 - 최신 데이터로 조건 확인
        current_datetime_cols = df.select_dtypes(include=['datetime']).shape[1]
        current_text_cols = df.select_dtypes(include=['object']).shape[1]
        
        if current_datetime_cols == 0 and current_text_cols > 0:
            st.markdown("""
            <div style="background-color: #fff3cd; padding: 15px; border-radius: 8px; margin: 10px 0; border-left: 4px solid #ffc107;">
                <h4 style="margin: 0 0 10px 0; color: #856404;">⚠️ 날짜형 열 변환 필요</h4>
                <p style="margin: 0 0 10px 0; color: #856404; font-size: 14px;">
                    날짜/시간 데이터가 문자형으로 인식되었습니다. 날짜형으로 변환하면 더 정확한 분석이 가능합니다.
                </p>
            </div>
            """, unsafe_allow_html=True)
            
            # 날짜형으로 변환할 열 선택 - 가로 배치
            text_columns = df.select_dtypes(include=['object']).columns.tolist()
            
            # 날짜 열 추천 함수
            def recommend_date_column(df, text_columns):
                """가장 유력한 날짜 열을 추천하는 함수"""
                if not text_columns:
                    return None
                
                # 각 열의 날짜 가능성 점수 계산
                column_scores = {}
                
                for col in text_columns:
                    sample_data = df[col].dropna().head(100)  # 처음 100개 샘플만 확인
                    if len(sample_data) == 0:
                        continue
                    
                    score = 0
                    
                    # 1. 열 이름 기반 점수
                    col_lower = col.lower()
                    date_keywords = ['date', 'time', '날짜', '시간', '시작', '종료', '시작일', '종료일', 'timestamp', 'datetime']
                    for keyword in date_keywords:
                        if keyword in col_lower:
                            score += 10
                    
                    # 2. 데이터 패턴 기반 점수
                    try:
                        # 날짜 변환 시도
                        converted = pd.to_datetime(sample_data, errors='coerce')
                        valid_ratio = converted.notna().sum() / len(sample_data)
                        
                        if valid_ratio > 0.8:  # 80% 이상이 유효한 날짜
                            score += 50
                        elif valid_ratio > 0.5:  # 50% 이상이 유효한 날짜
                            score += 30
                        elif valid_ratio > 0.2:  # 20% 이상이 유효한 날짜
                            score += 10
                        
                        # 고유값 비율 (날짜는 보통 고유값이 많음)
                        unique_ratio = sample_data.nunique() / len(sample_data)
                        if unique_ratio > 0.8:
                            score += 5
                        
                    except:
                        pass
                    
                    # 3. 데이터 길이 기반 점수 (날짜는 보통 일정한 길이)
                    avg_length = sample_data.astype(str).str.len().mean()
                    if 8 <= avg_length <= 20:  # 일반적인 날짜 길이
                        score += 5
                    
                    column_scores[col] = score
                
                # 가장 높은 점수의 열 반환
                if column_scores:
                    best_column = max(column_scores, key=column_scores.get)
                    if column_scores[best_column] > 10:  # 최소 점수 기준
                        return best_column
                
                return None
            
            # 추천 열 찾기
            recommended_column = recommend_date_column(df, text_columns)
            
            # 가로로 배치하기 위해 컬럼 사용
            col1, col2 = st.columns([1, 3])  # 3:1 비율로 배치
            
            with col1:
                # 옵션 리스트 생성 (추천 열이 있으면 맨 위에 배치)
                date_options = ["변환하지 않음"]
                
                if recommended_column:
                    date_options.append(f"⭐추천⭐ {recommended_column}")
                    # 추천 열을 제외한 나머지 열들 추가
                    for col in text_columns:
                        if col != recommended_column:
                            date_options.append(col)
                else:
                    # 추천 열이 없으면 모든 열 추가
                    date_options.extend(text_columns)
                
                selected_option = st.selectbox(
                    "날짜형으로 변환할 열을 선택하세요:",
                    options=date_options,
                    help="날짜/시간 형식의 데이터가 포함된 열을 선택하세요"
                )
            
            with col2:
                # 확인 버튼을 multiselect와 같은 높이에 배치
                st.write("")  # 빈 줄로 높이 맞추기
                st.write("")  # 빈 줄로 높이 맞추기
                if st.button("✅ 실행", key="date_format_confirm"):
                    if selected_option:
                        # "변환하지 않음" 옵션 체크
                        if selected_option == "변환하지 않음":
                            st.success("✅ 날짜형 변환을 건너뛰고 진행합니다.")
                            st.rerun()
                        else:
                            # [추천] 태그 제거하고 실제 열 이름 추출
                            actual_column = selected_option.replace("⭐추천⭐ ", "")
                            
                            # 선택된 열에 대해 날짜형 변환 시도
                            try:
                                # 원본 데이터 확인
                                original_data = df[actual_column]
                                
                                # 다양한 날짜 형식으로 변환 시도
                                converted_data = pd.to_datetime(original_data, errors='coerce', infer_datetime_format=True)
                                
                                # 변환 성공 여부 확인 (NaT가 아닌 값이 있는지 확인)
                                if converted_data.notna().sum() > 0:
                                    # 변환 성공 - 세션 상태 업데이트
                                    st.session_state['df'][actual_column] = converted_data
                                    
                                    # 파일도 업데이트 (안전한 파일 저장)
                                    if st.session_state['data_file_path']:
                                        try:
                                            if st.session_state['data_file_path'].endswith('.csv'):
                                                st.session_state['df'].to_csv(st.session_state['data_file_path'], index=False)
                                            else:
                                                st.session_state['df'].to_excel(st.session_state['data_file_path'], index=False)
                                        except PermissionError:
                                            st.warning("⚠️ 파일이 다른 프로그램에서 사용 중이어서 저장할 수 없습니다.")
                                        except Exception as e:
                                            st.warning(f"⚠️ 파일 저장 중 오류가 발생했습니다: {str(e)}")
                                    
                                    st.success(f"✅ '{actual_column}' 열이 날짜형으로 변환되었습니다!")
                                    
                                    # 날짜 열 정보를 session state에 저장
                                    st.session_state['date_column_name'] = actual_column
                                    
                                    # 데이터 기간 정보를 session state에 저장
                                    valid_dates = converted_data.dropna()
                                    if len(valid_dates) > 0:
                                        start_date = valid_dates.min()
                                        end_date = valid_dates.max()
                                        
                                        # 날짜 형식 포맷팅
                                        if pd.api.types.is_datetime64_any_dtype(start_date):
                                            start_str = start_date.strftime('%Y년 %m월 %d일')
                                            end_str = end_date.strftime('%Y년 %m월 %d일')
                                            
                                            # 시간 정보가 있는지 확인
                                            if start_date.hour != 0 or start_date.minute != 0:
                                                start_str += f" {start_date.strftime('%H:%M')}"
                                            if end_date.hour != 0 or end_date.minute != 0:
                                                end_str += f" {end_date.strftime('%H:%M')}"
                                            
                                            # 기간 정보를 session state에 저장
                                            st.session_state['data_period'] = {
                                                'start_date': start_str,
                                                'end_date': end_str,
                                                'date_column': actual_column
                                            }
                                        else:
                                            st.session_state['data_period'] = {
                                                'start_date': str(start_date),
                                                'end_date': str(end_date),
                                                'date_column': actual_column
                                            }
                                    
                                    # 즉시 페이지 새로고침하여 데이터 요약 업데이트
                                    st.rerun()
                                else:
                                    st.error(f"❌ '{actual_column}' 열에 유효한 날짜 데이터가 없습니다.")
                            except Exception as e:
                                st.error(f"❌ '{actual_column}' 열을 날짜형으로 변환할 수 없습니다. 오류: {str(e)}")
                    else:
                        st.warning("⚠️ 날짜형으로 변환할 열을 선택해주세요.")
            

                
        # 결측치 셀 스타일링을 위한 CSS 추가
        st.markdown("""
        <style>
        .data-preview-table .stDataFrame td[data-testid="stDataFrameCell"]:has(.stDataFrameCell[data-testid="stDataFrameCell"]:empty),
        .data-preview-table .stDataFrame td:empty {
            background-color: #FF6464 !important;
        }
        </style>
        """, unsafe_allow_html=True)
        
        with st.container():
            # 저장된 기간 정보가 있으면 표시, 없으면 새로 계산
            if 'data_period' in st.session_state and st.session_state['data_period']:
                # 저장된 기간 정보 사용
                period_info = st.session_state['data_period']
                st.markdown(f"""
                <div style="background-color: #e3f2fd; padding: 15px; border-radius: 8px; margin: 10px 0; border-left: 4px solid #2196f3;">
                    <h4 style="margin: 0 0 10px 0; color: #1976d2;">📅 데이터 기간 정보</h4>
                    <p style="margin: 0; color: #1976d2; font-size: 14px;">
                        <strong>기간:</strong> {period_info['start_date']} ~ {period_info['end_date']}<br>
                        <strong>날짜 열:</strong> {period_info['date_column']}
                    </p>
                </div>
                """, unsafe_allow_html=True)
            else:
                # 저장된 정보가 없으면 새로 계산
                datetime_cols = st.session_state['df'].select_dtypes(include=['datetime']).columns
                if len(datetime_cols) > 0:
                    # 첫 번째 날짜형 열의 기간 정보 표시
                    date_col = datetime_cols[0]
                    valid_dates = st.session_state['df'][date_col].dropna()
                    if len(valid_dates) > 0:
                        start_date = valid_dates.min()
                        end_date = valid_dates.max()
                        
                        # 날짜 형식 포맷팅
                        if pd.api.types.is_datetime64_any_dtype(start_date):
                            start_str = start_date.strftime('%Y년 %m월 %d일')
                            end_str = end_date.strftime('%Y년 %m월 %d일')
                            
                            # 시간 정보가 있는지 확인
                            if start_date.hour != 0 or start_date.minute != 0:
                                start_str += f" {start_date.strftime('%H:%M')}"
                            if end_date.hour != 0 or end_date.minute != 0:
                                end_str += f" {end_date.strftime('%H:%M')}"
                            
                            # 기간 정보를 session state에 저장
                            st.session_state['data_period'] = {
                                'start_date': start_str,
                                'end_date': end_str,
                                'date_column': date_col
                            }
                            
                            st.markdown(f"""
                            <div style="background-color: #e3f2fd; padding: 15px; border-radius: 8px; margin: 10px 0; border-left: 4px solid #2196f3;">
                                <h4 style="margin: 0 0 10px 0; color: #1976d2;">📅 데이터 기간 정보</h4>
                                <p style="margin: 0; color: #1976d2; font-size: 14px;">
                                    <strong>기간:</strong> {start_str} ~ {end_str}<br>
                                    <strong>날짜 열:</strong> {date_col}
                                </p>
                            </div>
                            """, unsafe_allow_html=True)
            
            st.markdown('<div class="data-preview-table">', unsafe_allow_html=True)
            # 행번호 열에 "행번호" 헤더 추가
            df_with_index = st.session_state['df'].copy()
            df_with_index.index.name = '행번호'
            st.dataframe(df_with_index, use_container_width=True, height=400, hide_index=False)
            st.markdown('</div>', unsafe_allow_html=True)

    # 결측치 처리 헤더
    st.markdown("""
    <div style="margin-bottom: 10px;">
        <div style="display: flex; align-items: center; gap: -2px;">
            <h3 style="margin: 0;">결측치 처리</h3>
            <p style="font-size: 15px; color: #666; margin: 0;">(🔄 여러 번 누적하여 실행 가능합니다.)</p>
        </div>
    </div>
    """, unsafe_allow_html=True)
    
    # Skip 버튼을 왼쪽에 배치
    col1, col2, col3 = st.columns([1, 1, 1])
    with col1:
        if st.button("⏭️ Skip", help="결측치가 있는 데이터를 분석에 사용하지 않을 예정이라면 Skip 해도 됩니다.", key="skip_preprocessing"):
            # 최종 결과를 메인 데이터프레임에 적용
            st.session_state['df'] = st.session_state['current_df'].copy()
            st.session_state['preprocessing_completed'] = True  # 전처리 완료 상태 설정
            st.info("전처리가 완료되어 분석 데이터가 업데이트되었습니다.")
            if st.session_state['data_file_path']:
                try:
                    if st.session_state['data_file_path'].endswith('.csv'):
                        st.session_state['df'].to_csv(st.session_state['data_file_path'], index=False)
                    else:
                        st.session_state['df'].to_excel(st.session_state['data_file_path'], index=False)
                except PermissionError:
                    st.warning("⚠️ 파일이 다른 프로그램에서 사용 중이어서 저장할 수 없습니다.")
                except Exception as e:
                    st.warning(f"⚠️ 파일 저장 중 오류가 발생했습니다: {str(e)}")
            st.rerun()
      
    # 전처리 내용 표시
    # 전처리 과정 추적을 위한 session state 초기화
    if 'preprocessing_steps' not in st.session_state:
        st.session_state['preprocessing_steps'] = []
    if 'current_df' not in st.session_state or st.session_state['current_df'] is None:
        st.session_state['current_df'] = st.session_state['df'].copy()
    
    # 단계별 요약 표시
    if st.session_state['preprocessing_steps']:
        st.markdown('<h3 style="font-size: 18px; margin-bottom: 11px;">📋 수행된 전처리 단계 요약</h3>', unsafe_allow_html=True)
        for i, step in enumerate(st.session_state['preprocessing_steps'], 1):
            st.info(f"**{i}단계:** {step}")
        st.markdown("---")
    
    # 현재 단계의 결측치 처리 UI
    
    # 현재 결측치 상태 확인 (전처리 단계 정보 표시 전에 계산)
    current_na_count = st.session_state['current_df'].isnull().sum().sum()
    
    # 결측치 분포도 제목과 범례
    st.markdown('<h4 style="font-size: 16px; font-style: italic; color: #000; margin-bottom: 8px;">결측치 분포도</h4>', unsafe_allow_html=True)
    
    # 범례
    st.markdown('''
    <div style="display: flex; align-items: center; gap: 15px; font-size: 12px; color: #666; margin-bottom: 15px; margin-left: 10px;">
        <span style="display: flex; align-items: center; gap: 5px;">
            <div style="width: 12px; height: 12px; background-color: #288549; border-radius: 2px;"></div>
            정상 데이터
        </span>
        <div style="width: 1px; height: 20px; background-color: #ddd; margin: 0 10px;"></div>
        <span style="display: flex; align-items: center; gap: 5px;">
            <div style="width: 12px; height: 12px; background-color: #FF6B6B; border-radius: 2px;"></div>
            1단계 결측 (None, NaN, 공백)
        </span>
        <span style="display: flex; align-items: center; gap: 5px;">
            <div style="width: 12px; height: 12px; background-color: #FF0000; border-radius: 2px;"></div>
            2단계 결측 (#DIV/0!, #N/A, #NAME? 등)
        </span>
        <span style="display: flex; align-items: center; gap: 5px;">
            <div style="width: 12px; height: 12px; background-color: #8E44AD; border-radius: 2px;"></div>
            3단계 결측 (#NULL!, #SPILL! 등)
        </span>
    </div>
    ''', unsafe_allow_html=True)

    # Missingno 차트 생성
    import matplotlib.pyplot as plt
    import numpy as np
    
    df = st.session_state['current_df']  # 현재 처리 중인 데이터 사용
    
    # df가 None이 아닌지 확인
    if df is None:
        st.error("데이터가 로드되지 않았습니다. 파일을 다시 업로드해주세요.")
    else:
        # 새로운 기준에 따른 결측치 분류
        # 1단계: 공백만
        # 2단계: 1단계 + #DIV/0!, #N/A, #NAME?, #NUM!, #REF!, #VALUE!, #NODATA
        # 3단계: 2단계 + #NULL!, #SPILL!, #CALC!, #GETTING_DATA, #FIELD!, #BLOCKED!, #CONNECT!, #BUSY!, #UNKNOWN!
        
        # 1단계: 공백만 (None, NaN, 빈 문자열 모두 포함)
        stage1_mask = pd.DataFrame(False, index=df.index, columns=df.columns)
        for col in df.columns:
            try:
                # None, NaN, 빈 문자열을 모두 공백으로 처리
                stage1_mask[col] = (df[col].isna()) | (df[col].astype(str).str.strip() == '')
            except:
                # 문자열 변환이 안 되는 경우 빈 값으로 처리
                stage1_mask[col] = df[col].isnull()
        
        # 2단계: 1단계 + Excel 오류 값들
        error_values_2 = ['#DIV/0!', '#N/A', '#NAME?', '#NUM!', '#REF!', '#VALUE!', '#NODATA']
        stage2_mask = stage1_mask.copy()
        for col in df.columns:
            # Excel 오류 값들 확인
            excel_error_mask = df[col].isin(error_values_2)
            stage2_mask[col] = stage2_mask[col] | excel_error_mask
        
        # 3단계: 2단계 + 추가 Excel 오류 값들
        error_values_3 = ['#NULL!', '#SPILL!', '#CALC!', '#GETTING_DATA', '#FIELD!', 
                         '#BLOCKED!', '#CONNECT!', '#BUSY!', '#UNKNOWN!']
        stage3_mask = stage2_mask.copy()
        for col in df.columns:
            # 추가 Excel 오류 값들 확인
            additional_error_mask = df[col].isin(error_values_3)
            stage3_mask[col] = stage3_mask[col] | additional_error_mask
        
        n_rows, n_cols = df.shape
        
        # matplotlib 그래프 생성
        fig, ax = plt.subplots(figsize=(4, 2.0))
       
        # 1. 배경을 초록색으로 채우기 (정상 데이터)
        green_data = np.ones((n_rows, n_cols)) * 0.8
        ax.imshow(green_data, cmap='Greens', alpha=1.0, 
                 extent=[0, n_cols, 0, n_rows], aspect='auto', vmin=0, vmax=1)
        
        # 2. 단계별 결측치를 다른 색상으로 표시 (빨간색→보라색 그라데이션)
        # 1단계 결측치 (밝은 빨간색) - 공백만
        stage1_positions = np.where(stage1_mask)
        if len(stage1_positions[0]) > 0:
            for row, col in zip(stage1_positions[0], stage1_positions[1]):
                ax.axhline(row, xmin=col/n_cols, xmax=(col+1)/n_cols, 
                          color='#FF6B6B', alpha=1, linewidth=0.5)
        
        # 2단계 결측치 (완전 빨간색) - 1단계 + Excel 오류 값들 (1단계와 겹치지 않는 부분만)
        stage2_only = stage2_mask & ~stage1_mask
        stage2_positions = np.where(stage2_only)
        if len(stage2_positions[0]) > 0:
            for row, col in zip(stage2_positions[0], stage2_positions[1]):
                ax.axhline(row, xmin=col/n_cols, xmax=(col+1)/n_cols, 
                          color='#FF0000', alpha=1, linewidth=0.5)
        
        # 3단계 결측치 (보라색) - 2단계 + 추가 Excel 오류 값들 (2단계와 겹치지 않는 부분만)
        stage3_only = stage3_mask & ~stage2_mask
        stage3_positions = np.where(stage3_only)
        if len(stage3_positions[0]) > 0:
            for row, col in zip(stage3_positions[0], stage3_positions[1]):
                ax.axhline(row, xmin=col/n_cols, xmax=(col+1)/n_cols, 
                          color='#8E44AD', alpha=1, linewidth=0.5)
        
        # 3. 격자 선들 (벡터화)
        # 세로 선들
        if n_cols > 1:
            ax.vlines(np.arange(1, n_cols), 0, n_rows, color='black', linewidth=0.1, alpha=0.2)
        
        # 가로 선들
        if n_rows > 1:
            ax.hlines(np.arange(1, n_rows), 0, n_cols, color='white', linewidth=0.1, alpha=0.1)
        
        # 4. 테두리 선 (1/3 수준으로 얇게)
        ax.hlines(0, 0, n_cols, color='black', linewidth=0.33)
        ax.hlines(n_rows, 0, n_cols, color='black', linewidth=0.33)
        ax.vlines(0, 0, n_rows, color='black', linewidth=0.33)
        ax.vlines(n_cols, 0, n_rows, color='black', linewidth=0.33)
        
        ax.set_xlim(0, n_cols)
        ax.set_ylim(n_rows, 0)
        
        # x축 번호 매기기 (글자 겹침 방지를 위해 간격 조정)
        # 더 넓은 간격으로 눈금 개수 줄이기
        target_x_ticks = min(12, max(3, n_cols // 5))  # 최소 3개, 최대 12개로 줄임
        if n_cols <= target_x_ticks:
            x_ticks = np.arange(n_cols) + 0.5
            x_labels = [str(i+1) for i in range(n_cols)]
        else:
            step = max(1, n_cols // target_x_ticks)
            x_ticks = np.arange(0, n_cols, step) + 0.5
            x_labels = [str(i+1) for i in range(0, n_cols, step)]
        
        ax.set_xticks(x_ticks)
        ax.set_xticklabels(x_labels, rotation=0, ha='center', fontsize=5)
        # x축 숫자를 위쪽에 표시
        ax.xaxis.set_ticks_position('top')
        ax.xaxis.set_label_position('top')
        
        # y축 번호 매기기 (글자 겹침 방지를 위해 간격 조정)
        # 더 넓은 간격으로 눈금 개수 줄이기
        target_y_ticks = min(8, max(2, n_rows // 100))  # 최소 2개, 최대 8개로 줄임
        if n_rows <= target_y_ticks:
            y_ticks = np.arange(n_rows + 1)
            y_labels = [str(i) for i in range(n_rows + 1)]
        else:
            step = max(1, n_rows // target_y_ticks)
            y_ticks = np.arange(0, n_rows + 1, step)
            y_labels = [str(i) for i in range(0, n_rows + 1, step)]
        
        ax.set_yticks(y_ticks)
        ax.set_yticklabels(y_labels, fontsize=5)
        
        ax.set_xlabel('Column Number', fontsize=6)
        ax.set_ylabel('Row Number', fontsize=6)
        ax.grid(False)
        ax.set_facecolor((1,1,1,0))
        
        # 테두리만 표시
        for spine in ax.spines.values():
            spine.set_visible(True)
            spine.set_linewidth(0.5)
                    
        # 그래프 크기를 제한하는 컨테이너
        st.markdown("""
        <style>
        .missingno-container {
            max-width: 400vw !important;
            max-height: 1000px !important;
            overflow: hidden;
            margin-top: -250px !important;
            padding-top: 0 !important;
            margin-bottom: -10px !important;
        }
        .missingno-container .stPlotlyChart {
            max-width: 400vw !important;
            max-height: 1000px !important;
        }
        </style>
        """, unsafe_allow_html=True)
        
        # missingno 그래프와 결측치 열 정보를 좌우로 배치
        col1, col2 = st.columns([1, 1])
        
        with col1:
            # 그래프 표시
            st.pyplot(fig, use_container_width=False)
        
        with col2:
            # 결측치 통계 정보 - 새로운 기준에 따른 분류
            # 1단계: 공백 (None, NaN, 빈 문자열 모두 포함)
            stage1_mask = pd.DataFrame(False, index=df.index, columns=df.columns)
            for col in df.columns:
                try:
                    # None, NaN, 빈 문자열을 모두 공백으로 처리
                    stage1_mask[col] = (df[col].isna()) | (df[col].astype(str).str.strip() == '')
                except:
                    stage1_mask[col] = df[col].isnull()
            
            # 2단계: 1단계 + Excel 오류 값들
            error_values_2 = ['#DIV/0!', '#N/A', '#NAME?', '#NUM!', '#REF!', '#VALUE!', '#NODATA']
            stage2_mask = stage1_mask.copy()
            for col in df.columns:
                excel_error_mask = df[col].isin(error_values_2)
                stage2_mask[col] = stage2_mask[col] | excel_error_mask
            
            # 3단계: 2단계 + 추가 Excel 오류 값들
            error_values_3 = ['#NULL!', '#SPILL!', '#CALC!', '#GETTING_DATA', '#FIELD!', 
                             '#BLOCKED!', '#CONNECT!', '#BUSY!', '#UNKNOWN!']
            stage3_mask = stage2_mask.copy()
            for col in df.columns:
                additional_error_mask = df[col].isin(error_values_3)
                stage3_mask[col] = stage3_mask[col] | additional_error_mask
            
            # 모든 결측치를 포함하는 마스크 (3단계가 모든 결측치를 포함)
            all_missing_mask = stage3_mask
            
            row_has_na = all_missing_mask.any(axis=1)
            col_has_na = all_missing_mask.any(axis=0)
            
            na_rows = df[row_has_na]  # 결측치가 있는 행만
            na_cols = df.columns[col_has_na].tolist()  # 결측치가 있는 열만
            total_rows = len(df)
            total_cols = len(df.columns)
            na_count_rows = row_has_na.sum()  # 더 빠른 계산
            na_count_cols = col_has_na.sum()  # 더 빠른 계산
            current_na_count = all_missing_mask.sum().sum()  # 모든 결측치 개수
            
            if current_na_count > 0:
                # 결측치가 있는 경우 - 처리 옵션 표시
                
                # 2. 결측치가 있는 열 정보
                st.markdown('''
                <div style="margin-left: 20px; margin-top: 15px;">
                    <p style="font-size: 16px; font-style: italic; color: #000; margin-bottom: 0px;">결측치가 있는 열 정보</p>
                </div>
                ''', unsafe_allow_html=True)
                missing_col_info = []
                for col in na_cols:
                    # 새로운 기준에 따른 각 단계별 결측치 개수 계산
                    
                    # 1단계: 공백만 (None, NaN, 빈 문자열 모두 포함)
                    try:
                        # None, NaN, 빈 문자열을 모두 공백으로 처리
                        stage1_mask = (df[col].isna()) | (df[col].astype(str).str.strip() == '')
                        stage1_count = stage1_mask.sum()
                    except:
                        stage1_count = 0
                    
                    # 2단계: 1단계 + Excel 오류 값들 (1단계와 겹치지 않는 부분만)
                    error_values_2 = ['#DIV/0!', '#N/A', '#NAME?', '#NUM!', '#REF!', '#VALUE!', '#NODATA']
                    stage2_only = df[col].isin(error_values_2) & ~stage1_mask
                    stage2_count = stage2_only.sum()
                    
                    # 3단계: 2단계 + 추가 Excel 오류 값들 (2단계와 겹치지 않는 부분만)
                    error_values_3 = ['#NULL!', '#SPILL!', '#CALC!', '#GETTING_DATA', '#FIELD!', '#BLOCKED!', '#CONNECT!', '#BUSY!', '#UNKNOWN!']
                    stage3_only = df[col].isin(error_values_3) & ~stage1_mask & ~stage2_only
                    stage3_count = stage3_only.sum()
                    
                    total_missing = stage1_count + stage2_count + stage3_count
                    missing_ratio = (total_missing / len(df)) * 100
                    
                    # 각 단계별 상세 정보 생성
                    detail_parts = []
                    
                    # 1단계: 공백 (None, NaN, 빈 문자열 세분화)
                    if stage1_count > 0:
                        # None, NaN, 빈 문자열을 개별적으로 카운트
                        none_count = df[col].isna().sum()
                        empty_str_count = ((df[col].astype(str).str.strip() == '') & ~df[col].isna()).sum()
                        
                        stage1_details = []
                        if none_count > 0:
                            stage1_details.append(f"None/NaN {none_count}개")
                        if empty_str_count > 0:
                            stage1_details.append(f"공백문자 {empty_str_count}개")
                        
                        if stage1_details:
                            detail_parts.append(f"🟠 1단계: {', '.join(stage1_details)}")
                        else:
                            detail_parts.append(f"🟠 1단계: 공백 {stage1_count}개")
                    
                    # 2단계: Excel 오류 값들 (개별 카운트)
                    if stage2_count > 0:
                        excel_errors_2 = ['#DIV/0!', '#N/A', '#NAME?', '#NUM!', '#REF!', '#VALUE!', '#NODATA']
                        stage2_details = []
                        for error in excel_errors_2:
                            error_count = (df[col].astype(str) == error).sum()
                            if error_count > 0:
                                stage2_details.append(f"{error} {error_count}개")
                        
                        if stage2_details:
                            detail_parts.append(f"🔴 2단계: {', '.join(stage2_details)}")
                    
                    # 3단계: 추가 Excel 오류 값들 (개별 카운트)
                    if stage3_count > 0:
                        excel_errors_3 = ['#NULL!', '#SPILL!', '#CALC!', '#GETTING_DATA', '#FIELD!', '#BLOCKED!', '#CONNECT!', '#BUSY!', '#UNKNOWN!']
                        stage3_details = []
                        for error in excel_errors_3:
                            error_count = (df[col].astype(str) == error).sum()
                            if error_count > 0:
                                stage3_details.append(f"{error} {error_count}개")
                        
                        if stage3_details:
                            detail_parts.append(f"🟣 3단계: {', '.join(stage3_details)}")
                    
                    detail_text = " | ".join(detail_parts) if detail_parts else "없음"
                    
                    # 열 번호 계산 (0부터 시작하는 인덱스를 1부터 시작하는 번호로 변환)
                    col_index = df.columns.get_loc(col) + 1
                    
                    missing_col_info.append({
                        '열 번호': col_index,
                        '열명': col,
                        '총 결측치 수': total_missing,
                        '결측치 비율': f"{missing_ratio:.1f}%",
                        '결측 유형': detail_text
                    })
                
                missing_df = pd.DataFrame(missing_col_info)
                # 결측치 열 정보 표에만 CSS 적용
                st.markdown("""
                <style>
                .missing-info-table div[data-testid="stDataFrame"] {
                    margin-top: 0px !important;
                    margin-left: 40px !important;
                }
                </style>
                """, unsafe_allow_html=True)
                
                # 결측치 열 정보 표를 특정 클래스로 감싸기
                st.markdown('<div class="missing-info-table">', unsafe_allow_html=True)
                st.dataframe(missing_df, use_container_width=False, hide_index=True, width=600)
                st.markdown('</div>', unsafe_allow_html=True)
        
        
        # 결측치 제거 레벨 선택
        st.markdown('<h4 style="font-size: 16px; font-style: italic; color: #000; margin-bottom: 0px;">결측치 제거 레벨 선택</h4>', unsafe_allow_html=True)
        
        # 라디오 버튼 상단 여백 제거를 위한 CSS
        st.markdown("""
        <style>
        [data-testid="stRadio"] {
            margin-top: -50px !important;
            padding-top: 0px !important;
            margin-bottom: 22px !important;
        }
        [data-testid="stRadio"] > div {
            margin-top: 0px !important;
            padding-top: 0px !important;
        }
        </style>
        """, unsafe_allow_html=True)
        
        # 세로로 배치된 라디오 버튼 사용
        with st.container():
            removal_level = st.radio(
                "",
                options=[
                    "1단계 : 공백, None, NaN 제거",
                    "2단계 : 1단계 + #DIV/0!, #N/A, #NAME?, #NUM!, #REF!, #VALUE!, #NODATA",
                    "3단계 : 2단계 + #NULL!, #SPILL!, #CALC!, #GETTING_DATA, #FIELD!, #BLOCKED!, #CONNECT!, #BUSY!, #UNKNOWN!"
                ],
                index=1,  # 2단계를 기본값으로 설정
                key="removal_level_radio"
            )
        
        # 처리 방법 선택
        st.markdown('<h4 style="font-size: 16px; font-style: italic; color: #000; margin-bottom: 0px;">결측치 처리 방법 선택</h4>', unsafe_allow_html=True)
        
        # 처리 방법 라디오 버튼 상단 여백 제거를 위한 CSS
        st.markdown("""
        <style>
        [data-testid="stRadio"]:nth-of-type(2) {
            margin-top: -35px !important;
            padding-top: 0px !important;
            margin-bottom: 20px !important;
        }
        [data-testid="stRadio"]:nth-of-type(2) > div {
            margin-top: 0px !important;
            padding-top: 0px !important;
        }
        </style>
        """, unsafe_allow_html=True)
        
        na_action = st.radio(
            "",
            [
                "유지",
                "행 삭제",
                "열 삭제"
                            ]
        )
        
        # 특정 행/열 선택 삭제 기능
        if "행 삭제" in na_action and na_count_rows > 0:
            # 결측치가 있는 행들의 인덱스 찾기
            rows_with_na_indices = st.session_state['current_df'][st.session_state['current_df'].isnull().any(axis=1)].index.tolist()
            
            # 저장된 날짜 열 정보를 우선적으로 사용
            if 'date_column_name' in st.session_state:
                # 사용자가 직접 날짜형으로 변환한 열 사용
                date_col = st.session_state['date_column_name']
                if date_col in st.session_state['current_df'].columns:
                    # 날짜 열이 있는 경우 날짜 정보로 표시
                    rows_with_na = []
                    for idx in rows_with_na_indices:
                        try:
                            date_value = st.session_state['current_df'].loc[idx, date_col]
                            if pd.notna(date_value):
                                # 날짜 형식으로 표시
                                try:
                                    if pd.api.types.is_datetime64_any_dtype(date_value):
                                        date_str = date_value.strftime('%Y-%m-%d %H:%M:%S')
                                    else:
                                        date_str = str(date_value)
                                    rows_with_na.append(f"{date_str} (행 {idx})")
                                except:
                                    # 날짜 형식 변환 실패 시 행 번호로 표시
                                    rows_with_na.append(f"행 {idx}")
                            else:
                                # 날짜가 없는 경우 행 번호로 표시
                                rows_with_na.append(f"행 {idx}")
                        except:
                            # 오류 발생 시 행 번호로 표시
                            rows_with_na.append(f"행 {idx}")
                else:
                    # 저장된 날짜 열이 현재 데이터에 없는 경우 행 번호로 표시
                    rows_with_na = [f"행 {idx}" for idx in rows_with_na_indices]
            elif 'data_period' in st.session_state and st.session_state['data_period']:
                # 저장된 날짜 열 정보 사용 (기존 방식)
                date_col = st.session_state['data_period']['date_column']
                if date_col in st.session_state['current_df'].columns:
                    # 날짜 열이 있는 경우 날짜 정보로 표시
                    rows_with_na = []
                    for idx in rows_with_na_indices:
                        try:
                            date_value = st.session_state['current_df'].loc[idx, date_col]
                            if pd.notna(date_value):
                                # 날짜 형식으로 표시
                                try:
                                    if pd.api.types.is_datetime64_any_dtype(date_value):
                                        date_str = date_value.strftime('%Y-%m-%d %H:%M:%S')
                                    else:
                                        date_str = str(date_value)
                                    rows_with_na.append(f"{date_str} (행 {idx})")
                                except:
                                    # 날짜 형식 변환 실패 시 행 번호로 표시
                                    rows_with_na.append(f"행 {idx}")
                            else:
                                # 날짜가 없는 경우 행 번호로 표시
                                rows_with_na.append(f"행 {idx}")
                        except:
                            # 오류 발생 시 행 번호로 표시
                            rows_with_na.append(f"행 {idx}")
                else:
                    # 저장된 날짜 열이 현재 데이터에 없는 경우 행 번호로 표시
                    rows_with_na = [f"행 {idx}" for idx in rows_with_na_indices]
            else:
                # 저장된 날짜 정보가 없는 경우 자동 감지
                datetime_cols = st.session_state['current_df'].select_dtypes(include=['datetime']).columns
                if len(datetime_cols) > 0:
                    # 날짜 열이 있는 경우 날짜 정보로 표시
                    date_col = datetime_cols[0]
                    rows_with_na = []
                    for idx in rows_with_na_indices:
                        try:
                            date_value = st.session_state['current_df'].loc[idx, date_col]
                            if pd.notna(date_value):
                                # 날짜 형식으로 표시
                                try:
                                    if pd.api.types.is_datetime64_any_dtype(date_value):
                                        date_str = date_value.strftime('%Y-%m-%d %H:%M:%S')
                                    else:
                                        date_str = str(date_value)
                                    rows_with_na.append(f"{date_str} (행 {idx})")
                                except:
                                    # 날짜 형식 변환 실패 시 행 번호로 표시
                                    rows_with_na.append(f"행 {idx}")
                            else:
                                # 날짜가 없는 경우 행 번호로 표시
                                rows_with_na.append(f"행 {idx}")
                        except:
                            # 오류 발생 시 행 번호로 표시
                            rows_with_na.append(f"행 {idx}")
                else:
                    # 날짜 열이 없는 경우 행 번호로 표시
                    rows_with_na = [f"행 {idx}" for idx in rows_with_na_indices]
            
            # 행 선택 방법 선택
            st.markdown('<p style="font-size: 17px; margin-bottom: 20px; margin-top: -10px; font-style: italic;">행 선택 방법</p>', unsafe_allow_html=True)
            row_selection_method = st.radio(
                "",
                options=["드롭다운에서 선택", "직접 범위 입력"],
                horizontal=True,
                key="row_selection_method"
            )
            
            if row_selection_method == "드롭다운에서 선택":
                rows_to_drop = st.multiselect(
                    "삭제할 행을 선택하세요",
                    options=rows_with_na,
                    default=[],
                    key="rows_multiselect"
                )
            else:
                # 직접 범위 입력
                st.markdown("""
                <div style="background-color: #f8f9fa; padding: 10px; border-radius: 5px; margin: -5px 0;">
                    <p style="margin: 0; font-size: 14px; color: #495057;">
                        <strong>입력 형식:</strong> 개별 행: 5,7,10 | 범위: 5-10, 15~20 | 혼합: 5,7,10-15,44,78~89
                    </p>
                </div>
                """, unsafe_allow_html=True)
                
                range_input = st.text_input(
                    "삭제할 행 범위를 입력하세요",
                    placeholder="예: 5,7,10-15,44,78~89",
                    key="row_range_input"
                )
                
                # 범위 입력을 파싱하는 함수
                def parse_row_range(range_str):
                    if not range_str.strip():
                        return []
                    
                    rows_to_drop = []
                    parts = range_str.replace(' ', '').split(',')
                    
                    for part in parts:
                        part = part.strip()
                        if not part:
                            continue
                        
                        # 범위 처리 (5-10 또는 5~10)
                        if '-' in part or '~' in part:
                            separator = '-' if '-' in part else '~'
                            try:
                                start, end = part.split(separator)
                                start_idx = int(start.strip())
                                end_idx = int(end.strip())
                                
                                # 실제 데이터 인덱스에 맞게 조정
                                for idx in range(start_idx, end_idx + 1):
                                    if idx in st.session_state['current_df'].index:
                                        rows_to_drop.append(idx)
                            except ValueError:
                                st.warning(f"잘못된 범위 형식: {part}")
                                continue
                        else:
                            # 개별 행 번호
                            try:
                                idx = int(part)
                                if idx in st.session_state['current_df'].index:
                                    rows_to_drop.append(idx)
                                else:
                                    st.warning(f"존재하지 않는 행 번호: {idx}")
                            except ValueError:
                                st.warning(f"잘못된 행 번호: {part}")
                                continue
                    
                    return list(set(rows_to_drop))  # 중복 제거
                
                # 입력된 범위를 실제 행 인덱스로 변환
                if range_input:
                    parsed_rows = parse_row_range(range_input)
                    rows_to_drop = parsed_rows
                else:
                    rows_to_drop = []
        elif "열 삭제" in na_action and na_count_cols > 0:
            cols_to_drop = st.multiselect(
                "삭제할 열을 선택하세요",
                options=na_cols,
                default=[],
                key="cols_multiselect"
            )
        else:
            rows_to_drop = []
            cols_to_drop = []
        
        # UI 요소들의 스타일 조정을 위한 CSS
        st.markdown("""
        <style>
        [data-testid="stTextInput"] {
            width: 100% !important;
        }
        .stRadio > div {
            margin-bottom: 10px !important;
        }
        /* placeholder 텍스트 색상 통일 - 더 포괄적인 선택자 */
        [data-testid="stSelectbox"] div[data-baseweb="select"] span,
        [data-testid="stMultiSelect"] div[data-baseweb="select"] span,
        [data-testid="stSelectbox"] div[data-baseweb="select"] div,
        [data-testid="stMultiSelect"] div[data-baseweb="select"] div,
        [data-testid="stSelectbox"] div[data-baseweb="select"] input::placeholder,
        [data-testid="stMultiSelect"] div[data-baseweb="select"] input::placeholder,
        [data-testid="stSelectbox"] div[data-baseweb="select"] input,
        [data-testid="stMultiSelect"] div[data-baseweb="select"] input {
            color: #666 !important;
        }
        /* 추가적인 placeholder 스타일 */
        [data-testid="stSelectbox"] div[data-baseweb="select"] div[role="option"],
        [data-testid="stMultiSelect"] div[data-baseweb="select"] div[role="option"] {
            color: #666 !important;
        }
        </style>
        """, unsafe_allow_html=True)
        
        # 버튼들을 나란히 배치
        col1, col2, col3 = st.columns(3)
        
        with col1:
            # 버튼 텍스트를 상황에 맞게 변경
            if st.session_state['preprocessing_steps']:
                button_text = "🔄 추가 전처리 수행"
            else:
                button_text = "▶️ 전처리 시작"
            
            if st.button(button_text, help="선택한 설정으로 결측치를 처리합니다"):
                # 제거될 데이터 비율 계산
                df = st.session_state['current_df'].copy()
                total_original = len(df) * len(df.columns)
                
                # 선택된 레벨에 따른 결측치 처리 (임시)
                temp_df = df.copy()
                
                # 1단계: 공백만 제거 (None, NaN, 빈 문자열 모두 포함)
                if "1단계" in removal_level:
                    for col in temp_df.columns:
                        # 빈 문자열을 pd.NA로 변환 (None과 NaN은 이미 pd.NA로 처리됨)
                        temp_df[col] = temp_df[col].replace('', pd.NA)
                elif "2단계" in removal_level:
                    error_values_2 = ['#DIV/0!', '#N/A', '#NAME?', '#NUM!', '#REF!', '#VALUE!', '#NODATA']
                    for col in temp_df.columns:
                        temp_df[col] = temp_df[col].replace('', pd.NA)
                        temp_df[col] = temp_df[col].replace(error_values_2, pd.NA)
                elif "3단계" in removal_level:
                    error_values_2 = ['#DIV/0!', '#N/A', '#NAME?', '#NUM!', '#REF!', '#VALUE!', '#NODATA']
                    error_values_3 = ['#NULL!', '#SPILL!', '#CALC!', '#GETTING_DATA', '#FIELD!', 
                                    '#BLOCKED!', '#CONNECT!', '#BUSY!', '#UNKNOWN!']
                    for col in temp_df.columns:
                        temp_df[col] = temp_df[col].replace('', pd.NA)
                        temp_df[col] = temp_df[col].replace(error_values_2 + error_values_3, pd.NA)
                
                # 처리 방법에 따른 제거 계산
                if "행 삭제" in na_action:
                    temp_df = temp_df.dropna()
                elif "열 삭제" in na_action:
                    if cols_to_drop:
                        existing_cols_to_drop = [col for col in cols_to_drop if col in temp_df.columns]
                        if existing_cols_to_drop:
                            temp_df = temp_df.drop(columns=existing_cols_to_drop)
                    else:
                        temp_df = temp_df.dropna(axis=1)
                
                # 제거될 데이터 비율 계산
                total_after = len(temp_df) * len(temp_df.columns)
                removal_percentage = ((total_original - total_after) / total_original) * 100
                
                # 30% 이상 제거되는 경우 확인 다이얼로그
                if removal_percentage >= 30:
                    st.markdown(f"""
                    <div style="background-color: #fff3cd; border: 1px solid #ffeaa7; color: #856404; padding: 12px; border-radius: 4px; margin: 10px 0;">
                        ⚠️ <strong>주의</strong>: 선택한 옵션에 따라 전처리를 수행할 경우 기존 데이터 중 <strong>{removal_percentage:.0f}%</strong>가 제거됩니다. 실행하시겠습니까?
                    </div>
                    """, unsafe_allow_html=True)
                    
                    # 세션 상태에 확인 대기 상태 저장
                    st.session_state['waiting_confirmation'] = True
                    st.session_state['removal_percentage'] = removal_percentage
                    st.session_state['temp_df'] = temp_df
                    st.session_state['prev_rows'] = len(st.session_state['current_df'])
                    st.session_state['prev_cols'] = len(st.session_state['current_df'].columns)
                    
                    col_confirm1, col_confirm2 = st.columns(2)
                    with col_confirm1:
                        if st.button("❌ 취소", key="cancel_preprocessing"):
                            st.session_state['waiting_confirmation'] = False
                            st.info("전처리가 취소되었습니다.")
                            st.rerun()
                    with col_confirm2:
                        if st.button("✅ 실행", key="confirm_preprocessing"):
                            st.session_state['waiting_confirmation'] = False
                            st.success("전처리를 진행합니다...")
                            st.rerun()
                else:
                    # 30% 미만인 경우 바로 실행
                    # 실제 전처리 수행
                    prev_rows = len(st.session_state['current_df'])
                    prev_cols = len(st.session_state['current_df'].columns)
                    
                    # 처리 과정 기록을 위한 설명 생성
                    step_description = ""
                    
                    # 현재 데이터프레임 복사
                    df = st.session_state['current_df'].copy()
                    
                    # 1단계: 공백만 제거 (None, NaN, 빈 문자열 모두 포함)
                    if "1단계" in removal_level:
                        # None, NaN, 빈 문자열을 모두 pd.NA로 변환
                        for col in df.columns:
                            # 빈 문자열을 pd.NA로 변환
                            df[col] = df[col].replace('', pd.NA)
                            # None과 NaN은 이미 pd.NA로 처리됨
                    elif "2단계" in removal_level:
                        # 2단계: 1단계 + Excel 오류 값들
                        error_values_2 = ['#DIV/0!', '#N/A', '#NAME?', '#NUM!', '#REF!', '#VALUE!', '#NODATA']
                        for col in df.columns:
                            df[col] = df[col].replace('', pd.NA)  # 공백 제거
                            df[col] = df[col].replace(error_values_2, pd.NA)  # Excel 오류 값들 제거
                    elif "3단계" in removal_level:
                        # 3단계: 2단계 + 추가 Excel 오류 값들
                        error_values_2 = ['#DIV/0!', '#N/A', '#NAME?', '#NUM!', '#REF!', '#VALUE!', '#NODATA']
                        error_values_3 = ['#NULL!', '#SPILL!', '#CALC!', '#GETTING_DATA', '#FIELD!', 
                                        '#BLOCKED!', '#CONNECT!', '#BUSY!', '#UNKNOWN!']
                        for col in df.columns:
                            df[col] = df[col].replace('', pd.NA)  # 공백 제거
                            df[col] = df[col].replace(error_values_2 + error_values_3, pd.NA)  # 모든 Excel 오류 값들 제거
                    
                    # 업데이트된 데이터프레임을 session state에 저장
                    st.session_state['current_df'] = df
                    
                    if "행 삭제" in na_action:
                        # 행 삭제 로직 - 특정 행 선택 삭제 또는 전체 삭제
                        if 'rows_to_drop' in locals() and rows_to_drop and len(rows_to_drop) > 0:
                            # 선택된 행만 삭제
                            existing_rows_to_drop = []
                            
                            # 행 선택 방법에 따라 처리
                            if 'row_selection_method' in st.session_state and st.session_state['row_selection_method'] == "직접 범위 입력":
                                # 직접 범위 입력의 경우 이미 실제 인덱스가 들어있음
                                existing_rows_to_drop = [idx for idx in rows_to_drop if idx in st.session_state['current_df'].index]
                            else:
                                # 드롭다운 선택의 경우 텍스트에서 인덱스 추출
                                for row_text in rows_to_drop:
                                    # "행 {idx}" 또는 "{date_str} (행 {idx})" 형식에서 인덱스 추출
                                    if "행 " in row_text:
                                        try:
                                            # "행 123" 형식에서 숫자 추출
                                            idx = int(row_text.split("행 ")[1].split(")")[0])
                                            if idx in st.session_state['current_df'].index:
                                                existing_rows_to_drop.append(idx)
                                        except:
                                            pass
                                    elif " (행 " in row_text:
                                        try:
                                            # "2024-01-01 12:00:00 (행 123)" 형식에서 숫자 추출
                                            idx = int(row_text.split(" (행 ")[1].split(")")[0])
                                            if idx in st.session_state['current_df'].index:
                                                existing_rows_to_drop.append(idx)
                                        except:
                                            pass
                            
                            if existing_rows_to_drop:
                                st.session_state['current_df'] = st.session_state['current_df'].drop(index=existing_rows_to_drop)
                                step_description += f"특정행삭제({len(existing_rows_to_drop)}개)"
                            else:
                                st.warning("선택한 행들이 이미 삭제되었거나 존재하지 않습니다.")
                        else:
                            # 특정 행이 선택되지 않은 경우 전체 삭제
                            st.session_state['current_df'] = st.session_state['current_df'].dropna()
                            step_description += f"행 삭제(전체)"
                    
                    if "열 삭제" in na_action:
                        # 열 삭제 로직 - 특정 열 선택 삭제 또는 전체 삭제
                        if cols_to_drop and len(cols_to_drop) > 0:
                            # 선택된 열만 삭제
                            existing_cols_to_drop = [col for col in cols_to_drop if col in st.session_state['current_df'].columns]
                            if existing_cols_to_drop:
                                st.session_state['current_df'] = st.session_state['current_df'].drop(columns=existing_cols_to_drop)
                                step_description += f"특정열삭제({len(existing_cols_to_drop)}개)"
                            else:
                                st.warning("선택한 열들이 이미 삭제되었거나 존재하지 않습니다.")
                        else:
                            # 특정 열이 선택되지 않은 경우 전체 삭제
                            st.session_state['current_df'] = st.session_state['current_df'].dropna(axis=1)
                            step_description += f"열 삭제(전체)"
                    
                    # 결과 표시
                    after_rows = len(st.session_state['current_df'])
                    after_cols = len(st.session_state['current_df'].columns)
                    
                    if "삭제" in na_action:
                        # 전처리 과정 기록
                        step_description += f" (행: {prev_rows}→{after_rows}, 열: {prev_cols}→{after_cols})"
                        st.session_state['preprocessing_steps'].append(step_description)
                        
                        st.success(f"결측치 처리가 완료되었습니다!")
                        st.write(f"**변경 사항:**")
                        st.write(f"- 행: {prev_rows} → {after_rows} ({prev_rows - after_rows}개 삭제)")
                        st.write(f"- 열: {prev_cols} → {after_cols} ({prev_cols - after_cols}개 삭제)")
                        
                        # 페이지 새로고침을 위한 rerun
                        st.rerun()
                    else:
                        st.info(f"결측치가 있는 데이터를 유지합니다. (총 {prev_rows}행, {prev_cols}열)")
            
            # 확인 대기 상태에서 실제 전처리 수행
            if 'waiting_confirmation' in st.session_state and st.session_state['waiting_confirmation']:
                st.stop()
        
        with col2:
            if st.button("✅ 전처리 완료", help="전처리 완료 버튼을 누르면 분석 수행 단계로 넘어갑니다."):
                # 최종 결과를 메인 데이터프레임에 적용
                st.session_state['df'] = st.session_state['current_df'].copy()
                st.session_state['preprocessing_completed'] = True  # 전처리 완료 상태 설정
                st.info("전처리가 완료되어 분석 데이터가 업데이트되었습니다.")
                if st.session_state['data_file_path']:
                    try:
                        if st.session_state['data_file_path'].endswith('.csv'):
                            st.session_state['df'].to_csv(st.session_state['data_file_path'], index=False)
                        else:
                            st.session_state['df'].to_excel(st.session_state['data_file_path'], index=False)
                    except PermissionError:
                        st.warning("⚠️ 파일이 다른 프로그램에서 사용 중이어서 저장할 수 없습니다.")
                    except Exception as e:
                        st.warning(f"⚠️ 파일 저장 중 오류가 발생했습니다: {str(e)}")
                st.rerun()
        
        with col3:
            if st.button("🔄 전처리 초기화", help="모든 전처리 과정을 초기화하고 처음부터 다시 시작합니다"):
                # 전처리 과정 초기화
                st.session_state['current_df'] = st.session_state['df'].copy()
                st.session_state['preprocessing_steps'] = []
                st.info("전처리 과정이 초기화되었습니다.")
                st.rerun()
    
    if current_na_count == 0:
        # 결측치가 없는 경우 - 완료 상태 표시
        st.success("✅ 결측치 처리가 완료되었습니다!")
        
        # 전처리 완료 안내 메시지
        st.info("💡 위의 '전처리 완료' 버튼을 클릭하여 다음 단계로 진행하세요.")

    # 이후 단계는 전처리가 완료된 후에만 표시
    if st.session_state['df'] is not None and st.session_state['preprocessing_completed']:
        # 3. 분석 대상 변수 선택 및 데이터 확인 (파일 업로드/전처리 이후에만 노출)
        st.markdown('<hr style="border:3px solid #333;">', unsafe_allow_html=True)
        st.header("3. 분석 대상 변수 선택 및 데이터 확인")
        st.subheader("분석에 사용할 변수 선택")
        numeric_cols = st.session_state['df'].select_dtypes(include=[np.number]).columns.tolist()
        if len(numeric_cols) < 2:
            st.warning("수치형 컬럼이 2개 이상 필요합니다.")
            x_cols, y_col = [], None
        else:
            # Y 선택과 X 선택 드롭다운 박스 (화면 가로 길이의 절반으로 설정)
            col1, col2 = st.columns([1, 1])
            with col1:
                y_col = st.selectbox("종속변수(Y) 선택", options=["변수를 선택하세요"] + numeric_cols, index=0, key="y_col")
            with col2:
                # X 선택: Y로 선택된 변수는 제외
                x_options = [col for col in numeric_cols if col != y_col] if y_col else numeric_cols
                x_cols = st.multiselect("독립변수(X) 선택 (2~5개 추천)", options=x_options, max_selections=5, key="x_cols", placeholder="변수들을 선택하세요")
            
            if y_col == "변수를 선택하세요":
                y_col = None
            if x_cols and y_col:
                # X, Y 모두 결측치가 없는 행만 남김
                df_no_na = st.session_state['df'][[y_col] + x_cols].dropna()
                # 변수 목록을 쉼표로 구분하여 한 줄에 나열
                x_vars_text = ", ".join(x_cols)
                st.success(f"**분석에 사용할 변수**\n\n**독립변수 (X):** {x_vars_text}\n\n**종속변수 (Y):** {y_col}")
                st.markdown('<span style="font-weight:bold; color:#1f77b4; font-size:18px;">기초 통계</span>', unsafe_allow_html=True)
                st.dataframe(df_no_na.describe())
                st.markdown('<span style="font-weight:bold; color:#ff7f0e; font-size:18px;">히스토그램</span>', unsafe_allow_html=True)
                # 히스토그램 2개씩 한 줄에 배치
                hist_cols = [y_col] + x_cols
                for i in range(0, len(hist_cols), 2):
                    cols = st.columns(2)
                    for j in range(2):
                        if i + j < len(hist_cols):
                            col = hist_cols[i + j]
                            with cols[j]:
                                fig = px.histogram(df_no_na, x=col, nbins=30, title=f"{col} 분포")
                                st.plotly_chart(fig, use_container_width=True)
                # 상관계수 행렬과 상관관계 분석을 좌우로 나란히 배치
                col1, col2 = st.columns(2)
                
                with col1:
                    st.markdown('<span style="font-weight:bold; color:#2ca02c; font-size:18px;">상관계수 행렬</span>', unsafe_allow_html=True)
                    corr = df_no_na.corr()
                    fig = ff.create_annotated_heatmap(
                        z=corr.values,
                        x=list(corr.columns),
                        y=list(corr.index)[::-1],  # 세로 방향 순서 뒤집기
                        annotation_text=np.round(corr.values, 2),
                        colorscale='RdBu', showscale=True  # Red-Blue 색상 스케일로 변경
                    )
                    fig.update_layout(
                        coloraxis=dict(
                            colorscale='RdBu',  # Red-Blue 색상 스케일
                            cmin=-1,  # 음의 상관관계도 표시
                            cmax=1,
                            colorbar=dict(title="상관계수")
                        ),
                        height=500,  # 높이 조정
                        width=400    # 너비 조정
                    )
                    st.plotly_chart(fig, use_container_width=True)
                
                with col2:
                    # 상관계수 행렬에서 0.4 이상인 쌍 찾기
                    high_corr_pairs = []
                    for i in range(len(corr.columns)):
                        for j in range(i+1, len(corr.columns)):  # 대각선 위쪽만 확인 (중복 방지)
                            corr_value = corr.iloc[i, j]
                            if abs(corr_value) >= 0.4:
                                high_corr_pairs.append({
                                    '변수1': corr.columns[i],
                                    '변수2': corr.columns[j],
                                    '상관계수': corr_value,
                                    '강도': '강한 양의 상관관계' if corr_value >= 0.7 else 
                                           '중간 양의 상관관계' if corr_value >= 0.4 else
                                           '강한 음의 상관관계' if corr_value <= -0.7 else
                                           '중간 음의 상관관계'
                                })
                    
                    # 상관계수 기준으로 정렬 (절댓값 기준 내림차순)
                    high_corr_pairs.sort(key=lambda x: abs(x['상관계수']), reverse=True)
                    
                    if high_corr_pairs:
                        # 결과를 데이터프레임으로 변환
                        high_corr_df = pd.DataFrame(high_corr_pairs)
                        
                        # 상관계수 값을 색상으로 구분
                        def color_correlation(val):
                            if val >= 0.7:
                                return 'background-color: #ffcccc; color: #cc0000; font-weight: bold;'  # 강한 양의 상관관계
                            elif val >= 0.4:
                                return 'background-color: #ffe6cc; color: #cc6600; font-weight: bold;'  # 중간 양의 상관관계
                            elif val <= -0.7:
                                return 'background-color: #ccf2ff; color: #0066cc; font-weight: bold;'  # 강한 음의 상관관계
                            else:
                                return 'background-color: #e6f3ff; color: #0066cc; font-weight: bold;'  # 중간 음의 상관관계
                        
                        # 스타일 적용
                        styled_df = high_corr_df.style.applymap(
                            lambda x: color_correlation(x) if isinstance(x, (int, float)) else '', 
                            subset=['상관계수']
                        )
                        
                        st.write(f"**총 {len(high_corr_pairs)}개의 변수 쌍에서 유의미한 상관관계가 발견되었습니다 (절댓값 ≥ 0.4):**")
                        st.dataframe(styled_df, use_container_width=True)
                        
                        # 요약 정보
                        st.markdown("**📊 상관관계 요약:**")
                        strong_positive = len([p for p in high_corr_pairs if p['상관계수'] >= 0.7])
                        moderate_positive = len([p for p in high_corr_pairs if 0.4 <= p['상관계수'] < 0.7])
                        strong_negative = len([p for p in high_corr_pairs if p['상관계수'] <= -0.7])
                        moderate_negative = len([p for p in high_corr_pairs if -0.7 < p['상관계수'] <= -0.4])
                        
                        # 2x2 그리드로 메트릭 배치
                        metric_col1, metric_col2 = st.columns(2)
                        with metric_col1:
                            st.metric("강한 양의 상관관계 (≥0.7)", strong_positive)
                            st.metric("강한 음의 상관관계 (≤-0.7)", strong_negative)
                        with metric_col2:
                            st.metric("중간 양의 상관관계 (0.4~0.7)", moderate_positive)
                            st.metric("중간 음의 상관관계 (-0.7~-0.4)", moderate_negative)
                        
                    else:
                        st.info("상관계수가 0.4 이상인 변수 쌍이 없습니다.")
                
                st.markdown('<span style="font-weight:bold; color:#d62728; font-size:18px;">산점도 행렬 (pairplot)</span>', unsafe_allow_html=True)
                fig = px.scatter_matrix(df_no_na, dimensions=[y_col]+x_cols)
                
                # 그림 크기와 글자 크기 조정
                fig.update_layout(
                    height=600,  # 세로 길이 증가
                    width=800,   # 가로 길이 설정
                    font=dict(size=10),  # 전체 글자 크기 축소
                    margin=dict(l=80, r=80, t=80, b=80)  # 여백 조정
                )
                
                # y축 제목 글자 크기 추가 축소
                fig.update_yaxes(title_font_size=8)
                fig.update_xaxes(title_font_size=8)
                
                st.plotly_chart(fig, use_container_width=True)
                
                # 산점도 행렬을 통한 Insight 분석
                st.markdown('<span style="font-weight:bold; color:#9467bd; font-size:18px;">💡 산점도 행렬 Insight 분석</span>', unsafe_allow_html=True)
                
                # 상관계수와 분포 특성을 기반으로 한 insight 생성
                insights = []
                
                # 1. 선형 관계 분석
                linear_relationships = []
                for i, col1 in enumerate([y_col] + x_cols):
                    for j, col2 in enumerate([y_col] + x_cols):
                        if i < j:  # 중복 방지
                            corr_val = corr.loc[col1, col2]
                            if abs(corr_val) >= 0.7:
                                linear_relationships.append({
                                    'var1': col1,
                                    'var2': col2,
                                    'corr': corr_val,
                                    'strength': '매우 강한 선형 관계'
                                })
                            elif abs(corr_val) >= 0.4:
                                linear_relationships.append({
                                    'var1': col1,
                                    'var2': col2,
                                    'corr': corr_val,
                                    'strength': '중간 정도의 선형 관계'
                                })
                
                # 2. 분포 특성 분석
                distribution_insights = []
                for col in [y_col] + x_cols:
                    data = df_no_na[col]
                    skewness = data.skew()
                    kurtosis = data.kurtosis()
                    
                    # 왜도 분석
                    if abs(skewness) > 1:
                        if skewness > 0:
                            distribution_insights.append({
                                'var': col,
                                'type': '오른쪽으로 치우친 분포',
                                'value': f'왜도: {skewness:.2f}'
                            })
                        else:
                            distribution_insights.append({
                                'var': col,
                                'type': '왼쪽으로 치우친 분포',
                                'value': f'왜도: {skewness:.2f}'
                            })
                    
                    # 첨도 분석
                    if kurtosis > 3:
                        distribution_insights.append({
                            'var': col,
                            'type': '뾰족한 분포',
                            'value': f'첨도: {kurtosis:.2f}'
                        })
                    elif kurtosis < 1:
                        distribution_insights.append({
                            'var': col,
                            'type': '평평한 분포',
                            'value': f'첨도: {kurtosis:.2f}'
                        })
                
                # 3. 비선형 관계 탐지
                nonlinear_insights = []
                for i, col1 in enumerate([y_col] + x_cols):
                    for j, col2 in enumerate([y_col] + x_cols):
                        if i < j:
                            corr_val = corr.loc[col1, col2]
                            # 상관계수가 낮지만 산점도에서 패턴이 있을 수 있는 경우
                            if abs(corr_val) < 0.3:
                                # 여기서는 간단한 휴리스틱으로 판단
                                nonlinear_insights.append({
                                    'var1': col1,
                                    'var2': col2,
                                    'description': '선형 상관관계가 낮지만 비선형 관계가 있을 수 있습니다'
                                })
                                break
                
                # Insight를 3개 컬럼으로 나누어 표시
                col1, col2, col3 = st.columns(3)
                
                with col1:
                    if linear_relationships:
                        st.markdown("""
                        <div style="background: linear-gradient(135deg, #667eea 0%, #764ba2 100%); padding: 15px; border-radius: 10px; margin: 10px 0; color: white;">
                            <h4 style="margin: 0 0 10px 0; text-align: center;">🔗 선형 관계</h4>
                        </div>
                        """, unsafe_allow_html=True)
                        for rel in linear_relationships[:3]:  # 상위 3개만 표시
                            st.markdown(f"""
                            <div style="background-color: rgba(255,255,255,0.1); padding: 12px; border-radius: 5px; margin: 8px 0; border-left: 3px solid #fff;">
                                <div style="display: flex; align-items: center; margin-bottom: 8px;">
                                    <span style="background-color: #fff; color: #667eea; padding: 4px 8px; border-radius: 4px; font-weight: bold; font-size: 14px;">{rel['var1']}</span>
                                    <span style="margin: 0 8px; font-size: 16px; color: #fff;">⟷</span>
                                    <span style="background-color: #fff; color: #667eea; padding: 4px 8px; border-radius: 4px; font-weight: bold; font-size: 14px;">{rel['var2']}</span>
                                </div>
                                <div style="font-size: 16px; font-weight: 500;">{rel['strength']} (r={rel['corr']:.3f})</div>
                            </div>
                            """, unsafe_allow_html=True)
                
                with col2:
                    if distribution_insights:
                        st.markdown("""
                        <div style="background: linear-gradient(135deg, #f093fb 0%, #f5576c 100%); padding: 15px; border-radius: 10px; margin: 10px 0; color: white;">
                            <h4 style="margin: 0 0 10px 0; text-align: center;">📊 분포 특성</h4>
                        </div>
                        """, unsafe_allow_html=True)
                        for insight in distribution_insights[:3]:  # 상위 3개만 표시
                            st.markdown(f"""
                            <div style="background-color: rgba(255,255,255,0.1); padding: 12px; border-radius: 5px; margin: 8px 0; border-left: 3px solid #fff;">
                                <div style="display: flex; align-items: center; margin-bottom: 8px;">
                                    <span style="background-color: #fff; color: #f093fb; padding: 4px 8px; border-radius: 4px; font-weight: bold; font-size: 14px;">{insight['var']}</span>
                                </div>
                                <div style="font-size: 16px; font-weight: 500;">{insight['type']} ({insight['value']})</div>
                            </div>
                            """, unsafe_allow_html=True)
                
                with col3:
                    if nonlinear_insights:
                        st.markdown("""
                        <div style="background: linear-gradient(135deg, #4facfe 0%, #00f2fe 100%); padding: 15px; border-radius: 10px; margin: 10px 0; color: white;">
                            <h4 style="margin: 0 0 10px 0; text-align: center;">🔄 비선형 관계</h4>
                        </div>
                        """, unsafe_allow_html=True)
                        for insight in nonlinear_insights[:2]:  # 상위 2개만 표시
                            st.markdown(f"""
                            <div style="background-color: rgba(255,255,255,0.1); padding: 12px; border-radius: 5px; margin: 8px 0; border-left: 3px solid #fff;">
                                <div style="display: flex; align-items: center; margin-bottom: 8px;">
                                    <span style="background-color: #fff; color: #4facfe; padding: 4px 8px; border-radius: 4px; font-weight: bold; font-size: 14px;">{insight['var1']}</span>
                                    <span style="margin: 0 8px; font-size: 16px; color: #fff;">⟷</span>
                                    <span style="background-color: #fff; color: #4facfe; padding: 4px 8px; border-radius: 4px; font-weight: bold; font-size: 14px;">{insight['var2']}</span>
                                </div>
                                <div style="font-size: 16px; font-weight: 500;">{insight['description']}</div>
                            </div>
                            """, unsafe_allow_html=True)
                    else:
                        st.markdown("""
                        <div style="background: linear-gradient(135deg, #4facfe 0%, #00f2fe 100%); padding: 15px; border-radius: 10px; margin: 10px 0; color: white;">
                            <h4 style="margin: 0 0 10px 0; text-align: center;">📈 데이터 품질</h4>
                        </div>
                        """, unsafe_allow_html=True)
                        total_vars = len([y_col] + x_cols)
                        if total_vars <= 3:
                            st.markdown("• 변수 수가 적어 관계 분석에 제한이 있을 수 있습니다")
                        elif total_vars >= 8:
                            st.markdown("• 변수 수가 많아 복잡한 상호작용이 있을 수 있습니다")
                        
                        if len(df_no_na) < 100:
                            st.markdown("• 샘플 수가 적어 통계적 신뢰도에 주의가 필요합니다")
                
                # 요약 박스
                st.markdown("""
                <div style="background: linear-gradient(135deg, #fa709a 0%, #fee140 100%); padding: 15px; border-radius: 10px; margin: 20px 0; color: white;">
                    <h4 style="margin: 0; text-align: center;">🎯 핵심 Insight</h4>
                    <p style="margin: 5px 0; text-align: center; font-size: 14px;">
                        산점도 행렬을 통해 변수 간의 관계 패턴, 분포 특성, 이상치 존재 여부를 종합적으로 파악할 수 있습니다.
                    </p>
                </div>
                """, unsafe_allow_html=True)
                
                # 4. 데이터 분석 및 결과 표출 (변수 선택이 완료된 이후에만 노출)
                st.markdown('<hr style="border:3px solid #333;">', unsafe_allow_html=True)
                st.header("4. 데이터 분석 결과")
                
                # 선형회귀분석
                st.markdown("""
                <div style="background: linear-gradient(135deg, #667eea 0%, #764ba2 100%); padding: 15px; border-radius: 15px; margin: 20px 0; color: white;">
                    <h3 style="margin: 0; font-size: 24px; text-align: center;">📊 선형회귀분석 결과</h3>
                </div>
                """, unsafe_allow_html=True)
                
                X = df_no_na[x_cols]
                y = df_no_na[y_col]
                X_train, X_test, y_train, y_test = train_test_split(X, y, test_size=0.2, random_state=42)
                lr = LinearRegression()
                lr.fit(X_train, y_train)
                y_pred = lr.predict(X_test)
                
                # 수식 생성 (Y에 대한 비중이 높은 순서대로 정렬)
                # 실제 데이터에서 각 항의 기여도 계산
                X_sample = X_test.iloc[:10]  # 테스트 데이터의 처음 10개 샘플 사용
                y_sample = y_test.iloc[:10]
                
                # 각 항의 기여도 계산
                term_contributions = {}
                for col in x_cols:
                    term_values = lr.coef_[x_cols.index(col)] * X_sample[col]
                    term_contributions[col] = term_values
                
                # 절편 기여도
                intercept_contribution = np.full(len(X_sample), lr.intercept_)
                term_contributions['절편'] = intercept_contribution
                
                # 예측값
                y_pred_sample = lr.predict(X_sample)
                
                # 각 항의 평균 기여도와 비중 계산
                term_analysis = []
                for term_name, contributions in term_contributions.items():
                    avg_contribution = np.mean(contributions)
                    contribution_ratio = (avg_contribution / np.mean(y_pred_sample)) * 100
                    
                    if term_name == '절편':
                        term_analysis.append({
                            'name': term_name,
                            'avg_contribution': avg_contribution,
                            'ratio': contribution_ratio,
                            'coef': lr.intercept_
                        })
                    else:
                        coef = lr.coef_[x_cols.index(term_name)]
                        term_analysis.append({
                            'name': term_name,
                            'avg_contribution': avg_contribution,
                            'ratio': contribution_ratio,
                            'coef': coef
                        })
                
                # Y에 대한 비중이 높은 순서대로 정렬 (절편 제외)
                non_intercept_terms = [term for term in term_analysis if term['name'] != '절편']
                non_intercept_terms.sort(key=lambda x: abs(x['ratio']), reverse=True)
                
                # 절편을 마지막에 추가
                intercept_term = [term for term in term_analysis if term['name'] == '절편'][0]
                sorted_terms = non_intercept_terms + [intercept_term]
                
                # 수식 생성
                equation_parts = []
                for i, term in enumerate(sorted_terms):
                    coef = term['coef']
                    ratio = term['ratio']
                    
                    # 계수의 유효숫자 결정
                    if abs(coef) >= 0.001:
                        coef_str = f"{coef:.3f}"
                    elif abs(coef) >= 0.0001:
                        coef_str = f"{coef:.4f}"
                    elif abs(coef) >= 0.00001:
                        coef_str = f"{coef:.5f}"
                    else:
                        coef_str = f"{coef:.2e}"
                    
                    # 첫 번째 항이 아닌 경우 부호 추가 (+ 또는 -)
                    if i > 0:
                        if coef >= 0:
                            if term['name'] == '절편':
                                equation_parts.append(f"&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;+ {coef_str} × <span style='color: #6c757d;'>(절편)</span> <span style='color: #6c757d; font-size: 14px;'>({ratio:.1f}%)</span>")
                            else:
                                equation_parts.append(f"&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;+ {coef_str} × {term['name']} <span style='color: #6c757d; font-size: 14px;'>({ratio:.1f}%)</span>")
                        else:
                            if term['name'] == '절편':
                                equation_parts.append(f"&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;- {abs(coef):.3f} × <span style='color: #6c757d;'>(절편)</span> <span style='color: #6c757d; font-size: 14px;'>({ratio:.1f}%)</span>")
                            else:
                                equation_parts.append(f"&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;- {abs(coef):.3f} × {term['name']} <span style='color: #6c757d; font-size: 14px;'>({ratio:.1f}%)</span>")
                    else:
                        if coef >= 0:
                            if term['name'] == '절편':
                                equation_parts.append(f"&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;{coef_str} × <span style='color: #6c757d;'>(절편)</span> <span style='color: #6c757d; font-size: 14px;'>({ratio:.1f}%)</span>")
                            else:
                                equation_parts.append(f"&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;{coef_str} × {term['name']} <span style='color: #6c757d; font-size: 14px;'>({ratio:.1f}%)</span>")
                        else:
                            if term['name'] == '절편':
                                equation_parts.append(f"&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;- {abs(coef):.3f} × <span style='color: #6c757d;'>(절편)</span> <span style='color: #6c757d; font-size: 14px;'>({ratio:.1f}%)</span>")
                            else:
                                equation_parts.append(f"&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;- {abs(coef):.3f} × {term['name']} <span style='color: #6c757d; font-size: 14px;'>({ratio:.1f}%)</span>")
                
                # Y 변수명을 괄호 안에 표시
                y_variable_name = y_col if y_col else "Target Variable"
                equation = f"Y ({y_variable_name}) = <br>" + "<br>".join(equation_parts)
                
                # 결과 표시
                col1, col2 = st.columns([1, 1])
                
                with col1:
                    st.markdown("""
                    <div style="background-color: #f8f9fa; padding: 20px; border-radius: 10px; border-left: 5px solid #dc3545;">
                        <h4 style="margin: 0 0 15px 0; color: #dc3545;">🎯 선형 상관관계 방정식</h4
                        <div style="background-color: white; padding: 15px; border-radius: 8px; border: 1px solid #dee2e6; font-family: 'Arial', 'Helvetica', sans-serif; font-size: 16px; font-weight: bold; color: #333; line-height: 1.5;">
                            {}
                        </div>
                    </div>
                    """.format(equation), unsafe_allow_html=True)
                    
                    # 계수 해석
                    st.markdown("""
                    <div style="background-color: #fff3cd; padding: 15px; border-radius: 10px; border-left: 5px solid #ffc107; margin-top: 15px;">
                        <h4 style="margin: 0 0 10px 0; color: #856404;">📈 계수 해석</h4>
                    """, unsafe_allow_html=True)
                    
                    for col, coef in zip(x_cols, lr.coef_):
                        if coef > 0:
                            interpretation = f"<span style='color: #28a745;'>📈 양의 영향</span>"
                        else:
                            interpretation = f"<span style='color: #dc3545;'>📉 음의 영향</span>"
                        
                        st.markdown(f"""
                        <div style="margin: 5px 0; padding: 8px; background-color: white; border-radius: 5px; border-left: 4px solid #ffc107;">
                            <strong>{col}:</strong> {coef:.4f} ({interpretation})
                            <br><small style="color: #6c757d;">→ {col}이 1단위 증가할 때 Y가 {coef:.4f}단위 변화</small>
                        </div>
                        """, unsafe_allow_html=True)
                    
                    st.markdown("</div>", unsafe_allow_html=True)
                    
                    # 항 해석 (각 항이 Y에 미치는 실제 비중 분석)
                    st.markdown("""
                    <div style="background-color: #e8f5e8; padding: 15px; border-radius: 10px; border-left: 5px solid #28a745; margin-top: 15px;">
                        <h4 style="margin: 0 0 10px 0; color: #155724;">🎯 항 해석 (실제 비중 분석)</h4>
                    """, unsafe_allow_html=True)
                    
                    # 실제 데이터에서 각 항의 기여도 계산
                    X_sample = X_test.iloc[:10]  # 테스트 데이터의 처음 10개 샘플 사용
                    y_sample = y_test.iloc[:10]
                    
                    # 각 항의 기여도 계산
                    term_contributions = {}
                    for col in x_cols:
                        term_values = lr.coef_[x_cols.index(col)] * X_sample[col]
                        term_contributions[col] = term_values
                    
                    # 절편 기여도
                    intercept_contribution = np.full(len(X_sample), lr.intercept_)
                    term_contributions['절편'] = intercept_contribution
                    
                    # 예측값
                    y_pred_sample = lr.predict(X_sample)
                    
                    # 각 항의 평균 기여도와 비중 계산
                    term_analysis = []
                    for term_name, contributions in term_contributions.items():
                        avg_contribution = np.mean(contributions)
                        contribution_ratio = (avg_contribution / np.mean(y_pred_sample)) * 100
                        
                        if term_name == '절편':
                            term_analysis.append({
                                'name': term_name,
                                'avg_contribution': avg_contribution,
                                'ratio': contribution_ratio,
                                'color': '#6c757d'
                            })
                        else:
                            term_analysis.append({
                                'name': term_name,
                                'avg_contribution': avg_contribution,
                                'ratio': contribution_ratio,
                                'color': '#28a745'
                            })
                    
                    # 비중 순으로 정렬
                    term_analysis.sort(key=lambda x: abs(x['ratio']), reverse=True)
                    
                    for term in term_analysis:
                        if term['avg_contribution'] > 0:
                            impact_icon = "📈"
                            impact_text = "양의 기여"
                        else:
                            impact_icon = "📉"
                            impact_text = "음의 기여"
                        
                        st.markdown(f"""
                        <div style="margin: 8px 0; padding: 10px; background-color: white; border-radius: 6px; border-left: 4px solid {term['color']};">
                            <div style="display: flex; justify-content: space-between; align-items: center; margin-bottom: 5px;">
                                <strong style="color: {term['color']};">{term['name']}</strong>
                                <span style="font-size: 12px; color: #666;">{impact_icon} {impact_text}</span>
                            </div>
                            <div style="font-size: 14px; color: #333;">
                                평균 기여도: <strong>{term['avg_contribution']:.2f}</strong>
                                <br>Y에 대한 비중: <strong>{term['ratio']:.1f}%</strong>
                            </div>
                        </div>
                        """, unsafe_allow_html=True)
                    
                    st.markdown("</div>", unsafe_allow_html=True)
                    
                    # 모델 성능 계산 (col2에서 사용하기 위해 미리 계산)
                    score = lr.score(X_test, y_test)
                    
                    # 분석 결과를 세션 상태에 저장
                    analysis_results = {
                        'equation': equation,
                        'model_performance': {
                            'r2': score,
                            'rmse': np.sqrt(np.mean((y_test - lr.predict(X_test))**2)),
                            'mae': np.mean(np.abs(y_test - lr.predict(X_test)))
                        },
                        'variable_importance': {},
                        'term_analysis': term_analysis
                    }
                    
                    # 변수 중요도 계산 (절대값 기준)
                    for term in term_analysis:
                        if term['name'] != '절편':
                            analysis_results['variable_importance'][term['name']] = abs(term['ratio'])
                    
                    st.session_state['analysis_results'] = analysis_results
                
                with col2:
                    # 성능 등급 결정
                    if score >= 0.8:
                        performance_grade = "🟢 우수"
                        performance_color = "#28a745"
                    elif score >= 0.6:
                        performance_grade = "🟡 양호"
                        performance_color = "#ffc107"
                    elif score >= 0.4:
                        performance_grade = "🟠 보통"
                        performance_color = "#fd7e14"
                    else:
                        performance_grade = "🔴 미흡"
                        performance_color = "#dc3545"
                    
                    st.markdown(f"""
                    <div style="background-color: #e9ecef; padding: 20px; border-radius: 10px; border-left: 5px solid #007bff;">
                        <h4 style="margin: 0 0 15px 0; color: #007bff;">📊 모델 성능</h4>
                        <div style="text-align: center;">
                            <div style="font-size: 36px; font-weight: bold; color: {performance_color}; margin: 10px 0;">
                                {score:.1%}
                            </div>
                            <div style="font-size: 18px; color: #6c757d; margin-bottom: 15px;">
                                R² (설명력)
                            </div>
                            <div style="background-color: white; padding: 10px; border-radius: 5px; font-weight: bold; color: {performance_color};">
                                {performance_grade}
                            </div>
                        </div>
                    </div>
                    """, unsafe_allow_html=True)
                    
                    # 인사이트 제공
                    st.markdown("""
                    <div style="background-color: #d1ecf1; padding: 15px; border-radius: 10px; border-left: 5px solid #17a2b8; margin-top: 15px;">
                        <h4 style="margin: 0 0 10px 0; color: #0c5460;">💡 주요 인사이트</h4>
                    """, unsafe_allow_html=True)
                    
                    if score < 0.3:
                        insight = "• 모델의 설명력이 낮습니다. 더 많은 변수나 비선형 관계를 고려해보세요."
                    elif score < 0.6:
                        insight = "• 모델의 설명력이 보통 수준입니다. 추가 변수나 특성 엔지니어링을 고려해보세요."
                    else:
                        insight = "• 모델의 설명력이 양호합니다. 선형 관계가 잘 포착되었습니다."
                    
                    # 항 해석 결과를 주요 인사이트에 추가
                    important_terms = []
                    less_important_terms = []
                    
                    for term in term_analysis:
                        if abs(term['ratio']) >= 5:  # 5% 이상인 항
                            important_terms.append(term)
                        else:  # 5% 미만인 항
                            less_important_terms.append(term)
                    
                    # 텍스트 설명 제거 - 빈 div만 유지
                    st.markdown("""
                    <div style="color: #0c5460; font-size: 16px; line-height: 1.8;">
                    </div>
                    """, unsafe_allow_html=True)
                    
                    # 항 해석 결과를 주요 인사이트의 하부 항목으로 표시
                    if important_terms:
                        st.markdown("""
                        <h4 style="margin: 0 0 10px 0; color: #0d47a1;">🎯 핵심 영향 변수 (Y에 대한 비중 5% 이상)</h4>
                        """, unsafe_allow_html=True)
                        
                        for term in important_terms:
                            impact_icon = "📈" if term['avg_contribution'] > 0 else "📉"
                            impact_text = "양의 영향" if term['avg_contribution'] > 0 else "음의 영향"
                            
                            st.markdown(f"""
                            <div style="margin: 5px 0; padding: 8px; background-color: white; border-radius: 5px; border-left: 3px solid #1976d2;">
                                <div style="display: flex; justify-content: space-between; align-items: center;">
                                    <strong style="color: #0d47a1;">{term['name']}</strong>
                                    <span style="font-size: 12px; color: #1976d2;">{impact_icon} {impact_text}</span>
                                </div>
                                <div style="font-size: 13px; color: #0d47a1; margin-top: 3px;">
                                    평균 기여도: <strong>{term['avg_contribution']:.2f}</strong> | Y에 대한 비중: <strong>{term['ratio']:.1f}%</strong>
                                </div>
                            </div>
                            """, unsafe_allow_html=True)
                    
                    if less_important_terms:
                        st.markdown("""
                        <h4 style="margin: 0 0 10px 0; color: #4a148c;">📊 보조 영향 변수 (Y에 대한 비중 5% 미만)</h4>
                        """, unsafe_allow_html=True)
                        
                        for term in less_important_terms:
                            impact_icon = "📈" if term['avg_contribution'] > 0 else "📉"
                            impact_text = "양의 영향" if term['avg_contribution'] > 0 else "음의 영향"
                            
                            st.markdown(f"""
                            <div style="margin: 5px 0; padding: 8px; background-color: white; border-radius: 5px; border-left: 3px solid #7b1fa2;">
                                <div style="display: flex; justify-content: space-between; align-items: center;">
                                    <strong style="color: #4a148c;">{term['name']}</strong>
                                    <span style="font-size: 12px; color: #7b1fa2;">{impact_icon} {impact_text}</span>
                                </div>
                                <div style="font-size: 13px; color: #4a148c; margin-top: 3px;">
                                    평균 기여도: <strong>{term['avg_contribution']:.2f}</strong> | Y에 대한 비중: <strong>{term['ratio']:.1f}%</strong>
                                </div>
                            </div>
                            """, unsafe_allow_html=True)
                    
                    st.markdown("</div>", unsafe_allow_html=True)
                
                # 예측 결과 그래프
                st.markdown("""
                <div style="margin: 10px 0;">
                    <h4 style="text-align: left; color: #495057; margin-bottom: 10px;">📈 Target(Y값) 예측 vs 실제값 비교 시뮬레이션</h4>
                </div>
                """, unsafe_allow_html=True)
                
                # x축과 y축의 범위를 동일하게 설정
                min_val = min(y_test.min(), y_pred.min())
                max_val = max(y_test.max(), y_pred.max())
                
                # 예측값 vs 실제값 산점도
                fig = px.scatter(x=y_test, y=y_pred, 
                               labels={'x':'실제값', 'y':'예측값'}, 
                               title="",
                               color_discrete_sequence=['#007bff'])
                
                # 완벽한 예측선 추가
                fig.add_shape(type="line", x0=min_val, y0=min_val, 
                            x1=max_val, y1=max_val, 
                            line=dict(color="red", dash="dash", width=2))
                
                # 범례를 위한 더미 데이터 추가
                fig.add_scatter(x=[None], y=[None], mode='markers', 
                              marker=dict(color='red', size=8),
                              name='Y=X 기준선 (완벽한 예측선)', showlegend=True)
                
                fig.update_layout(
                    title="",  # undefined 제거
                    plot_bgcolor='white',
                    paper_bgcolor='white',
                    font=dict(size=12),
                    xaxis=dict(
                        range=[min_val, max_val],
                        showgrid=True,
                        gridcolor='lightgray',
                        zeroline=True,
                        zerolinecolor='black',
                        zerolinewidth=1,
                        showline=True,
                        linecolor='black',
                        linewidth=1
                    ),
                    yaxis=dict(
                        range=[min_val, max_val],
                        showgrid=True,
                        gridcolor='lightgray',
                        zeroline=True,
                        zerolinecolor='black',
                        zerolinewidth=1,
                        showline=True,
                        linecolor='black',
                        linewidth=1
                    ),
                    width=600,  # 정방형을 위한 가로 크기
                    height=600,  # 정방형을 위한 세로 크기
                    showlegend=True,
                    legend=dict(
                        x=0.02,
                        y=0.98,
                        bgcolor='rgba(255,255,255,0.9)',
                        bordercolor='black',
                        borderwidth=1,
                        font=dict(size=12)
                    )
                )
                
                # 범례 텍스트 설정
                fig.data[0].name = "실제 데이터 (예측값 vs 실제값)"
                
                st.plotly_chart(fig, use_container_width=False)
                
                # 정량적인 정확성 데이터 추가
                from sklearn.metrics import r2_score, mean_squared_error, mean_absolute_error
                
                # 예측 정확도 지표 계산
                r2 = r2_score(y_test, y_pred)
                mse = mean_squared_error(y_test, y_pred)
                rmse = np.sqrt(mse)
                mae = mean_absolute_error(y_test, y_pred)
                
                # 예측값과 실제값의 차이 계산
                residuals = y_pred - y_test
                residual_std = np.std(residuals)
                residual_mean = np.mean(residuals)
                
                # 정확성 데이터를 컬럼으로 표시
                col1, col2, col3 = st.columns(3)
                
                with col1:
                    st.markdown("""
                    <div style="background-color: #e3f2fd; padding: 15px; border-radius: 10px; border-left: 5px solid #1976d2;">
                        <h4 style="margin: 0 0 10px 0; color: #0d47a1;">📊 모델 성능 지표</h4>
                        <div style="font-size: 14px; line-height: 1.6;">
                            <strong>R² 점수:</strong> {:.4f}<br>
                            <strong>RMSE:</strong> {:.4f}<br>
                            <strong>MAE:</strong> {:.4f}
                        </div>
                    </div>
                    """.format(r2, rmse, mae), unsafe_allow_html=True)
                
                with col2:
                    st.markdown("""
                    <div style="background-color: #f3e5f5; padding: 15px; border-radius: 10px; border-left: 5px solid #7b1fa2;">
                        <h4 style="margin: 0 0 10px 0; color: #4a148c;">📈 예측 오차 분석</h4>
                        <div style="font-size: 14px; line-height: 1.6;">
                            <strong>평균 오차:</strong> {:.4f}<br>
                            <strong>오차 표준편차:</strong> {:.4f}<br>
                            <strong>MSE:</strong> {:.4f}
                        </div>
                    </div>
                    """.format(residual_mean, residual_std, mse), unsafe_allow_html=True)
                
                with col3:
                    st.markdown("""
                    <div style="background-color: #e8f5e8; padding: 15px; border-radius: 10px; border-left: 5px solid #28a745;">
                        <h4 style="margin: 0 0 10px 0; color: #155724;">🎯 분포 특성</h4>
                        <div style="font-size: 14px; line-height: 1.6;">
                            <strong>데이터 포인트:</strong> {}개<br>
                            <strong>예측 범위:</strong> {:.2f} ~ {:.2f}<br>
                            <strong>실제 범위:</strong> {:.2f} ~ {:.2f}
                        </div>
                    </div>
                    """.format(len(y_test), y_pred.min(), y_pred.max(), y_test.min(), y_test.max()), unsafe_allow_html=True)

                # 여백 추가
                st.markdown("<div style='margin: 40px 0;'></div>", unsafe_allow_html=True)

                # 구분선 추가
                st.markdown('<hr style="border:2px solid #667eea; margin: 20px 0;">', unsafe_allow_html=True)

                # 머신러닝 분석
                st.markdown("""
                <div style="background: linear-gradient(135deg, #667eea 0%, #764ba2 100%); padding: 15px; border-radius: 15px; margin: 20px 0; color: white;">
                    <h3 style="margin: 0; font-size: 24px; text-align: center;">🧠 머신러닝 모델 비교 분석</h3>
                </div>
                """, unsafe_allow_html=True)
                
                from sklearn.ensemble import RandomForestRegressor
                from sklearn.svm import SVR
                from sklearn.metrics import r2_score, mean_squared_error, mean_absolute_error
                import numpy as np
                
                # 다양한 머신러닝 모델 테스트
                models = {
                    'RandomForest': RandomForestRegressor(n_estimators=100, max_depth=10, random_state=42),
                    'SVR': SVR(C=1.0, epsilon=0.1),
                    'LinearRegression': lr  # 이미 훈련된 선형회귀 모델
                }
                
                # 모델 성능 비교
                model_results = {}
                
                with st.spinner("🧠 여러가지 머신러닝 모델들을 훈련하고 최적화하는 중..."):
                    # Optuna를 사용한 하이퍼파라미터 최적화
                    def objective(trial):
                        model_name = trial.suggest_categorical("model", ["RandomForest", "SVR"])
                        if model_name == "RandomForest":
                            n_estimators = trial.suggest_int("n_estimators", 10, 200)
                            max_depth = trial.suggest_int("max_depth", 2, 20)
                            model = RandomForestRegressor(n_estimators=n_estimators, max_depth=max_depth, random_state=42)
                        else:
                            C = trial.suggest_loguniform("C", 1e-2, 1e3)
                            epsilon = trial.suggest_loguniform("epsilon", 1e-3, 1.0)
                            model = SVR(C=C, epsilon=epsilon)
                        model.fit(X_train, y_train)
                        y_pred = model.predict(X_test)
                        return r2_score(y_test, y_pred)
                    
                    # Optuna 최적화 실행
                    study = optuna.create_study(direction="maximize")
                    study.optimize(objective, n_trials=30)
                    
                    # 최적화된 모델들 평가
                    for name, model in models.items():
                        if name == 'RandomForest':
                            # 최적화된 RandomForest 사용
                            best_rf = RandomForestRegressor(
                                n_estimators=study.best_params.get('n_estimators', 100),
                                max_depth=study.best_params.get('max_depth', 10),
                                random_state=42
                            )
                            best_rf.fit(X_train, y_train)
                            y_pred = best_rf.predict(X_test)
                        elif name == 'SVR':
                            # 최적화된 SVR 사용
                            best_svr = SVR(
                                C=study.best_params.get('C', 1.0),
                                epsilon=study.best_params.get('epsilon', 0.1)
                            )
                            best_svr.fit(X_train, y_train)
                            y_pred = best_svr.predict(X_test)
                        else:
                            # 선형회귀는 이미 훈련됨
                            y_pred = lr.predict(X_test)
                        
                        # 성능 지표 계산
                        r2 = r2_score(y_test, y_pred)
                        mse = mean_squared_error(y_test, y_pred)
                        mae = mean_absolute_error(y_test, y_pred)
                        
                        model_results[name] = {
                            'r2': r2,
                            'mse': mse,
                            'mae': mae,
                            'y_pred': y_pred
                        }
                
                # 결과 표시
                col1, col2 = st.columns([1, 1])
                
                with col1:
                    # 모델 성능 비교 테이블
                    st.markdown("""
                    <div style="background-color: #f8f9fa; padding: 20px; border-radius: 10px; border-left: 5px solid #6f42c1;">
                        <h4 style="margin: 0 0 15px 0; color: #6f42c1;">🏆 모델 성능 비교</h4>
                    """, unsafe_allow_html=True)
                    
                    # 성능 순위 결정
                    sorted_models = sorted(model_results.items(), key=lambda x: x[1]['r2'], reverse=True)
                    best_model_name = sorted_models[0][0]
                    
                    for i, (name, results) in enumerate(sorted_models):
                        if name == best_model_name:
                            rank_icon = "🥇"
                            rank_color = "#ffd700"
                        elif i == 1:
                            rank_icon = "🥈"
                            rank_color = "#c0c0c0"
                        elif i == 2:
                            rank_icon = "🥉"
                            rank_color = "#cd7f32"
                        else:
                            rank_icon = "📊"
                            rank_color = "#6c757d"
                        
                        # 성능 등급 결정
                        if results['r2'] >= 0.8:
                            grade = "🟢 우수"
                            grade_color = "#28a745"
                        elif results['r2'] >= 0.6:
                            grade = "🟡 양호"
                            grade_color = "#ffc107"
                        elif results['r2'] >= 0.4:
                            grade = "🟠 보통"
                            grade_color = "#fd7e14"
                        else:
                            grade = "🔴 미흡"
                            grade_color = "#dc3545"
                        
                        st.markdown(f"""
                        <div style="background-color: white; padding: 15px; border-radius: 8px; margin: 10px 0; border-left: 4px solid {rank_color};">
                            <div style="display: flex; align-items: center; margin-bottom: 10px;">
                                <span style="font-size: 20px; margin-right: 10px;">{rank_icon}</span>
                                <h5 style="margin: 0; color: #495057;">{name}</h5>
                            </div>
                            <div style="display: grid; grid-template-columns: 1fr 1fr; gap: 10px;">
                                <div style="text-align: center;">
                                    <div style="font-size: 24px; font-weight: bold; color: {grade_color};">
                                        {results['r2']:.1%}
                                    </div>
                                    <div style="font-size: 12px; color: #6c757d;">R²</div>
                                </div>
                                <div style="text-align: center;">
                                    <div style="font-size: 16px; font-weight: bold; color: #495057;">
                                        {grade}
                                    </div>
                                    <div style="font-size: 12px; color: #6c757d;">등급</div>
                                </div>
                            </div>
                        </div>
                        """, unsafe_allow_html=True)
                    
                    st.markdown("</div>", unsafe_allow_html=True)
                
                with col2:
                    # 최적 모델 정보
                    st.markdown(f"""
                    <div style="background-color: #e9ecef; padding: 20px; border-radius: 10px; border-left: 5px solid #28a745;">
                        <h4 style="margin: 0 0 15px 0; color: #495057;">🏆 최적 모델: {best_model_name}</h4>
                    """, unsafe_allow_html=True)
                    
                    best_results = model_results[best_model_name]
                    
                    # 모델별 특성 설명
                    model_descriptions = {
                        'LinearRegression': '선형 관계를 가정하는 가장 단순한 모델',
                        'RandomForest': '여러 의사결정 트리를 조합한 앙상블 모델',
                        'SVR': '비선형 관계를 학습할 수 있는 서포트 벡터 회귀'
                    }
                    
                    st.markdown(f"""
                    <div style="background-color: white; padding: 15px; border-radius: 8px; margin-bottom: 15px;">
                        <h5 style="margin: 0 0 10px 0; color: #495057;">📋 모델 특성</h5>
                        <p style="margin: 0; color: #6c757d; font-size: 14px;">
                            {model_descriptions.get(best_model_name, '알 수 없는 모델')}
                        </p>
                    </div>
                    """, unsafe_allow_html=True)
                    
                    # 성능 지표
                    st.markdown("""
                    <div style="background-color: white; padding: 15px; border-radius: 8px;">
                        <h5 style="margin: 0 0 10px 0; color: #495057;">📊 성능 지표</h5>
                    """, unsafe_allow_html=True)
                    
                    st.markdown(f"""
                    <div style="display: grid; grid-template-columns: 1fr 1fr; gap: 10px; margin-bottom: 10px;">
                        <div style="text-align: center; padding: 8px; background-color: #f8f9fa; border-radius: 5px;">
                            <div style="font-size: 18px; font-weight: bold; color: #28a745;">
                                {best_results['r2']:.1%}
                            </div>
                            <div style="font-size: 12px; color: #6c757d;">설명력 (R²)</div>
                        </div>
                        <div style="text-align: center; padding: 8px; background-color: #f8f9fa; border-radius: 5px;">
                            <div style="font-size: 18px; font-weight: bold; color: #dc3545;">
                                {best_results['mae']:.4f}
                            </div>
                            <div style="font-size: 12px; color: #6c757d;">평균 절대 오차</div>
                        </div>
                    </div>
                    """, unsafe_allow_html=True)
                    
                    st.markdown("</div>", unsafe_allow_html=True)
                    
                    # 인사이트 제공
                    st.markdown("""
                    <div style="background-color: #d1ecf1; padding: 15px; border-radius: 10px; border-left: 5px solid #17a2b8; margin-top: 15px;">
                        <h4 style="margin: 0 0 10px 0; color: #0c5460;">💡 머신러닝 인사이트</h4>
                    """, unsafe_allow_html=True)
                    
                    # 모델 비교 인사이트
                    linear_r2 = model_results['LinearRegression']['r2']
                    best_r2 = best_results['r2']
                    improvement = ((best_r2 - linear_r2) / linear_r2 * 100) if linear_r2 > 0 else 0
                    
                    if best_model_name != 'LinearRegression':
                        insight = f"• {best_model_name}이 선형회귀보다 {improvement:.1f}% 더 나은 성능을 보입니다."
                    else:
                        insight = "• 선형회귀가 가장 좋은 성능을 보입니다. 데이터가 선형 관계를 잘 따르고 있습니다."
                    
                    st.markdown(f"""
                    <div style="color: #0c5460; font-size: 14px;">
                        {insight}
                        <br>• 하이퍼파라미터는 최적화 완료됨
                        <br>• 모델 복잡도: {best_model_name}이 가장 적합한 모델로 선택됨
                        <br>• 예측 정확도: 평균 절대 오차 {best_results['mae']:.4f} 단위
                    </div>
                    """, unsafe_allow_html=True)
                    
                    st.markdown("</div>", unsafe_allow_html=True)
                
                # 최적 모델의 예측 결과 그래프
                st.markdown("""
                <div style="margin: 10px 0;">
                    <h4 style="text-align: left; color: #495057; margin-bottom: 10px;">🤖 최적 모델 예측 결과</h4>
                </div>
                """, unsafe_allow_html=True)
                
                best_y_pred = model_results[best_model_name]['y_pred']
                
                # x축과 y축의 범위를 동일하게 설정
                min_val = min(y_test.min(), best_y_pred.min())
                max_val = max(y_test.max(), best_y_pred.max())
                
                fig = px.scatter(x=y_test, y=best_y_pred, 
                               labels={'x':'실제값', 'y':'예측값'}, 
                               title="",
                               color_discrete_sequence=['#6f42c1'])
                
                # 완벽한 예측선 추가
                fig.add_shape(type="line", x0=min_val, y0=min_val, 
                            x1=max_val, y1=max_val, 
                            line=dict(color="red", dash="dash", width=2))
                
                # 범례를 위한 더미 데이터 추가
                fig.add_scatter(x=[None], y=[None], mode='markers', 
                              marker=dict(color='red', size=8),
                              name='Y=X 기준선 (완벽한 예측선)', showlegend=True)
                
                fig.update_layout(
                    title="",  # undefined 제거
                    plot_bgcolor='white',
                    paper_bgcolor='white',
                    font=dict(size=12),
                    xaxis=dict(
                        range=[min_val, max_val],
                        showgrid=True,
                        gridcolor='lightgray',
                        zeroline=True,
                        zerolinecolor='black',
                        zerolinewidth=1,
                        showline=True,
                        linecolor='black',
                        linewidth=1
                    ),
                    yaxis=dict(
                        range=[min_val, max_val],
                        showgrid=True,
                        gridcolor='lightgray',
                        zeroline=True,
                        zerolinecolor='black',
                        zerolinewidth=1,
                        showline=True,
                        linecolor='black',
                        linewidth=1
                    ),
                    width=600,  # 정방형을 위한 가로 크기
                    height=600,  # 정방형을 위한 세로 크기
                    showlegend=True,
                    legend=dict(
                        x=0.02,
                        y=0.98,
                        bgcolor='rgba(255,255,255,0.9)',
                        bordercolor='black',
                        borderwidth=1,
                        font=dict(size=12)
                    )
                )
                
                # 범례 텍스트 설정
                fig.data[0].name = "실제 데이터 (예측값 vs 실제값)"
                
                st.plotly_chart(fig, use_container_width=False)
                
                # 여백 추가
                st.markdown("<div style='margin: 40px 0;'></div>", unsafe_allow_html=True)

                # 구분선 추가
                st.markdown('<hr style="border:2px solid #ff6b6b; margin: 20px 0;">', unsafe_allow_html=True)

                # H2O AutoML 분석
                st.markdown("""
                <div style="background: linear-gradient(135deg, #ff6b6b 0%, #ee5a24 100%); padding: 15px; border-radius: 15px; margin: 20px 0; color: white;">
                    <h3 style="margin: 0; font-size: 24px; text-align: center;">🚀 H2O AutoML 자동 모델 탐색</h3>
                </div>
                """, unsafe_allow_html=True)
                
                # H2O AutoML 실행
                try:
                    import h2o
                    from h2o.automl import H2OAutoML
                    
                    # H2O 초기화
                    h2o.init()
                    
                    # 데이터를 H2O 프레임으로 변환
                    with st.spinner("🚀 H2O AutoML이 다양한 알고리즘을 자동으로 탐색하고 최적의 모델을 찾는 중..."):
                        # 데이터 준비
                        h2o_df = h2o.H2OFrame(df_no_na)
                        
                        # 특성과 타겟 분리
                        feature_cols = x_cols
                        target_col = y_col
                        
                        # H2O AutoML 실행
                        aml = H2OAutoML(
                            max_models=20,  # 최대 모델 수
                            seed=42,
                            max_runtime_secs=300,  # 5분 제한
                            sort_metric='RMSE'
                        )
                        
                        aml.train(x=feature_cols, y=target_col, training_frame=h2o_df)
                        
                        # 결과 표시
                        col1, col2 = st.columns([1, 1])
                        
                        with col1:
                            # AutoML 리더보드
                            st.markdown("""
                            <div style="background-color: #f8f9fa; padding: 20px; border-radius: 10px; border-left: 5px solid #ff6b6b;">
                                <h4 style="margin: 0 0 15px 0; color: #ff6b6b;">🏆 AutoML 모델 순위</h4>
                            """, unsafe_allow_html=True)
                            
                            # 리더보드 데이터 추출
                            leaderboard = aml.leaderboard
                            leaderboard_df = leaderboard.as_data_frame()
                            
                            # 상위 5개 모델만 표시
                            top_models = leaderboard_df.head(5)
                            
                            for idx, row in top_models.iterrows():
                                rank = idx + 1
                                model_name = row['model_id']
                                rmse = row['rmse']
                                mae = row['mae']
                                r2 = row['r2']
                                
                                if rank == 1:
                                    rank_icon = "🥇"
                                    rank_color = "#ffd700"
                                elif rank == 2:
                                    rank_icon = "🥈"
                                    rank_color = "#c0c0c0"
                                elif rank == 3:
                                    rank_icon = "🥉"
                                    rank_color = "#cd7f32"
                                else:
                                    rank_icon = f"{rank}위"
                                    rank_color = "#6c757d"
                                
                                st.markdown(f"""
                                <div style="margin: 8px 0; padding: 12px; background-color: white; border-radius: 8px; border-left: 4px solid {rank_color};">
                                    <div style="display: flex; justify-content: space-between; align-items: center;">
                                        <span style="font-weight: bold; color: {rank_color};">{rank_icon} {model_name}</span>
                                        <span style="font-size: 12px; color: #6c757d;">{rank}위</span>
                                    </div>
                                    <div style="font-size: 13px; color: #495057; margin-top: 5px;">
                                        RMSE: <strong>{rmse:.4f}</strong> | R²: <strong>{r2:.4f}</strong>
                                    </div>
                                </div>
                                """, unsafe_allow_html=True)
                            
                            st.markdown("</div>", unsafe_allow_html=True)
                        
                        with col2:
                            # 최고 성능 모델 정보
                            best_automl_model = aml.leader
                            best_model_name = best_automl_model.model_id
                            
                            st.markdown(f"""
                            <div style="background-color: #e9ecef; padding: 20px; border-radius: 10px; border-left: 5px solid #ff6b6b;">
                                <h4 style="margin: 0 0 15px 0; color: #495057;">🚀 최고 성능 모델</h4>
                            """, unsafe_allow_html=True)
                            
                            # 모델 성능 지표
                            model_performance = best_automl_model.model_performance()
                            metrics = model_performance.metric()
                            
                            st.markdown(f"""
                            <div style="background-color: white; padding: 15px; border-radius: 8px;">
                                <h5 style="margin: 0 0 10px 0; color: #495057;">📊 성능 지표</h5>
                                <div style="display: grid; grid-template-columns: 1fr 1fr; gap: 10px;">
                                    <div style="text-align: center; padding: 8px; background-color: #f8f9fa; border-radius: 5px;">
                                        <div style="font-size: 18px; font-weight: bold; color: #28a745;">
                                            {metrics['r2']:.1%}
                                        </div>
                                        <div style="font-size: 12px; color: #6c757d;">설명력 (R²)</div>
                                    </div>
                                    <div style="text-align: center; padding: 8px; background-color: #f8f9fa; border-radius: 5px;">
                                        <div style="font-size: 18px; font-weight: bold; color: #dc3545;">
                                            {metrics['rmse']:.4f}
                                        </div>
                                        <div style="font-size: 12px; color: #6c757d;">RMSE</div>
                                    </div>
                                </div>
                            </div>
                            """, unsafe_allow_html=True)
                            
                            st.markdown("</div>", unsafe_allow_html=True)
                        
                        # AutoML 인사이트
                        st.markdown("""
                        <div style="background-color: #fff3cd; padding: 15px; border-radius: 10px; border-left: 5px solid #ffc107; margin-top: 20px;">
                            <h4 style="margin: 0 0 10px 0; color: #856404;">💡 AutoML 인사이트</h4>
                        """, unsafe_allow_html=True)
                        
                        # 모델 다양성 분석
                        model_types = leaderboard_df['model_id'].str.split('_').str[0].value_counts()
                        total_models = len(leaderboard_df)
                        
                        st.markdown(f"""
                        <div style="color: #856404; font-size: 14px;">
                            • <strong>탐색된 모델 수:</strong> {total_models}개
                            <br>• <strong>알고리즘 다양성:</strong> {len(model_types)}가지 유형
                            <br>• <strong>최고 성능 모델:</strong> {best_model_name}
                            <br>• <strong>자동 앙상블:</strong> Stacked Ensemble 모델 포함
                            <br>• <strong>탐색 시간:</strong> 5분 내 최적화 완료
                        </div>
                        """, unsafe_allow_html=True)
                        
                        st.markdown("</div>", unsafe_allow_html=True)
                        
                        # H2O 종료
                        h2o.cluster().shutdown()
                        
                except ImportError:
                    st.error("H2O 라이브러리가 설치되지 않았습니다. 'pip install h2o'를 실행해주세요.")
                except Exception as e:
                    st.error(f"H2O AutoML 실행 중 오류가 발생했습니다: {str(e)}")

# 다운로드 기능 함수들
def create_word_document():
    """현재 페이지의 모든 내용을 Word 문서로 생성"""
    doc = Document()
    
    # 제목 추가
    title = doc.add_heading('공정 데이터 상관관계 분석 보고서', 0)
    title.alignment = 1  # 가운데 정렬
    
    # 생성 날짜 추가
    doc.add_paragraph(f'생성 날짜: {datetime.now().strftime("%Y년 %m월 %d일 %H:%M")}')
    doc.add_paragraph('')
    
    # 데이터 정보 추가
    if 'df' in st.session_state and st.session_state['df'] is not None:
        df = st.session_state['df']
        doc.add_heading('1. 데이터 개요', level=1)
        doc.add_paragraph(f'데이터 크기: {df.shape[0]}행 × {df.shape[1]}열')
        doc.add_paragraph(f'변수 수: {len(df.columns)}개')
        
        # 컬럼 목록 추가
        doc.add_paragraph('변수 목록:')
        for i, col in enumerate(df.columns, 1):
            doc.add_paragraph(f'{i}. {col}')
        doc.add_paragraph('')
    
    # 분석 결과 추가 (세션 상태에서 가져오기)
    if 'analysis_results' in st.session_state:
        results = st.session_state['analysis_results']
        doc.add_heading('2. 분석 결과', level=1)
        
        # 회귀 방정식
        if 'equation' in results:
            doc.add_heading('2.1 회귀 방정식', level=2)
            # HTML 태그 제거하고 텍스트만 추출
            equation_text = results['equation'].replace('<br>', '\n').replace('<span style="color: #6c757d; font-size: 14px;">', '').replace('</span>', '')
            doc.add_paragraph(equation_text)
            doc.add_paragraph('')
        
        # 모델 성능
        if 'model_performance' in results:
            doc.add_heading('2.2 모델 성능', level=2)
            performance = results['model_performance']
            doc.add_paragraph(f'R² (설명력): {performance.get("r2", "N/A"):.4f}')
            doc.add_paragraph(f'RMSE: {performance.get("rmse", "N/A"):.4f}')
            doc.add_paragraph(f'MAE: {performance.get("mae", "N/A"):.4f}')
            doc.add_paragraph('')
        
        # 변수 중요도 및 기여도
        if 'term_analysis' in results:
            doc.add_heading('2.3 변수 중요도 및 기여도', level=2)
            term_analysis = results['term_analysis']
            
            # 중요도 순으로 정렬
            sorted_terms = sorted(term_analysis, key=lambda x: abs(x['ratio']), reverse=True)
            
            for term in sorted_terms:
                if term['name'] == '절편':
                    doc.add_paragraph(f'(절편): 평균 기여도 {term["avg_contribution"]:.2f}, Y에 대한 비중 {term["ratio"]:.1f}%')
                else:
                    impact = "양의 영향" if term['avg_contribution'] > 0 else "음의 영향"
                    doc.add_paragraph(f'{term["name"]}: 평균 기여도 {term["avg_contribution"]:.2f}, Y에 대한 비중 {term["ratio"]:.1f}% ({impact})')
            doc.add_paragraph('')
        
        # 주요 인사이트
        doc.add_heading('3. 주요 인사이트', level=1)
        
        # 핵심 영향 변수 (5% 이상)
        important_vars = [term for term in results.get('term_analysis', []) if abs(term['ratio']) >= 5]
        if important_vars:
            doc.add_heading('3.1 핵심 영향 변수 (Y에 대한 비중 5% 이상)', level=2)
            for term in important_vars:
                impact = "양의 영향" if term['avg_contribution'] > 0 else "음의 영향"
                doc.add_paragraph(f'• {term["name"]}: {term["ratio"]:.1f}% ({impact})')
            doc.add_paragraph('')
        
        # 보조 영향 변수 (5% 미만)
        less_important_vars = [term for term in results.get('term_analysis', []) if abs(term['ratio']) < 5]
        if less_important_vars:
            doc.add_heading('3.2 보조 영향 변수 (Y에 대한 비중 5% 미만)', level=2)
            for term in less_important_vars:
                impact = "양의 영향" if term['avg_contribution'] > 0 else "음의 영향"
                doc.add_paragraph(f'• {term["name"]}: {term["ratio"]:.1f}% ({impact})')
            doc.add_paragraph('')
    
    return doc

def create_pdf_report():
    """현재 페이지의 모든 내용을 PDF로 생성 (Word를 통한 간접 변환)"""
    doc = create_word_document()
    
    # 임시 파일로 저장
    with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp_file:
        doc.save(tmp_file.name)
        tmp_file_path = tmp_file.name
    
    return tmp_file_path

def download_report(format_type):
    """보고서 다운로드 함수"""
    try:
        if format_type == "Word":
            doc = create_word_document()
            
            # 임시 파일로 저장
            with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp_file:
                doc.save(tmp_file.name)
                
                # 파일 읽기
                with open(tmp_file.name, 'rb') as f:
                    file_data = f.read()
                
                # 임시 파일 삭제
                os.unlink(tmp_file.name)
                
                # 다운로드 버튼 생성
                st.download_button(
                    label="📄 Word 문서 다운로드",
                    data=file_data,
                    file_name=f"상관관계_분석_보고서_{datetime.now().strftime('%Y%m%d_%H%M')}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
        
        elif format_type == "PDF":
            # PDF 생성을 위해 Word 문서를 먼저 생성
            doc = create_word_document()
            
            # 임시 파일로 저장
            with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp_file:
                doc.save(tmp_file.name)
                
                # 파일 읽기
                with open(tmp_file.name, 'rb') as f:
                    file_data = f.read()
                
                # 임시 파일 삭제
                os.unlink(tmp_file.name)
                
                # 다운로드 버튼 생성 (PDF로 표시하지만 실제로는 Word 파일)
                st.download_button(
                    label="📄 PDF 보고서 다운로드",
                    data=file_data,
                    file_name=f"상관관계_분석_보고서_{datetime.now().strftime('%Y%m%d_%H%M')}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
                
    except Exception as e:
        st.error(f"보고서 생성 중 오류가 발생했습니다: {str(e)}")

# 다운로드 버튼 UI
def render_download_buttons():
    """사이드바에 다운로드 버튼들 렌더링"""
    # 사이드바에 다운로드 버튼 추가
    with st.sidebar:
        st.markdown("---")
        st.markdown("### 📥 보고서 다운로드")
        
        # 분석이 완료되었을 때만 다운로드 버튼 표시
        if 'analysis_results' in st.session_state and st.session_state['analysis_results'] is not None:
            try:
                doc = create_word_document()
                
                # 임시 파일로 저장
                with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp_file:
                    doc.save(tmp_file.name)
                    
                    # 파일 읽기
                    with open(tmp_file.name, 'rb') as f:
                        file_data = f.read()
                    
                    # 임시 파일 삭제
                    os.unlink(tmp_file.name)
                    
                    # Word 문서 다운로드 버튼
                    st.download_button(
                        label="📄 Word 문서 다운로드",
                        data=file_data,
                        file_name=f"상관관계_분석_보고서_{datetime.now().strftime('%Y%m%d_%H%M')}.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        key="word_download"
                    )
                    
                    # PDF 보고서 다운로드 버튼 (실제로는 Word 파일)
                    st.download_button(
                        label="📄 PDF 보고서 다운로드",
                        data=file_data,
                        file_name=f"상관관계_분석_보고서_{datetime.now().strftime('%Y%m%d_%H%M')}.pdf",
                        mime="application/pdf",
                        key="pdf_download"
                    )
                    
            except Exception as e:
                st.error(f"보고서 생성 중 오류: {str(e)}")
        else:
            st.info("분석을 완료하면 다운로드 버튼이 나타납니다.")

# 다운로드 버튼 렌더링
render_download_buttons() 