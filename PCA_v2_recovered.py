import streamlit as st
import pandas as pd
import numpy as np
import plotly.express as px
import plotly.figure_factory as ff
from sklearn.linear_model import LinearRegression
from sklearn.model_selection import train_test_split
import optuna
import os
import io
import base64
from docx import Document
from docx.shared import Inches
import tempfile
from datetime import datetime
import plotly.graph_objects as go
from selenium import webdriver
from selenium.webdriver.chrome.options import Options
from selenium.webdriver.common.by import By
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
import time
from PIL import Image
import plotly.io as pio
import chardet

st.set_page_config(page_title="Correlation Analyzer", layout="wide")



st.markdown("""
<div style="margin: 10px 0;">
    <h1 style="margin: 0; padding: 10px 0; text-align: left; color: #000000; font-size: 42px; font-weight: bold;">
        공정 데이터 상관관계 분석 도우미
    </h1>
</div>
""", unsafe_allow_html=True)

# 1. 파일 업로드
st.header("1. 파일 업로드")
uploaded_file = st.file_uploader("CSV 또는 Excel 파일을 업로드하세요", type=["csv", "xlsx"])
st.info("분석할 파일을 업로드하세요. Data는 회사 내부 서버에 저장됩니다.") 

if 'df' not in st.session_state:
    st.session_state['df'] = None
if 'data_file_path' not in st.session_state:
    st.session_state['data_file_path'] = None
if 'preprocessing_steps' not in st.session_state:
    st.session_state['preprocessing_steps'] = []
if 'current_df' not in st.session_state:
    st.session_state['current_df'] = None
if 'data_period' not in st.session_state:
    st.session_state['data_period'] = None
if 'preprocessing_completed' not in st.session_state:
    st.session_state['preprocessing_completed'] = False

# 파일이 업로드되면 모든 상태 초기화
if uploaded_file is not None:
    # 새로운 파일이 업로드되었는지 확인
    current_file_name = uploaded_file.name
    if 'uploaded_file_name' not in st.session_state or st.session_state['uploaded_file_name'] != current_file_name:
        # 새로운 파일이므로 모든 상태 초기화
        st.session_state['uploaded_file_name'] = current_file_name
        st.session_state['preprocessing_steps'] = []
        st.session_state['current_df'] = None
        st.session_state['df'] = None
        st.session_state['data_file_path'] = None
        st.session_state['data_period'] = None  # 기간 정보도 초기화
        st.session_state['preprocessing_completed'] = False  # 전처리 완료 상태도 초기화
        # 변수 선택 관련 상태도 초기화
        if 'y_col' in st.session_state:
            del st.session_state['y_col']
        if 'x_cols' in st.session_state:
            del st.session_state['x_cols']

if uploaded_file is not None:
    if uploaded_file.name.endswith('.csv'):
        # 더 강력한 인코딩 감지 및 CSV 파일 읽기
        # 파일 내용을 바이트로 읽어서 인코딩 감지
        uploaded_file.seek(0)  # 파일 포인터를 처음으로
        raw_data = uploaded_file.read()
        uploaded_file.seek(0)  # 파일 포인터를 다시 처음으로
        
        # chardet를 사용한 인코딩 감지
        detected_encoding = chardet.detect(raw_data)
        detected_encoding_name = detected_encoding['encoding']
        confidence = detected_encoding['confidence']
        
        # 인코딩 목록 (우선순위 순)
        encodings_to_try = [
            'utf-8-sig',  # BOM이 있는 UTF-8
            'utf-8',      # 일반 UTF-8
            'cp949',      # 한글 Windows
            'euc-kr',     # 한글 EUC-KR
            'latin1',     # ISO-8859-1
            'iso-8859-1', # ISO-8859-1
            'cp1252',     # Windows-1252
            'ascii'       # ASCII
        ]
        
        # 감지된 인코딩이 있으면 맨 앞에 추가
        if detected_encoding_name and confidence > 0.7:
            encodings_to_try.insert(0, detected_encoding_name)
        
        df = None
        successful_encoding = None
        
        for encoding in encodings_to_try:
            try:
                uploaded_file.seek(0)  # 파일 포인터를 처음으로
                df = pd.read_csv(uploaded_file, encoding=encoding, on_bad_lines='skip')
                successful_encoding = encoding
                break
            except UnicodeDecodeError:
                continue
            except Exception as e:
                # 다른 오류도 처리
                continue
        
        if df is None:
            # 모든 인코딩이 실패한 경우, 더 관대한 설정으로 시도
            try:
                uploaded_file.seek(0)
                df = pd.read_csv(uploaded_file, encoding='utf-8', on_bad_lines='skip', 
                                 sep=None, engine='python')
                successful_encoding = 'utf-8 (with errors)'
            except:
                try:
                    uploaded_file.seek(0)
                    df = pd.read_csv(uploaded_file, encoding='latin1', on_bad_lines='skip', engine='python')
                    successful_encoding = 'latin1 (with errors)'
                except:
                    try:
                        uploaded_file.seek(0)
                        # 더 관대한 설정으로 시도
                        df = pd.read_csv(uploaded_file, encoding='utf-8', on_bad_lines='skip', 
                                       sep=None, engine='python')
                        successful_encoding = 'utf-8 (auto-detect)'
                    except Exception as e:
                        st.error(f"❌ 파일을 읽을 수 없습니다. 오류: {str(e)}")
                        st.error("💡 해결 방법:")
                        st.error("1. 파일을 UTF-8로 저장한 후 다시 업로드해주세요.")
                        st.error("2. 파일에 특수문자나 깨진 문자가 있는지 확인해주세요.")
                        st.error("3. 파일이 올바른 CSV 형식인지 확인해주세요.")
                        st.error("4. 파일이 텍스트 형식인지 확인해주세요.")
                        st.stop()
        
        # 성공한 인코딩 정보 표시
        if successful_encoding:
            st.success(f"✅ 파일이 성공적으로 로드되었습니다! (인코딩: {successful_encoding})")
        
        # 구분자 감지 및 재시도 (필요한 경우)
        if df is not None and len(df.columns) == 1:
            # 한 개의 열만 있다면 구분자 문제일 가능성이 높음
            
            # 일반적인 구분자들 시도
            delimiters = [',', ';', '\t', '|', ':', ' ']
            
            for delimiter in delimiters:
                try:
                    uploaded_file.seek(0)
                    temp_df = pd.read_csv(uploaded_file, encoding=successful_encoding, delimiter=delimiter, on_bad_lines='skip')
                    if len(temp_df.columns) > 1:
                        df = temp_df
                        break
                except:
                    continue
    else:
        df = pd.read_excel(uploaded_file)
    
    # 원본 데이터 보존
    st.session_state['original_df'] = df.copy()  # 원본 데이터 보존
    st.session_state['df'] = df.copy()  # 작업용 데이터
    st.session_state['current_df'] = df.copy()  # 전처리용 데이터
    

    
    # 업로드된 파일을 서버에 저장 (안전한 파일 저장)
    import os
    import time
    
    # 파일명에 타임스탬프 추가하여 중복 방지
    timestamp = int(time.time())
    base_name = uploaded_file.name.rsplit('.', 1)[0]
    extension = uploaded_file.name.rsplit('.', 1)[1] if '.' in uploaded_file.name else ''
    
    save_path = f"uploaded_{base_name}_{timestamp}.{extension}"
    
    try:
        if uploaded_file.name.endswith('.csv'):
            df.to_csv(save_path, index=False)
        else:
            df.to_excel(save_path, index=False)
        st.session_state['data_file_path'] = save_path
    except PermissionError:
        # 권한 오류 시 임시 파일명 사용
        temp_save_path = f"temp_uploaded_{timestamp}.{extension}"
        try:
            if uploaded_file.name.endswith('.csv'):
                df.to_csv(temp_save_path, index=False)
            else:
                df.to_excel(temp_save_path, index=False)
            st.session_state['data_file_path'] = temp_save_path
            st.warning(f"⚠️ 원본 파일명으로 저장할 수 없어 임시 파일명으로 저장했습니다: {temp_save_path}")
        except Exception as e:
            st.error(f"❌ 파일 저장 중 오류가 발생했습니다: {str(e)}")
            st.session_state['data_file_path'] = None
    except Exception as e:
        st.error(f"❌ 파일 저장 중 오류가 발생했습니다: {str(e)}")
        st.session_state['data_file_path'] = None
    
    # 굵은 가로 구분선 추가
    st.markdown('<hr style="border:3px solid #333;">', unsafe_allow_html=True)
    
    # 스크롤을 조정하여 결측치 처리 섹션이 화면 중앙 위쪽에 오도록 함
    st.markdown("""
    <script>
    // 페이지 로드 후 스크롤 조정
    function adjustScroll() {
        // 결측치 처리 섹션을 찾아서 스크롤 조정
        const elements = document.querySelectorAll('h1, h2, h3, h4, h5, h6');
        for (let element of elements) {
            if (element.textContent.includes('결측치 처리')) {
                // 화면 높이의 20% 지점으로 스크롤 (더 위쪽으로)
                const targetPosition = element.offsetTop - window.innerHeight * 0.2;
                window.scrollTo({
                    top: targetPosition,
                    behavior: 'smooth'
                });
                break;
            }
        }
    }
    
    // 여러 시점에서 스크롤 조정 시도
    window.addEventListener('load', function() {
        setTimeout(adjustScroll, 500);
        setTimeout(adjustScroll, 1500);
        setTimeout(adjustScroll, 2500);
    });
    
    // DOM 변경 감지
    const observer = new MutationObserver(function(mutations) {
        mutations.forEach(function(mutation) {
            if (mutation.type === 'childList') {
                setTimeout(adjustScroll, 100);
            }
        });
    });
    
    observer.observe(document.body, {
        childList: true,
        subtree: true
    });
    </script>
    """, unsafe_allow_html=True)
    
    st.header("2. 데이터 전처리")
    # st.subheader("컬럼별 데이터형 변환")
    # dtype_map = {col: str(st.session_state['df'][col].dtype) for col in st.session_state['df'].columns}
    # new_dtypes = {}
    # dtype_options = {"문자열": "object", "숫자": "float64", "날짜": "datetime64[ns]"}
    # cols = st.columns(len(st.session_state['df'].columns))
    # for i, col in enumerate(st.session_state['df'].columns):
    #     with cols[i]:
    #         dtype = st.selectbox(f"{col}", options=list(dtype_options.keys()), index=list(dtype_options.values()).index(dtype_map[col]) if dtype_map[col] in dtype_options.values() else 0, key=f"dtype_{col}")
    #         new_dtypes[col] = dtype_options[dtype]
    # if st.button("데이터형 변환 적용"):
    #     for col, dtype in new_dtypes.items():
    #         try:
    #             if dtype == "datetime64[ns]":
    #                 st.session_state['df'][col] = pd.to_datetime(st.session_state['df'][col], errors='coerce')
    #             else:
    #                 st.session_state['df'][col] = st.session_state['df'][col].astype(dtype)
    #         except Exception as e:
    #             st.warning(f"{col} 변환 실패: {e}")
    #     st.success("데이터형 변환 완료!")

    # 데이터 미리보기
    with st.spinner('Loading ...'):
        st.subheader("데이터 미리보기")
        
        # 데이터 요약 정보 표시 - 항상 최신 데이터 사용
        df = st.session_state['df']
        original_df = st.session_state.get('original_df', df)
        
        total_rows, total_cols = df.shape
        original_rows, original_cols = original_df.shape
        numeric_cols = df.select_dtypes(include=['number']).shape[1]
        text_cols = df.select_dtypes(include=['object']).shape[1]
        datetime_cols = df.select_dtypes(include=['datetime']).shape[1]
        missing_total = df.isnull().sum().sum()
        missing_ratio = (missing_total / (total_rows * total_cols)) * 100
        
        # 데이터 손실 여부 확인
        data_loss = original_rows - total_rows
        loss_percentage = (data_loss / original_rows) * 100 if original_rows > 0 else 0
        

        
        # 원본 데이터로 복원 버튼 (데이터 손실이 있을 때만 표시)
        if data_loss > 0:
            col1, col2 = st.columns([1, 3])
            with col1:
                if st.button("🔄 원본 데이터로 복원", help="전처리 과정을 초기화하고 원본 데이터로 되돌립니다."):
                    st.session_state['df'] = st.session_state['original_df'].copy()
                    st.session_state['current_df'] = st.session_state['original_df'].copy()
                    st.session_state['preprocessing_steps'] = []
                    st.session_state['preprocessing_completed'] = False
                    st.success("✅ 원본 데이터로 복원되었습니다!")
                    st.rerun()
            with col2:
                st.info("💡 데이터 손실이 감지되었습니다. '원본 데이터로 복원' 버튼을 클릭하여 처음부터 다시 시작할 수 있습니다.")
        
        # 날짜형 열 변환 기능 추가 - 최신 데이터로 조건 확인
        current_datetime_cols = df.select_dtypes(include=['datetime']).shape[1]
        current_text_cols = df.select_dtypes(include=['object']).shape[1]
        
        if current_datetime_cols == 0 and current_text_cols > 0:
            st.markdown("""
            <div style="background-color: #fff3cd; padding: 15px; border-radius: 8px; margin: 10px 0; border-left: 4px solid #ffc107;">
                <h4 style="margin: 0 0 10px 0; color: #856404;">⚠️ 날짜형 열 변환 필요</h4>
                <p style="margin: 0 0 10px 0; color: #856404; font-size: 14px;">
                    날짜/시간 데이터가 문자형으로 인식되었습니다. 날짜형으로 변환하면 더 정확한 분석이 가능합니다.
                </p>
            </div>
            """, unsafe_allow_html=True)
            
            # 날짜형으로 변환할 열 선택 - 가로 배치
            text_columns = df.select_dtypes(include=['object']).columns.tolist()
            
            # 날짜 열 추천 함수
            def recommend_date_column(df, text_columns):
                """가장 유력한 날짜 열을 추천하는 함수"""
                if not text_columns:
                    return None
                
                # 각 열의 날짜 가능성 점수 계산
                column_scores = {}
                
                for col in text_columns:
                    sample_data = df[col].dropna().head(100)  # 처음 100개 샘플만 확인
                    if len(sample_data) == 0:
                        continue
                    
                    score = 0
                    
                    # 1. 열 이름 기반 점수
                    col_lower = col.lower()
                    date_keywords = ['date', 'time', '날짜', '시간', '시작', '종료', '시작일', '종료일', 'timestamp', 'datetime']
                    for keyword in date_keywords:
                        if keyword in col_lower:
                            score += 10
                    
                    # 2. 데이터 패턴 기반 점수
                    try:
                        # 날짜 변환 시도
                        converted = pd.to_datetime(sample_data, errors='coerce')
                        valid_ratio = converted.notna().sum() / len(sample_data)
                        
                        if valid_ratio > 0.8:  # 80% 이상이 유효한 날짜
                            score += 50
                        elif valid_ratio > 0.5:  # 50% 이상이 유효한 날짜
                            score += 30
                        elif valid_ratio > 0.2:  # 20% 이상이 유효한 날짜
                            score += 10
                        
                        # 고유값 비율 (날짜는 보통 고유값이 많음)
                        unique_ratio = sample_data.nunique() / len(sample_data)
                        if unique_ratio > 0.8:
                            score += 5
                        
                    except:
                        pass
                    
                    # 3. 데이터 길이 기반 점수 (날짜는 보통 일정한 길이)
                    avg_length = sample_data.astype(str).str.len().mean()
                    if 8 <= avg_length <= 20:  # 일반적인 날짜 길이
                        score += 5
                    
                    column_scores[col] = score
                
                # 가장 높은 점수의 열 반환
                if column_scores:
                    best_column = max(column_scores, key=column_scores.get)
                    if column_scores[best_column] > 10:  # 최소 점수 기준
                        return best_column
                
                return None
            
            # 추천 열 찾기
            recommended_column = recommend_date_column(df, text_columns)
            
            # 가로로 배치하기 위해 컬럼 사용
            col1, col2 = st.columns([1, 1])
            
            with col1:
                # 추천 열이 있으면 기본값으로 설정
                default_index = 0
                if recommended_column:
                    try:
                        default_index = text_columns.index(recommended_column) + 1  # +1은 "변환하지 않음" 옵션 때문
                    except:
                        pass
                
                date_col = st.selectbox(
                    "날짜형으로 변환할 열 선택",
                    options=["변환하지 않음"] + text_columns,
                    index=default_index,
                    help="날짜/시간 데이터가 포함된 열을 선택하세요. 추천된 열이 자동으로 선택됩니다."
                )
            
            with col2:
                if date_col != "변환하지 않음":
                    # 날짜 형식 선택
                    date_format = st.selectbox(
                        "날짜 형식 선택",
                        options=["자동 감지", "YYYY-MM-DD", "YYYY/MM/DD", "DD-MM-YYYY", "DD/MM/YYYY", "기타"],
                        help="날짜 데이터의 형식을 선택하세요. '자동 감지'를 권장합니다."
                    )
                    
                    if st.button("날짜형 변환 적용"):
                        try:
                            if date_format == "자동 감지":
                                # 자동 감지 시도
                                converted_data = pd.to_datetime(df[date_col], errors='coerce')
                            else:
                                # 수동 형식 지정
                                format_map = {
                                    "YYYY-MM-DD": "%Y-%m-%d",
                                    "YYYY/MM/DD": "%Y/%m/%d", 
                                    "DD-MM-YYYY": "%d-%m-%Y",
                                    "DD/MM/YYYY": "%d/%m/%Y"
                                }
                                
                                if date_format in format_map:
                                    converted_data = pd.to_datetime(df[date_col], format=format_map[date_format], errors='coerce')
                                else:
                                    # 기타 형식의 경우 자동 감지
                                    converted_data = pd.to_datetime(df[date_col], errors='coerce')
                            
                            # 변환 결과 확인
                            valid_dates = converted_data.dropna()
                            conversion_ratio = len(valid_dates) / len(df) * 100
                            
                            if conversion_ratio > 50:  # 50% 이상 변환 성공
                                st.session_state['df'][date_col] = converted_data
                                st.success(f"✅ {date_col} 열이 날짜형으로 변환되었습니다! (성공률: {conversion_ratio:.1f}%)")
                                
                                # 기간 정보 저장
                                if len(valid_dates) > 0:
                                    start_date = valid_dates.min()
                                    end_date = valid_dates.max()
                                    st.session_state['data_period'] = {
                                        'start': start_date,
                                        'end': end_date,
                                        'column': date_col
                                    }
                                
                                st.rerun()
                            else:
                                st.warning(f"⚠️ 변환 성공률이 낮습니다 ({conversion_ratio:.1f}%). 다른 형식을 시도해보세요.")
                                
                        except Exception as e:
                            st.error(f"❌ 날짜 변환 중 오류가 발생했습니다: {str(e)}")
                
                # 기간 정보가 있으면 표시
                if st.session_state.get('data_period'):
                    period_info = st.session_state['data_period']
                    start_str = period_info['start'].strftime('%Y-%m-%d %H:%M:%S')
                    end_str = period_info['end'].strftime('%Y-%m-%d %H:%M:%S')
                    date_col = period_info['column']
                    
                    st.markdown(f"""
                    <div style="background-color: #e3f2fd; padding: 15px; border-radius: 8px; margin: 10px 0; border-left: 4px solid #2196f3;">
                        <h4 style="margin: 0 0 10px 0; color: #1976d2;">📅 데이터 기간 정보</h4>
                        <p style="margin: 0; color: #1976d2; font-size: 14px;">
                            <strong>기간:</strong> {start_str} ~ {end_str}<br>
                            <strong>날짜 열:</strong> {date_col}
                        </p>
                    </div>
                    """, unsafe_allow_html=True)
            
            st.markdown('<div class="data-preview-table">', unsafe_allow_html=True)
            # 행번호 열에 "행번호" 헤더 추가하고 인덱스를 초기화
            df_with_index = st.session_state['df'].copy()
            df_with_index = df_with_index.reset_index(drop=True)  # 인덱스 초기화
            df_with_index.index.name = '행번호'
            
            # 데이터 미리보기 정보 추가
            st.markdown(f"""
            <div style="background-color: #e8f5e8; padding: 10px; border-radius: 5px; margin: 10px 0; border-left: 4px solid #4caf50;">
                <p style="margin: 0; font-size: 14px; color: #2e7d32;">
                    📋 <strong>미리보기 정보:</strong> 전체 {len(df_with_index)}행 중 처음 20행 표시 (인덱스: 0~{min(19, len(df_with_index)-1)})
                </p>
            </div>
            """, unsafe_allow_html=True)
            
            st.dataframe(df_with_index, use_container_width=True, height=400, hide_index=False)
            st.markdown('</div>', unsafe_allow_html=True)

    # 결측치 처리 헤더
    st.markdown("""
    <div style="margin-bottom: 10px;">
        <div style="display: flex; align-items: center; gap: -2px;">
            <h3 style="margin: 0;">결측치 처리</h3>
            <p style="font-size: 15px; color: #666; margin: 0;">(🔄 여러 번 누적하여 실행 가능합니다.)</p>
        </div>
    </div>
    """, unsafe_allow_html=True)
    
    # Skip 버튼을 왼쪽에 배치
    col1, col2, col3 = st.columns([1, 1, 1])
    with col1:
        if st.button("⏭️ Skip", help="결측치가 있는 데이터를 분석에 사용하지 않을 예정이라면 Skip 해도 됩니다.", key="skip_preprocessing"):
            # 최종 결과를 메인 데이터프레임에 적용
            st.session_state['df'] = st.session_state['current_df'].copy()
            st.session_state['preprocessing_completed'] = True  # 전처리 완료 상태 설정
            st.info("전처리가 완료되어 분석 데이터가 업데이트되었습니다.")
            if st.session_state['data_file_path']:
                try:
                    if st.session_state['data_file_path'].endswith('.csv'):
                        st.session_state['df'].to_csv(st.session_state['data_file_path'], index=False)
                    else:
                        st.session_state['df'].to_excel(st.session_state['data_file_path'], index=False)
                except PermissionError:
                    st.warning("⚠️ 파일이 다른 프로그램에서 사용 중이어서 저장할 수 없습니다.")
                except Exception as e:
                    st.warning(f"⚠️ 파일 저장 중 오류가 발생했습니다: {str(e)}")
            st.rerun()
      
    # 전처리 내용 표시
    # 전처리 과정 추적을 위한 session state 초기화
    if 'preprocessing_steps' not in st.session_state:
        st.session_state['preprocessing_steps'] = []
    if 'current_df' not in st.session_state or st.session_state['current_df'] is None:
        st.session_state['current_df'] = st.session_state['df'].copy()
    
    # 단계별 요약 표시
    if st.session_state['preprocessing_steps']:
        st.markdown('<h3 style="font-size: 18px; margin-bottom: 11px;">📋 수행된 전처리 단계 요약</h3>', unsafe_allow_html=True)
        for i, step in enumerate(st.session_state['preprocessing_steps'], 1):
            st.info(f"**{i}단계:** {step}")
        st.markdown("---")
    
    # 현재 단계의 결측치 처리 UI
    
    # 현재 결측치 상태 확인 (전처리 단계 정보 표시 전에 계산)
    current_na_count = st.session_state['current_df'].isnull().sum().sum()
    
    # 결측치 분포도 제목과 범례
    st.markdown('<h4 style="font-size: 16px; font-style: italic; color: #000; margin-bottom: 8px;">결측치 분포도</h4>', unsafe_allow_html=True)
    
    # 범례
    st.markdown('''
    <div style="display: flex; align-items: center; gap: 15px; font-size: 12px; color: #666; margin-bottom: 15px; margin-left: 10px;">
        <span style="display: flex; align-items: center; gap: 5px;">
            <div style="width: 12px; height: 12px; background-color: #288549; border-radius: 2px;"></div>
            정상 데이터
        </span>
        <div style="width: 1px; height: 20px; background-color: #ddd; margin: 0 10px;"></div>
        <span style="display: flex; align-items: center; gap: 5px;">
            <div style="width: 12px; height: 12px; background-color: #FFD700; border-radius: 2px;"></div>
            1단계 결측 (공백, None, NaN)
        </span>
        <span style="display: flex; align-items: center; gap: 5px;">
            <div style="width: 12px; height: 12px; background-color: #FF0000; border-radius: 2px;"></div>
            2단계 결측 (#DIV/0!, #N/A, #NAME? 등)
        </span>
        <span style="display: flex; align-items: center; gap: 5px;">
            <div style="width: 12px; height: 12px; background-color: #8E44AD; border-radius: 2px;"></div>
            3단계 결측 (#NULL!, #SPILL! 등)
        </span>
    </div>
    ''', unsafe_allow_html=True)

    # Missingno 차트 생성
    import matplotlib.pyplot as plt
    import numpy as np
    
    df = st.session_state['current_df']  # 현재 처리 중인 데이터 사용
    
    # 전처리 완료 후 분포도 업데이트 플래그 확인
    update_viz = st.session_state.get('update_visualization', False)
    if update_viz:
        st.session_state['update_visualization'] = False  # 플래그 초기화
        st.info("🔄 분포도가 업데이트되었습니다.")
        st.info(f"현재 데이터 크기: {df.shape[0]}행 x {df.shape[1]}열")
    
    # df가 None이 아닌지 확인
    if df is None:
        st.error("데이터가 로드되지 않았습니다. 파일을 다시 업로드해주세요.")
    else:
        # 새로운 기준에 따른 결측치 분류
        # 1단계: 공백만
        # 2단계: 1단계 + #DIV/0!, #N/A, #NAME?, #NUM!, #REF!, #VALUE!, #NODATA
        # 3단계: 2단계 + #NULL!, #SPILL!, #CALC!, #GETTING_DATA, #FIELD!, #BLOCKED!, #CONNECT!, #BUSY!, #UNKNOWN!
        
        # 1단계: 공백만 (None, NaN, 빈 문자열 모두 포함)
        stage1_mask = pd.DataFrame(False, index=df.index, columns=df.columns)
        for col in df.columns:
            try:
                # None, NaN, 빈 문자열을 모두 공백으로 처리
                stage1_mask[col] = (df[col].isna()) | (df[col].astype(str).str.strip() == '')
            except:
                # 문자열 변환이 안 되는 경우 빈 값으로 처리
                stage1_mask[col] = df[col].isnull()
        
        # 2단계: 1단계 + Excel 오류 값들
        error_values_2 = ['#DIV/0!', '#N/A', '#NAME?', '#NUM!', '#REF!', '#VALUE!', '#NODATA']
        stage2_mask = stage1_mask.copy()
        for col in df.columns:
            # Excel 오류 값들 확인
            excel_error_mask = df[col].isin(error_values_2)
            stage2_mask[col] = stage2_mask[col] | excel_error_mask
        
        # 3단계: 2단계 + 추가 Excel 오류 값들
        error_values_3 = ['#NULL!', '#SPILL!', '#CALC!', '#GETTING_DATA', '#FIELD!', 
                         '#BLOCKED!', '#CONNECT!', '#BUSY!', '#UNKNOWN!']
        stage3_mask = stage2_mask.copy()
        for col in df.columns:
            # 추가 Excel 오류 값들 확인
            additional_error_mask = df[col].isin(error_values_3)
            stage3_mask[col] = stage3_mask[col] | additional_error_mask
        
        n_rows, n_cols = df.shape
        
        # matplotlib 그래프 생성
        fig, ax = plt.subplots(figsize=(4, 2.0))
       
        # 1. 배경을 초록색으로 채우기 (정상 데이터)
        green_data = np.ones((n_rows, n_cols)) * 0.8
        ax.imshow(green_data, cmap='Greens', alpha=1.0, 
                 extent=[0, n_cols, 0, n_rows], aspect='auto', vmin=0, vmax=1)
        
        # 2. 단계별 결측치를 다른 색상으로 표시 (빨간색→보라색 그라데이션)
        # 1단계 결측치 (노란색) - 공백만
        stage1_positions = np.where(stage1_mask)
        if len(stage1_positions[0]) > 0:
            for row, col in zip(stage1_positions[0], stage1_positions[1]):
                ax.axhline(row, xmin=col/n_cols, xmax=(col+1)/n_cols, 
                          color='#FFD700', alpha=1, linewidth=0.5)
        
        # 2단계 결측치 (완전 빨간색) - 1단계 + Excel 오류 값들 (1단계와 겹치지 않는 부분만)
        stage2_only = stage2_mask & ~stage1_mask
        stage2_positions = np.where(stage2_only)
        if len(stage2_positions[0]) > 0:
            for row, col in zip(stage2_positions[0], stage2_positions[1]):
                ax.axhline(row, xmin=col/n_cols, xmax=(col+1)/n_cols, 
                          color='#FF0000', alpha=1, linewidth=0.5)
        
        # 3단계 결측치 (보라색) - 2단계 + 추가 Excel 오류 값들 (2단계와 겹치지 않는 부분만)
        stage3_only = stage3_mask & ~stage2_mask
        stage3_positions = np.where(stage3_only)
        if len(stage3_positions[0]) > 0:
            for row, col in zip(stage3_positions[0], stage3_positions[1]):
                ax.axhline(row, xmin=col/n_cols, xmax=(col+1)/n_cols, 
                          color='#8E44AD', alpha=1, linewidth=0.5)
        
        # 3. 격자 선들 (벡터화)
        # 세로 선들
        if n_cols > 1:
            ax.vlines(np.arange(1, n_cols), 0, n_rows, color='black', linewidth=0.1, alpha=0.2)
        
        # 가로 선들
        if n_rows > 1:
            ax.hlines(np.arange(1, n_rows), 0, n_cols, color='white', linewidth=0.1, alpha=0.1)
        
        # 4. 테두리 선 (1/3 수준으로 얇게)
        ax.hlines(0, 0, n_cols, color='black', linewidth=0.33)
        ax.hlines(n_rows, 0, n_cols, color='black', linewidth=0.33)
        ax.vlines(0, 0, n_rows, color='black', linewidth=0.33)
        ax.vlines(n_cols, 0, n_rows, color='black', linewidth=0.33)
        
        ax.set_xlim(0, n_cols)
        ax.set_ylim(n_rows, 0)
        
        # x축 번호 매기기 (글자 겹침 방지를 위해 간격 조정)
        # 더 넓은 간격으로 눈금 개수 줄이기
        target_x_ticks = min(12, max(3, n_cols // 5))  # 최소 3개, 최대 12개로 줄임
        if n_cols <= target_x_ticks:
            x_ticks = np.arange(n_cols) + 0.5
            x_labels = [str(i+1) for i in range(n_cols)]
        else:
            step = max(1, n_cols // target_x_ticks)
            x_ticks = np.arange(0, n_cols, step) + 0.5
            x_labels = [str(i+1) for i in range(0, n_cols, step)]
        
        ax.set_xticks(x_ticks)
        ax.set_xticklabels(x_labels, rotation=0, ha='center', fontsize=5)
        # x축 숫자를 위쪽에 표시
        ax.xaxis.set_ticks_position('top')
        ax.xaxis.set_label_position('top')
        
        # y축 번호 매기기 (글자 겹침 방지를 위해 간격 조정)
        # 더 넓은 간격으로 눈금 개수 줄이기
        target_y_ticks = min(8, max(2, n_rows // 100))  # 최소 2개, 최대 8개로 줄임
        if n_rows <= target_y_ticks:
            y_ticks = np.arange(n_rows + 1)
            y_labels = [str(i) for i in range(n_rows + 1)]
        else:
            step = max(1, n_rows // target_y_ticks)
            y_ticks = np.arange(0, n_rows + 1, step)
            y_labels = [str(i) for i in range(0, n_rows + 1, step)]
        
        ax.set_yticks(y_ticks)
        ax.set_yticklabels(y_labels, fontsize=5)
        
        ax.set_xlabel('Column Number', fontsize=6)
        ax.set_ylabel('Row Number', fontsize=6)
        ax.grid(False)
        ax.set_facecolor((1,1,1,0))
        
        # 테두리만 표시
        for spine in ax.spines.values():
            spine.set_visible(True)
            spine.set_linewidth(0.5)
                    
        # 그래프 크기를 제한하는 컨테이너
        st.markdown("""
        <style>
        .missingno-container {
            max-width: 400vw !important;
            max-height: 1000px !important;
            overflow: hidden;
            margin-top: -250px !important;
            padding-top: 0 !important;
            margin-bottom: -10px !important;
        }
        .missingno-container .stPlotlyChart {
            max-width: 400vw !important;
            max-height: 1000px !important;
        }
        </style>
        """, unsafe_allow_html=True)
        
        # missingno 그래프와 결측치 열 정보를 좌우로 배치
        col1, col2 = st.columns([1, 1])
        
        with col1:
            # 그래프 표시
            st.pyplot(fig, use_container_width=False)
        
        with col2:
            # 결측치 통계 정보 - 새로운 기준에 따른 분류
            # 1단계: 공백 (None, NaN, 빈 문자열 모두 포함)
            stage1_mask = pd.DataFrame(False, index=df.index, columns=df.columns)
            for col in df.columns:
                try:
                    # None, NaN, 빈 문자열을 모두 공백으로 처리
                    stage1_mask[col] = (df[col].isna()) | (df[col].astype(str).str.strip() == '')
                except:
                    stage1_mask[col] = df[col].isnull()
            
            # 2단계: 1단계 + Excel 오류 값들
            error_values_2 = ['#DIV/0!', '#N/A', '#NAME?', '#NUM!', '#REF!', '#VALUE!', '#NODATA']
            stage2_mask = stage1_mask.copy()
            for col in df.columns:
                excel_error_mask = df[col].isin(error_values_2)
                stage2_mask[col] = stage2_mask[col] | excel_error_mask
            
            # 3단계: 2단계 + 추가 Excel 오류 값들
            error_values_3 = ['#NULL!', '#SPILL!', '#CALC!', '#GETTING_DATA', '#FIELD!', 
                             '#BLOCKED!', '#CONNECT!', '#BUSY!', '#UNKNOWN!']
            stage3_mask = stage2_mask.copy()
            for col in df.columns:
                additional_error_mask = df[col].isin(error_values_3)
                stage3_mask[col] = stage3_mask[col] | additional_error_mask
            
            # 모든 결측치를 포함하는 마스크 (3단계가 모든 결측치를 포함)
            all_missing_mask = stage3_mask
            
            row_has_na = all_missing_mask.any(axis=1)
            col_has_na = all_missing_mask.any(axis=0)
            
            na_rows = df[row_has_na]  # 결측치가 있는 행만
            na_cols = df.columns[col_has_na].tolist()  # 결측치가 있는 열만
            total_rows = len(df)
            total_cols = len(df.columns)
            na_count_rows = row_has_na.sum()  # 더 빠른 계산
            na_count_cols = col_has_na.sum()  # 더 빠른 계산
            current_na_count = all_missing_mask.sum().sum()  # 모든 결측치 개수
            
            if current_na_count > 0:
                # 결측치가 있는 경우 - 처리 옵션 표시
                
                # 2. 결측치가 있는 열 정보
                st.markdown('''
                <div style="margin-left: 20px; margin-top: 15px;">
                    <p style="font-size: 16px; font-style: italic; color: #000; margin-bottom: 0px;">결측치가 있는 열 정보</p>
                </div>
                ''', unsafe_allow_html=True)
                missing_col_info = []
                for col in na_cols:
                    # 새로운 기준에 따른 각 단계별 결측치 개수 계산
                    
                    # 1단계: 공백만 (None, NaN, 빈 문자열 모두 포함)
                    try:
                        # None, NaN, 빈 문자열을 모두 공백으로 처리
                        stage1_mask = (df[col].isna()) | (df[col].astype(str).str.strip() == '')
                        stage1_count = stage1_mask.sum()
                    except:
                        stage1_count = 0
                    
                    # 2단계: 1단계 + Excel 오류 값들 (1단계와 겹치지 않는 부분만)
                    error_values_2 = ['#DIV/0!', '#N/A', '#NAME?', '#NUM!', '#REF!', '#VALUE!', '#NODATA']
                    stage2_only = df[col].isin(error_values_2) & ~stage1_mask
                    stage2_count = stage2_only.sum()
                    
                    # 3단계: 2단계 + 추가 Excel 오류 값들 (2단계와 겹치지 않는 부분만)
                    error_values_3 = ['#NULL!', '#SPILL!', '#CALC!', '#GETTING_DATA', '#FIELD!', '#BLOCKED!', '#CONNECT!', '#BUSY!', '#UNKNOWN!']
                    stage3_only = df[col].isin(error_values_3) & ~stage1_mask & ~stage2_only
                    stage3_count = stage3_only.sum()
                    
                    total_missing = stage1_count + stage2_count + stage3_count
                    missing_ratio = (total_missing / len(df)) * 100
                    
                    # 각 단계별 상세 정보 생성
                    detail_parts = []
                    
                    # 1단계: 공백 (None, NaN, 빈 문자열 세분화)
                    if stage1_count > 0:
                        # None, NaN, 빈 문자열을 개별적으로 카운트
                        none_count = df[col].isna().sum()
                        empty_str_count = ((df[col].astype(str).str.strip() == '') & ~df[col].isna()).sum()
                        
                        stage1_details = []
                        if none_count > 0:
                            stage1_details.append(f"None/NaN {none_count}개")
                        if empty_str_count > 0:
                            stage1_details.append(f"공백문자 {empty_str_count}개")
                        
                        if stage1_details:
                            detail_parts.append(f"🟡 1단계: {', '.join(stage1_details)}")
                        else:
                            detail_parts.append(f"🟡 1단계: 공백 {stage1_count}개")
                    
                    # 2단계: Excel 오류 값들 (개별 카운트)
                    if stage2_count > 0:
                        excel_errors_2 = ['#DIV/0!', '#N/A', '#NAME?', '#NUM!', '#REF!', '#VALUE!', '#NODATA']
                        stage2_details = []
                        for error in excel_errors_2:
                            error_count = (df[col].astype(str) == error).sum()
                            if error_count > 0:
                                stage2_details.append(f"{error} {error_count}개")
                        
                        if stage2_details:
                            detail_parts.append(f"🔴 2단계: {', '.join(stage2_details)}")
                    
                    # 3단계: 추가 Excel 오류 값들 (개별 카운트)
                    if stage3_count > 0:
                        excel_errors_3 = ['#NULL!', '#SPILL!', '#CALC!', '#GETTING_DATA', '#FIELD!', '#BLOCKED!', '#CONNECT!', '#BUSY!', '#UNKNOWN!']
                        stage3_details = []
                        for error in excel_errors_3:
                            error_count = (df[col].astype(str) == error).sum()
                            if error_count > 0:
                                stage3_details.append(f"{error} {error_count}개")
                        
                        if stage3_details:
                            detail_parts.append(f"🟣 3단계: {', '.join(stage3_details)}")
                    
                    detail_text = " | ".join(detail_parts) if detail_parts else "없음"
                    
                    # 열 번호 계산 (0부터 시작하는 인덱스를 1부터 시작하는 번호로 변환)
                    col_index = df.columns.get_loc(col) + 1
                    
                    missing_col_info.append({
                        '열 번호': col_index,
                        '열명': col,
                        '총 결측치 수': total_missing,
                        '결측치 비율': f"{missing_ratio:.1f}%",
                        '결측 유형': detail_text
                    })
                
                missing_df = pd.DataFrame(missing_col_info)
                # 결측치 열 정보 표에만 CSS 적용
                st.markdown("""
                <style>
                .missing-info-table div[data-testid="stDataFrame"] {
                    margin-top: 0px !important;
                    margin-left: 40px !important;
                }
                </style>
                """, unsafe_allow_html=True)
                
                # 결측치 열 정보 표를 특정 클래스로 감싸기
                st.markdown('<div class="missing-info-table">', unsafe_allow_html=True)
                st.dataframe(missing_df, use_container_width=False, hide_index=True, width=600)
                st.markdown('</div>', unsafe_allow_html=True)
        
        
        # 결측치 제거 레벨 선택
        st.markdown('<h4 style="font-size: 16px; font-style: italic; color: #000; margin-bottom: 0px;">결측치 제거 레벨 선택</h4>', unsafe_allow_html=True)
        
        # 라디오 버튼 상단 여백 제거를 위한 CSS
        st.markdown("""
        <style>
        [data-testid="stRadio"] {
            margin-top: -50px !important;
            padding-top: 0px !important;
            margin-bottom: 22px !important;
        }
        [data-testid="stRadio"] > div {
            margin-top: 0px !important;
            padding-top: 0px !important;
        }
        </style>
        """, unsafe_allow_html=True)
        
        # 세로로 배치된 라디오 버튼 사용
        with st.container():
            removal_level = st.radio(
                "",
                options=[
                    "1단계 : 공백, None, NaN 제거",
                    "2단계 : 1단계 + #DIV/0!, #N/A, #NAME?, #NUM!, #REF!, #VALUE!, #NODATA",
                    "3단계 : 2단계 + #NULL!, #SPILL!, #CALC!, #GETTING_DATA, #FIELD!, #BLOCKED!, #CONNECT!, #BUSY!, #UNKNOWN!"
                ],
                index=1,  # 2단계를 기본값으로 설정
                key="removal_level_radio"
            )
        
        # 처리 방법 선택
        st.markdown('<h4 style="font-size: 16px; font-style: italic; color: #000; margin-bottom: 0px;">결측치 처리 방법 선택</h4>', unsafe_allow_html=True)
        
        # 처리 방법 라디오 버튼 상단 여백 제거를 위한 CSS
        st.markdown("""
        <style>
        [data-testid="stRadio"]:nth-of-type(2) {
            margin-top: -35px !important;
            padding-top: 0px !important;
            margin-bottom: 20px !important;
        }
        [data-testid="stRadio"]:nth-of-type(2) > div {
            margin-top: 0px !important;
            padding-top: 0px !important;
        }
        </style>
        """, unsafe_allow_html=True)
        
        na_action = st.radio(
            "",
            [
                "유지",
                "행 삭제",
                "열 삭제"
                            ]
        )
        
        # 특정 행/열 선택 삭제 기능
        if ("행 삭제" in na_action) and na_count_rows > 0:
            
            # 결측치가 있는 행들의 인덱스 찾기
            rows_with_na_indices = st.session_state['current_df'][st.session_state['current_df'].isnull().any(axis=1)].index.tolist()
            
            # 저장된 날짜 열 정보를 우선적으로 사용
            if 'date_column_name' in st.session_state:
                # 사용자가 직접 날짜형으로 변환한 열 사용
                date_col = st.session_state['date_column_name']
                if date_col in st.session_state['current_df'].columns:
                    # 날짜 열이 있는 경우 날짜 정보로 표시
                    rows_with_na = []
                    for idx in rows_with_na_indices:
                        try:
                            date_value = st.session_state['current_df'].loc[idx, date_col]
                            if pd.notna(date_value):
                                # 날짜 형식으로 표시
                                try:
                                    if pd.api.types.is_datetime64_any_dtype(date_value):
                                        date_str = date_value.strftime('%Y-%m-%d %H:%M:%S')
                                    else:
                                        date_str = str(date_value)
                                    rows_with_na.append(f"{date_str} (행 {idx})")
                                except:
                                    # 날짜 형식 변환 실패 시 행 번호로 표시
                                    rows_with_na.append(f"행 {idx}")
                            else:
                                # 날짜가 없는 경우 행 번호로 표시
                                rows_with_na.append(f"행 {idx}")
                        except:
                            # 오류 발생 시 행 번호로 표시
                            rows_with_na.append(f"행 {idx}")
                else:
                    # 저장된 날짜 열이 현재 데이터에 없는 경우 행 번호로 표시
                    rows_with_na = [f"행 {idx}" for idx in rows_with_na_indices]
            elif 'data_period' in st.session_state and st.session_state['data_period']:
                # 저장된 날짜 열 정보 사용 (기존 방식)
                date_col = st.session_state['data_period']['date_column']
                if date_col in st.session_state['current_df'].columns:
                    # 날짜 열이 있는 경우 날짜 정보로 표시
                    rows_with_na = []
                    for idx in rows_with_na_indices:
                        try:
                            date_value = st.session_state['current_df'].loc[idx, date_col]
                            if pd.notna(date_value):
                                # 날짜 형식으로 표시
                                try:
                                    if pd.api.types.is_datetime64_any_dtype(date_value):
                                        date_str = date_value.strftime('%Y-%m-%d %H:%M:%S')
                                    else:
                                        date_str = str(date_value)
                                    rows_with_na.append(f"{date_str} (행 {idx})")
                                except:
                                    # 날짜 형식 변환 실패 시 행 번호로 표시
                                    rows_with_na.append(f"행 {idx}")
                            else:
                                # 날짜가 없는 경우 행 번호로 표시
                                rows_with_na.append(f"행 {idx}")
                        except:
                            # 오류 발생 시 행 번호로 표시
                            rows_with_na.append(f"행 {idx}")
                else:
                    # 저장된 날짜 열이 현재 데이터에 없는 경우 행 번호로 표시
                    rows_with_na = [f"행 {idx}" for idx in rows_with_na_indices]
            else:
                # 저장된 날짜 정보가 없는 경우 자동 감지
                datetime_cols = st.session_state['current_df'].select_dtypes(include=['datetime']).columns
                if len(datetime_cols) > 0:
                    # 날짜 열이 있는 경우 날짜 정보로 표시
                    date_col = datetime_cols[0]
                    rows_with_na = []
                    for idx in rows_with_na_indices:
                        try:
                            date_value = st.session_state['current_df'].loc[idx, date_col]
                            if pd.notna(date_value):
                                # 날짜 형식으로 표시
                                try:
                                    if pd.api.types.is_datetime64_any_dtype(date_value):
                                        date_str = date_value.strftime('%Y-%m-%d %H:%M:%S')
                                    else:
                                        date_str = str(date_value)
                                    rows_with_na.append(f"{date_str} (행 {idx})")
                                except:
                                    # 날짜 형식 변환 실패 시 행 번호로 표시
                                    rows_with_na.append(f"행 {idx}")
                            else:
                                # 날짜가 없는 경우 행 번호로 표시
                                rows_with_na.append(f"행 {idx}")
                        except:
                            # 오류 발생 시 행 번호로 표시
                            rows_with_na.append(f"행 {idx}")
                else:
                    # 날짜 열이 없는 경우 행 번호로 표시
                    rows_with_na = [f"행 {idx}" for idx in rows_with_na_indices]
            
            # 행 선택 방법 선택
            st.markdown('<p style="font-size: 17px; margin-bottom: 20px; margin-top: -10px; font-style: italic;">행 선택 방법</p>', unsafe_allow_html=True)
            row_selection_method = st.radio(
                "",
                options=["드롭다운에서 선택", "직접 범위 입력"],
                horizontal=True,
                key="row_selection_method"
            )
            
            if row_selection_method == "드롭다운에서 선택":
                rows_to_drop = st.multiselect(
                    "삭제할 행을 선택하세요",
                    options=rows_with_na,
                    default=[],
                    key="rows_multiselect"
                )
            else:
                # 직접 범위 입력
                st.markdown("""
                <div style="background-color: #f8f9fa; padding: 10px; border-radius: 5px; margin: -5px 0;">
                    <p style="margin: 0; font-size: 14px; color: #495057;">
                        <strong>입력 형식:</strong> 개별 행: 5,7,10 | 범위: 5-10, 15~20 | 혼합: 5,7,10-15,44,78~89 (⚠️ 전처리를 거듭하더라도 원본의 행 번호를 기준으로 입력해야 합니다.)
                    </p>
                </div>
                """, unsafe_allow_html=True)
                
                range_input = st.text_input(
                    "삭제할 행 범위를 입력하세요",
                    placeholder="예: 5,7,10-15,44,78~89",
                    key="row_range_input"
                )
                
                # 범위 입력을 파싱하는 함수
                def parse_row_range(range_str):
                    if not range_str.strip():
                        return []
                    
                    rows_to_drop = []
                    parts = range_str.replace(' ', '').split(',')
                    
                    for part in parts:
                        part = part.strip()
                        if not part:
                            continue
                        
                        # 범위 처리 (5-10 또는 5~10)
                        if '-' in part or '~' in part:
                            separator = '-' if '-' in part else '~'
                            try:
                                start, end = part.split(separator)
                                start_idx = int(start.strip())
                                end_idx = int(end.strip())
                                
                                # 실제 데이터 인덱스에 맞게 조정
                                for idx in range(start_idx, end_idx + 1):
                                    if idx in st.session_state['current_df'].index:
                                        rows_to_drop.append(idx)
                            except ValueError:
                                st.warning(f"잘못된 범위 형식: {part}")
                                continue
                        else:
                            # 개별 행 번호
                            try:
                                idx = int(part)
                                if idx in st.session_state['current_df'].index:
                                    rows_to_drop.append(idx)
                                else:
                                    st.warning(f"존재하지 않는 행 번호: {idx}")
                            except ValueError:
                                st.warning(f"잘못된 행 번호: {part}")
                                continue
                    
                    return list(set(rows_to_drop))  # 중복 제거
                
                # 입력된 범위를 실제 행 인덱스로 변환
                if range_input:
                    parsed_rows = parse_row_range(range_input)
                    rows_to_drop = parsed_rows
                    if parsed_rows:
                        st.info(f"🔍 파싱된 행 범위: {len(parsed_rows)}개 행 (첫 5개: {parsed_rows[:5]})")
                else:
                    rows_to_drop = []
        elif ("열 삭제" in na_action) and na_count_cols > 0:
            # 결측치가 있는 열 전체 옵션을 포함한 드롭다운 메뉴
            col_options = ["결측치가 있는 열 전체"] + na_cols
            
            cols_to_drop = st.multiselect(
                "삭제할 열을 선택하세요",
                options=col_options,
                default=[],
                key="cols_multiselect"
            )
        else:
            rows_to_drop = []
            cols_to_drop = []
        
        # UI 요소들의 스타일 조정을 위한 CSS
        st.markdown("""
        <style>
        [data-testid="stTextInput"] {
            width: 100% !important;
        }
        .stRadio > div {
            margin-bottom: 10px !important;
        }
        /* placeholder 텍스트 색상 통일 - 더 포괄적인 선택자 */
        [data-testid="stSelectbox"] div[data-baseweb="select"] span,
        [data-testid="stMultiSelect"] div[data-baseweb="select"] span,
        [data-testid="stSelectbox"] div[data-baseweb="select"] div,
        [data-testid="stMultiSelect"] div[data-baseweb="select"] div,
        [data-testid="stSelectbox"] div[data-baseweb="select"] input::placeholder,
        [data-testid="stMultiSelect"] div[data-baseweb="select"] input::placeholder,
        [data-testid="stSelectbox"] div[data-baseweb="select"] input,
        [data-testid="stMultiSelect"] div[data-baseweb="select"] input {
            color: #666 !important;
        }
        /* 추가적인 placeholder 스타일 */
        [data-testid="stSelectbox"] div[data-baseweb="select"] div[role="option"],
        [data-testid="stMultiSelect"] div[data-baseweb="select"] div[role="option"] {
            color: #666 !important;
        }
        </style>
        """, unsafe_allow_html=True)
        
        # 버튼들을 나란히 배치
        col1, col2, col3 = st.columns(3)
        
        with col1:
            # 버튼 텍스트를 상황에 맞게 변경
            if st.session_state['preprocessing_steps']:
                button_text = "🔄 추가 전처리 수행"
            else:
                button_text = "▶️ 전처리 시작"
            
            if st.button(button_text, help="선택한 설정으로 결측치를 처리합니다"):
                # 제거될 데이터 비율 계산
                df = st.session_state['current_df'].copy()
                total_original = len(df) * len(df.columns)
                
                # 선택된 레벨에 따른 결측치 처리 (임시)
                temp_df = df.copy()
                
                # 1단계: 공백만 제거 (None, NaN, 빈 문자열 모두 포함)
                if "1단계" in removal_level:
                    for col in temp_df.columns:
                        # 빈 문자열을 pd.NA로 변환 (None과 NaN은 이미 pd.NA로 처리됨)
                        temp_df[col] = temp_df[col].replace('', pd.NA)
                elif "2단계" in removal_level:
                    error_values_2 = ['#DIV/0!', '#N/A', '#NAME?', '#NUM!', '#REF!', '#VALUE!', '#NODATA']
                    for col in temp_df.columns:
                        temp_df[col] = temp_df[col].replace('', pd.NA)
                        temp_df[col] = temp_df[col].replace(error_values_2, pd.NA)
                elif "3단계" in removal_level:
                    error_values_2 = ['#DIV/0!', '#N/A', '#NAME?', '#NUM!', '#REF!', '#VALUE!', '#NODATA']
                    error_values_3 = ['#NULL!', '#SPILL!', '#CALC!', '#GETTING_DATA', '#FIELD!', 
                                    '#BLOCKED!', '#CONNECT!', '#BUSY!', '#UNKNOWN!']
                    for col in temp_df.columns:
                        temp_df[col] = temp_df[col].replace('', pd.NA)
                        temp_df[col] = temp_df[col].replace(error_values_2 + error_values_3, pd.NA)
                
                # 처리 방법에 따른 제거 계산
                if "행 삭제" in na_action:
                    # 사용자가 특정 행을 선택한 경우
                    if 'rows_to_drop' in locals() and rows_to_drop and len(rows_to_drop) > 0:
                        # 선택된 행만 삭제하는 경우
                        selected_rows_to_drop = []
                        
                        # 행 선택 방법에 따라 처리
                        if 'row_selection_method' in st.session_state and st.session_state['row_selection_method'] == "직접 범위 입력":
                            # 직접 범위 입력의 경우 이미 실제 인덱스가 들어있음
                            selected_rows_to_drop = [idx for idx in rows_to_drop if idx in temp_df.index]
                        else:
                            # 드롭다운 선택의 경우 텍스트에서 인덱스 추출
                            for row_text in rows_to_drop:
                                if "행 " in row_text:
                                    try:
                                        idx = int(row_text.split("행 ")[1])
                                        if idx in temp_df.index:
                                            selected_rows_to_drop.append(idx)
                                    except (ValueError, IndexError):
                                        continue
                        
                        # 선택된 행만 삭제
                        if selected_rows_to_drop:
                            temp_df = temp_df.drop(index=selected_rows_to_drop)
                            # 선택된 행 수만큼만 제거 비율 계산
                            removal_percentage = (len(selected_rows_to_drop) / len(df)) * 100
                        else:
                            # 선택된 행이 없으면 전체 결측치 기준으로 계산
                            temp_df = temp_df.dropna()
                            total_after = len(temp_df) * len(temp_df.columns)
                            removal_percentage = ((total_original - total_after) / total_original) * 100
                    else:
                        # 특정 행 선택이 없으면 전체 결측치 기준으로 계산
                        temp_df = temp_df.dropna()
                        total_after = len(temp_df) * len(temp_df.columns)
                        removal_percentage = ((total_original - total_after) / total_original) * 100
                elif "열 삭제" in na_action:
                    if cols_to_drop:
                        # "결측치가 있는 열 전체" 옵션이 선택된 경우
                        if "결측치가 있는 열 전체" in cols_to_drop:
                            # 모든 결측치가 있는 열 삭제 (3단계 기준 적용)
                            # 3단계 결측치 마스크 생성 (분포도와 동일한 기준)
                            stage3_mask = pd.DataFrame(False, index=temp_df.index, columns=temp_df.columns)
                            
                            # 1단계: 공백만 (None, NaN, 빈 문자열 모두 포함)
                            for col in temp_df.columns:
                                try:
                                    stage3_mask[col] = (temp_df[col].isna()) | (temp_df[col].astype(str).str.strip() == '')
                                except:
                                    stage3_mask[col] = temp_df[col].isnull()
                            
                            # 2단계: 1단계 + Excel 오류 값들
                            error_values_2 = ['#DIV/0!', '#N/A', '#NAME?', '#NUM!', '#REF!', '#VALUE!', '#NODATA']
                            for col in temp_df.columns:
                                excel_error_mask = temp_df[col].isin(error_values_2)
                                stage3_mask[col] = stage3_mask[col] | excel_error_mask
                            
                            # 3단계: 2단계 + 추가 Excel 오류 값들
                            error_values_3 = ['#NULL!', '#SPILL!', '#CALC!', '#GETTING_DATA', '#FIELD!', 
                                            '#BLOCKED!', '#CONNECT!', '#BUSY!', '#UNKNOWN!']
                            for col in temp_df.columns:
                                additional_error_mask = temp_df[col].isin(error_values_3)
                                stage3_mask[col] = stage3_mask[col] | additional_error_mask
                            
                            # 결측치가 있는 열 찾기
                            cols_with_na = temp_df.columns[stage3_mask.any(axis=0)].tolist()
                            
                            if cols_with_na:
                                temp_df = temp_df.drop(columns=cols_with_na)
                                # 선택된 열 수만큼만 제거 비율 계산
                                removal_percentage = (len(cols_with_na) / len(df.columns)) * 100
                            else:
                                temp_df = temp_df.dropna(axis=1)
                                total_after = len(temp_df) * len(temp_df.columns)
                                removal_percentage = ((total_original - total_after) / total_original) * 100
                        else:
                            existing_cols_to_drop = [col for col in cols_to_drop if col in temp_df.columns]
                            if existing_cols_to_drop:
                                temp_df = temp_df.drop(columns=existing_cols_to_drop)
                                # 선택된 열 수만큼만 제거 비율 계산
                                removal_percentage = (len(existing_cols_to_drop) / len(df.columns)) * 100
                            else:
                                temp_df = temp_df.dropna(axis=1)
                                total_after = len(temp_df) * len(temp_df.columns)
                                removal_percentage = ((total_original - total_after) / total_original) * 100
                    else:
                        temp_df = temp_df.dropna(axis=1)
                        total_after = len(temp_df) * len(temp_df.columns)
                        removal_percentage = ((total_original - total_after) / total_original) * 100

                
                # 30% 이상 제거되는 경우 확인 다이얼로그
                if removal_percentage >= 30:
                    st.markdown(f"""
                    <div style="background-color: #fff3cd; border: 1px solid #ffeaa7; color: #856404; padding: 12px; border-radius: 4px; margin: 10px 0;">
                        ⚠️ <strong>주의</strong>: 선택한 옵션에 따라 전처리를 수행할 경우 기존 데이터 중 <strong>{removal_percentage:.0f}%</strong>가 제거됩니다. 실행하시겠습니까?
                    </div>
                    """, unsafe_allow_html=True)
                    
                    # 세션 상태에 확인 대기 상태 저장
                    st.session_state['waiting_confirmation'] = True
                    st.session_state['removal_percentage'] = removal_percentage
                    st.session_state['temp_df'] = temp_df
                    st.session_state['prev_rows'] = len(st.session_state['current_df'])
                    st.session_state['prev_cols'] = len(st.session_state['current_df'].columns)
                    
                    col_confirm1, col_confirm2 = st.columns(2)
                    with col_confirm1:
                        if st.button("❌ 취소", key="cancel_preprocessing"):
                            st.session_state['waiting_confirmation'] = False
                            st.info("전처리가 취소되었습니다.")
                            st.rerun()
                    with col_confirm2:
                        if st.button("✅ 실행", key="confirm_preprocessing"):
                            # 실제 전처리 수행
                            if 'temp_df' in st.session_state:
                                prev_rows = st.session_state.get('prev_rows', len(st.session_state['current_df']))
                                prev_cols = st.session_state.get('prev_cols', len(st.session_state['current_df'].columns))
                                
                                # temp_df를 current_df로 적용
                                st.session_state['current_df'] = st.session_state['temp_df'].copy()
                                
                                # 결과 표시
                                after_rows = len(st.session_state['current_df'])
                                after_cols = len(st.session_state['current_df'].columns)
                                
                                # 전처리 과정 기록
                                step_description = f"결측치처리 (행: {prev_rows}→{after_rows}, 열: {prev_cols}→{after_cols})"
                                st.session_state['preprocessing_steps'].append(step_description)
                                
                                st.success(f"전처리가 완료되었습니다!")
                                st.write(f"**변경 사항:**")
                                st.write(f"- 행: {prev_rows} → {after_rows} ({prev_rows - after_rows}개 삭제)")
                                st.write(f"- 열: {prev_cols} → {after_cols} ({prev_cols - after_cols}개 삭제)")
                                
                                # 세션 상태 정리
                                if 'temp_df' in st.session_state:
                                    del st.session_state['temp_df']
                                if 'waiting_confirmation' in st.session_state:
                                    del st.session_state['waiting_confirmation']
                                if 'removal_percentage' in st.session_state:
                                    del st.session_state['removal_percentage']
                                if 'prev_rows' in st.session_state:
                                    del st.session_state['prev_rows']
                                if 'prev_cols' in st.session_state:
                                    del st.session_state['prev_cols']
                                
                                # 전처리 완료 후 분포도 업데이트를 위한 플래그 설정 (실제 변경이 있을 때만)
                                if prev_rows != after_rows or prev_cols != after_cols:
                                    st.session_state['update_visualization'] = True
                                # 페이지 새로고침
                                st.rerun()
                            else:
                                st.error("전처리 데이터를 찾을 수 없습니다.")
                                st.session_state['waiting_confirmation'] = False
                                st.rerun()
                else:
                    # 30% 미만인 경우 바로 실행
                    # 실제 전처리 수행
                    prev_rows = len(st.session_state['current_df'])
                    prev_cols = len(st.session_state['current_df'].columns)
                    
                    # 처리 과정 기록을 위한 설명 생성
                    step_description = ""
                    
                    # 현재 데이터프레임 복사
                    df = st.session_state['current_df'].copy()
                    
                    # 1단계: 공백만 제거 (None, NaN, 빈 문자열 모두 포함)
                    if "1단계" in removal_level:
                        # None, NaN, 빈 문자열을 모두 pd.NA로 변환
                        for col in df.columns:
                            # 빈 문자열을 pd.NA로 변환
                            df[col] = df[col].replace('', pd.NA)
                            # None과 NaN은 이미 pd.NA로 처리됨
                    elif "2단계" in removal_level:
                        # 2단계: 1단계 + Excel 오류 값들
                        error_values_2 = ['#DIV/0!', '#N/A', '#NAME?', '#NUM!', '#REF!', '#VALUE!', '#NODATA']
                        for col in df.columns:
                            df[col] = df[col].replace('', pd.NA)  # 공백 제거
                            df[col] = df[col].replace(error_values_2, pd.NA)  # Excel 오류 값들 제거
                    elif "3단계" in removal_level:
                        # 3단계: 2단계 + 추가 Excel 오류 값들
                        error_values_2 = ['#DIV/0!', '#N/A', '#NAME?', '#NUM!', '#REF!', '#VALUE!', '#NODATA']
                        error_values_3 = ['#NULL!', '#SPILL!', '#CALC!', '#GETTING_DATA', '#FIELD!', 
                                        '#BLOCKED!', '#CONNECT!', '#BUSY!', '#UNKNOWN!']
                        for col in df.columns:
                            df[col] = df[col].replace('', pd.NA)  # 공백 제거
                            df[col] = df[col].replace(error_values_2 + error_values_3, pd.NA)  # 모든 Excel 오류 값들 제거
                    
                    # 업데이트된 데이터프레임을 session state에 저장
                    st.session_state['current_df'] = df
                    
                    if "행 삭제" in na_action:
                        # 행 삭제 로직 - 특정 행 선택 삭제 또는 전체 삭제
                        if 'rows_to_drop' in locals() and rows_to_drop and len(rows_to_drop) > 0:
                            # 선택된 행만 삭제
                            existing_rows_to_drop = []
                            
                            # 행 선택 방법에 따라 처리
                            if 'row_selection_method' in st.session_state and st.session_state['row_selection_method'] == "직접 범위 입력":
                                # 직접 범위 입력의 경우 이미 실제 인덱스가 들어있음
                                existing_rows_to_drop = [idx for idx in rows_to_drop if idx in st.session_state['current_df'].index]
                            else:
                                # 드롭다운 선택의 경우 텍스트에서 인덱스 추출
                                for row_text in rows_to_drop:
                                    # "행 {idx}" 또는 "{date_str} (행 {idx})" 형식에서 인덱스 추출
                                    if "행 " in row_text:
                                        try:
                                            # "행 123" 형식에서 숫자 추출
                                            idx = int(row_text.split("행 ")[1].split(")")[0])
                                            if idx in st.session_state['current_df'].index:
                                                existing_rows_to_drop.append(idx)
                                        except:
                                            pass
                                    elif " (행 " in row_text:
                                        try:
                                            # "2024-01-01 12:00:00 (행 123)" 형식에서 숫자 추출
                                            idx = int(row_text.split(" (행 ")[1].split(")")[0])
                                            if idx in st.session_state['current_df'].index:
                                                existing_rows_to_drop.append(idx)
                                        except:
                                            pass
                            
                            if existing_rows_to_drop:
                                st.session_state['current_df'] = st.session_state['current_df'].drop(index=existing_rows_to_drop)
                                step_description += f"특정행삭제({len(existing_rows_to_drop)}개)"
                                # 행 삭제 후 분포도 강제 업데이트
                                st.session_state['update_visualization'] = True
                                st.info(f"✅ {len(existing_rows_to_drop)}개 행이 삭제되었습니다. 분포도가 업데이트됩니다.")
                            else:
                                st.warning("선택한 행들이 이미 삭제되었거나 존재하지 않습니다.")
                        else:
                            # 특정 행이 선택되지 않은 경우 전체 삭제
                            st.session_state['current_df'] = st.session_state['current_df'].dropna()
                            step_description += f"행 삭제(전체)"
                    
                    if "열 삭제" in na_action:
                        # 열 삭제 로직 - 특정 열 선택 삭제 또는 전체 삭제
                        if cols_to_drop and len(cols_to_drop) > 0:
                            # "결측치가 있는 열 전체" 옵션이 선택된 경우
                            if "결측치가 있는 열 전체" in cols_to_drop:
                                # 모든 결측치가 있는 열 삭제 (3단계 기준 적용)
                                df = st.session_state['current_df']
                                
                                # 3단계 결측치 마스크 생성 (분포도와 동일한 기준)
                                stage3_mask = pd.DataFrame(False, index=df.index, columns=df.columns)
                                
                                # 1단계: 공백만 (None, NaN, 빈 문자열 모두 포함)
                                for col in df.columns:
                                    try:
                                        stage3_mask[col] = (df[col].isna()) | (df[col].astype(str).str.strip() == '')
                                    except:
                                        stage3_mask[col] = df[col].isnull()
                                
                                # 2단계: 1단계 + Excel 오류 값들
                                error_values_2 = ['#DIV/0!', '#N/A', '#NAME?', '#NUM!', '#REF!', '#VALUE!', '#NODATA']
                                for col in df.columns:
                                    excel_error_mask = df[col].isin(error_values_2)
                                    stage3_mask[col] = stage3_mask[col] | excel_error_mask
                                
                                # 3단계: 2단계 + 추가 Excel 오류 값들
                                error_values_3 = ['#NULL!', '#SPILL!', '#CALC!', '#GETTING_DATA', '#FIELD!', 
                                                '#BLOCKED!', '#CONNECT!', '#BUSY!', '#UNKNOWN!']
                                for col in df.columns:
                                    additional_error_mask = df[col].isin(error_values_3)
                                    stage3_mask[col] = stage3_mask[col] | additional_error_mask
                                
                                # 결측치가 있는 열 찾기
                                cols_with_na = df.columns[stage3_mask.any(axis=0)].tolist()
                                
                                if cols_with_na:
                                    st.session_state['current_df'] = st.session_state['current_df'].drop(columns=cols_with_na)
                                    step_description += f"결측치열전체삭제({len(cols_with_na)}개)"
                                else:
                                    st.warning("결측치가 있는 열이 없습니다.")
                            else:
                                # 선택된 열만 삭제
                                existing_cols_to_drop = [col for col in cols_to_drop if col in st.session_state['current_df'].columns]
                                if existing_cols_to_drop:
                                    st.session_state['current_df'] = st.session_state['current_df'].drop(columns=existing_cols_to_drop)
                                    step_description += f"특정열삭제({len(existing_cols_to_drop)}개)"
                                else:
                                    st.warning("선택한 열들이 이미 삭제되었거나 존재하지 않습니다.")
                        else:
                            # 특정 열이 선택되지 않은 경우 전체 삭제
                            st.session_state['current_df'] = st.session_state['current_df'].dropna(axis=1)
                            step_description += f"열 삭제(전체)"
                    

                    
                    # 결과 표시
                    after_rows = len(st.session_state['current_df'])
                    after_cols = len(st.session_state['current_df'].columns)
                    
                    if "삭제" in na_action:
                        # 전처리 과정 기록
                        step_description += f" (행: {prev_rows}→{after_rows}, 열: {prev_cols}→{after_cols})"
                        st.session_state['preprocessing_steps'].append(step_description)
                        
                        st.success(f"결측치 처리가 완료되었습니다!")
                        st.write(f"**변경 사항:**")
                        st.write(f"- 행: {prev_rows} → {after_rows} ({prev_rows - after_rows}개 삭제)")
                        st.write(f"- 열: {prev_cols} → {after_cols} ({prev_cols - after_cols}개 삭제)")
                        
                        # 전처리 완료 후 분포도 업데이트를 위한 플래그 설정 (실제 변경이 있을 때만)
                        if prev_rows != after_rows or prev_cols != after_cols:
                            st.session_state['update_visualization'] = True
                        # 페이지 새로고침을 위한 rerun
                        st.rerun()
                    else:
                        st.info(f"결측치가 있는 데이터를 유지합니다. (총 {prev_rows}행, {prev_cols}열)")
            

        
        with col2:
            if st.button("✅ 전처리 완료", help="전처리 완료 버튼을 누르면 분석 수행 단계로 넘어갑니다."):
                # 최종 결과를 메인 데이터프레임에 적용
                st.session_state['df'] = st.session_state['current_df'].copy()
                st.session_state['preprocessing_completed'] = True  # 전처리 완료 상태 설정
                # 전처리 완료 후 분포도 업데이트를 위한 플래그 설정
                st.session_state['update_visualization'] = True
                st.info("전처리가 완료되어 분석 데이터가 업데이트되었습니다.")
                if st.session_state['data_file_path']:
                    try:
                        if st.session_state['data_file_path'].endswith('.csv'):
                            st.session_state['df'].to_csv(st.session_state['data_file_path'], index=False)
                        else:
                            st.session_state['df'].to_excel(st.session_state['data_file_path'], index=False)
                    except PermissionError:
                        st.warning("⚠️ 파일이 다른 프로그램에서 사용 중이어서 저장할 수 없습니다.")
                    except Exception as e:
                        st.warning(f"⚠️ 파일 저장 중 오류가 발생했습니다: {str(e)}")
                st.rerun()
        
        with col3:
            if st.button("🔄 전처리 초기화", help="모든 전처리 과정을 초기화하고 원본 데이터로 되돌립니다"):
                # 전처리 과정 초기화 - 원본 데이터로 복원
                if 'original_df' in st.session_state:
                    st.session_state['current_df'] = st.session_state['original_df'].copy()
                    st.session_state['df'] = st.session_state['original_df'].copy()
                else:
                    st.session_state['current_df'] = st.session_state['df'].copy()
                st.session_state['preprocessing_steps'] = []
                # 전처리 초기화 후 분포도 업데이트를 위한 플래그 설정
                st.session_state['update_visualization'] = True
                st.info("전처리 과정이 초기화되었습니다.")
                st.rerun()
    
    if current_na_count == 0:
        # 결측치가 없는 경우 - 완료 상태 표시
        st.success("✅ 결측치 처리가 완료되었습니다!")
        
        # 전처리 완료 안내 메시지
        st.info("💡 위의 '전처리 완료' 버튼을 클릭하여 다음 단계로 진행하세요.")

    # 이후 단계는 전처리가 완료된 후에만 표시
    if st.session_state['df'] is not None and st.session_state['preprocessing_completed']:
        # 3. 분석 대상 변수 선택 및 데이터 확인 (파일 업로드/전처리 이후에만 노출)
        st.markdown('<hr style="border:3px solid #333;">', unsafe_allow_html=True)
        st.header("3. 분석 대상 변수 선택 및 데이터 확인")
        st.subheader("분석에 사용할 변수 선택")
        numeric_cols = st.session_state['df'].select_dtypes(include=[np.number]).columns.tolist()
        if len(numeric_cols) < 2:
            st.warning("수치형 컬럼이 2개 이상 필요합니다.")
            x_cols, y_col = [], None
        else:
            # Y 선택과 X 선택 드롭다운 박스 (화면 가로 길이의 절반으로 설정)
            col1, col2 = st.columns([1, 1])
            with col1:
                y_col = st.selectbox("종속변수(Y) 선택", options=["변수를 선택하세요"] + numeric_cols, index=0, key="y_col")
            with col2:
                # X 선택: Y로 선택된 변수는 제외
                x_options = [col for col in numeric_cols if col != y_col] if y_col else numeric_cols
                x_cols = st.multiselect("독립변수(X) 선택 (2~5개 추천)", options=x_options, max_selections=5, key="x_cols", placeholder="변수들을 선택하세요")
            
            if y_col == "변수를 선택하세요":
                y_col = None
            if x_cols and y_col:
                # X, Y 모두 결측치가 없는 행만 남김
                df_no_na = st.session_state['df'][[y_col] + x_cols].dropna()
                # 변수 목록을 쉼표로 구분하여 한 줄에 나열
                x_vars_text = ", ".join(x_cols)
                st.success(f"**분석에 사용할 변수**\n\n**독립변수 (X):** {x_vars_text}\n\n**종속변수 (Y):** {y_col}")
                st.markdown('<span style="font-weight:bold; color:#1f77b4; font-size:18px;">기초 통계</span>', unsafe_allow_html=True)
                st.dataframe(df_no_na.describe())
                st.markdown('<span style="font-weight:bold; color:#ff7f0e; font-size:18px;">히스토그램</span>', unsafe_allow_html=True)
                # 히스토그램 2개씩 한 줄에 배치
                hist_cols = [y_col] + x_cols
                histogram_charts = []  # 차트 저장용 리스트
                
                for i in range(0, len(hist_cols), 2):
                    cols = st.columns(2)
                    for j in range(2):
                        if i + j < len(hist_cols):
                            col = hist_cols[i + j]
                            with cols[j]:
                                fig = px.histogram(df_no_na, x=col, nbins=30, title=f"{col} 분포")
                                st.plotly_chart(fig, use_container_width=True)
                                histogram_charts.append(fig)  # 차트 저장
                
                # 히스토그램 차트들을 session state에 저장
                st.session_state['histogram_charts'] = histogram_charts
                # 상관계수 행렬과 상관관계 분석을 좌우로 나란히 배치
                col1, col2 = st.columns(2)
                
                with col1:
                    st.markdown('<span style="font-weight:bold; color:#2ca02c; font-size:18px;">상관계수 행렬</span>', unsafe_allow_html=True)
                    corr = df_no_na.corr()
                    fig = ff.create_annotated_heatmap(
                        z=corr.values,
                        x=list(corr.columns),
                        y=list(corr.index)[::-1],  # 세로 방향 순서 뒤집기
                        annotation_text=np.round(corr.values, 2),
                        colorscale='RdBu', showscale=True  # Red-Blue 색상 스케일로 변경
                    )
                    fig.update_layout(
                        coloraxis=dict(
                            colorscale='RdBu',  # Red-Blue 색상 스케일
                            cmin=-1,  # 음의 상관관계도 표시
                            cmax=1,
                            colorbar=dict(title="상관계수")
                        ),
                        height=500,  # 높이 조정
                        width=400    # 너비 조정
                    )
                    st.plotly_chart(fig, use_container_width=True)
                    # 상관관계 히트맵을 session state에 저장
                    st.session_state['correlation_heatmap'] = fig
                
                with col2:
                    # 상관계수 행렬에서 0.4 이상인 쌍 찾기
                    high_corr_pairs = []
                    for i in range(len(corr.columns)):
                        for j in range(i+1, len(corr.columns)):  # 대각선 위쪽만 확인 (중복 방지)
                            corr_value = corr.iloc[i, j]
                            if abs(corr_value) >= 0.4:
                                high_corr_pairs.append({
                                    '변수1': corr.columns[i],
                                    '변수2': corr.columns[j],
                                    '상관계수': corr_value,
                                    '강도': '강한 양의 상관관계' if corr_value >= 0.7 else 
                                           '중간 양의 상관관계' if corr_value >= 0.4 else
                                           '강한 음의 상관관계' if corr_value <= -0.7 else
                                           '중간 음의 상관관계'
                                })
                    
                    # 상관계수 기준으로 정렬 (절댓값 기준 내림차순)
                    high_corr_pairs.sort(key=lambda x: abs(x['상관계수']), reverse=True)
                    
                    if high_corr_pairs:
                        # 결과를 데이터프레임으로 변환
                        high_corr_df = pd.DataFrame(high_corr_pairs)
                        
                        # 상관계수를 소수점 2째자리로 반올림하고 표시 형식 제한
                        high_corr_df['상관계수'] = high_corr_df['상관계수'].round(2).apply(lambda x: f"{x:.2f}")
                        
                        # 상관계수 값을 색상으로 구분
                        def color_correlation(val):
                            if val >= 0.7:
                                return 'background-color: #ffcccc; color: #cc0000; font-weight: bold;'  # 강한 양의 상관관계
                            elif val >= 0.4:
                                return 'background-color: #ffe6cc; color: #cc6600; font-weight: bold;'  # 중간 양의 상관관계
                            elif val <= -0.7:
                                return 'background-color: #ccf2ff; color: #0066cc; font-weight: bold;'  # 강한 음의 상관관계
                            else:
                                return 'background-color: #e6f3ff; color: #0066cc; font-weight: bold;'  # 중간 음의 상관관계
                        
                        # 스타일 적용
                        styled_df = high_corr_df.style.applymap(
                            lambda x: color_correlation(float(x)) if isinstance(x, str) and x.replace('.', '').replace('-', '').isdigit() else '', 
                            subset=['상관계수']
                        )
                        
                        st.write(f"**총 {len(high_corr_pairs)}개의 변수 쌍에서 유의미한 상관관계가 발견되었습니다:**")
                        st.dataframe(styled_df, use_container_width=True, hide_index=True)
                        
                        # 요약 정보
                        st.markdown("**📊 상관관계 요약:**")
                        strong_positive = len([p for p in high_corr_pairs if p['상관계수'] >= 0.7])
                        moderate_positive = len([p for p in high_corr_pairs if 0.4 <= p['상관계수'] < 0.7])
                        strong_negative = len([p for p in high_corr_pairs if p['상관계수'] <= -0.7])
                        moderate_negative = len([p for p in high_corr_pairs if -0.7 < p['상관계수'] <= -0.4])
                        
                        # 2x2 그리드로 메트릭 배치
                        metric_col1, metric_col2 = st.columns(2)
                        with metric_col1:
                            st.metric("강한 양의 상관관계 (≥0.7)", strong_positive)
                            st.metric("강한 음의 상관관계 (≤-0.7)", strong_negative)
                        with metric_col2:
                            st.metric("중간 양의 상관관계 (0.4~0.7)", moderate_positive)
                            st.metric("중간 음의 상관관계 (-0.7~-0.4)", moderate_negative)
                        
                    else:
                        st.info("상관계수가 0.4 이상인 변수 쌍이 없습니다.")
                    
                    # 높은 상관관계 쌍을 session state에 저장
                    st.session_state['high_correlation_pairs'] = high_corr_pairs
                
                st.markdown('<span style="font-weight:bold; color:#d62728; font-size:18px;">산점도 행렬 (pairplot)</span>', unsafe_allow_html=True)
                fig = px.scatter_matrix(df_no_na, dimensions=[y_col]+x_cols)
                
                # 그림 크기와 글자 크기 조정
                fig.update_layout(
                    height=600,  # 세로 길이 증가
                    width=800,   # 가로 길이 설정
                    font=dict(size=10),  # 전체 글자 크기 축소
                    margin=dict(l=80, r=80, t=80, b=80)  # 여백 조정
                )
                
                # y축 제목 글자 크기 추가 축소
                fig.update_yaxes(title_font_size=8)
                fig.update_xaxes(title_font_size=8)
                
                st.plotly_chart(fig, use_container_width=True)
                # 산점도 차트를 session state에 저장
                st.session_state['scatter_charts'] = [fig]
                
                # 산점도 행렬을 통한 Insight 분석
                st.markdown('<span style="font-weight:bold; color:#9467bd; font-size:18px;">💡 산점도 행렬 Insight 분석</span>', unsafe_allow_html=True)
                
                # 상관계수와 분포 특성을 기반으로 한 insight 생성
                insights = []
                
                # 1. 선형 관계 분석
                linear_relationships = []
                for i, col1 in enumerate([y_col] + x_cols):
                    for j, col2 in enumerate([y_col] + x_cols):
                        if i < j:  # 중복 방지
                            corr_val = corr.loc[col1, col2]
                            if abs(corr_val) >= 0.7:
                                linear_relationships.append({
                                    'var1': col1,
                                    'var2': col2,
                                    'corr': corr_val,
                                    'strength': '매우 강한 선형 관계'
                                })
                            elif abs(corr_val) >= 0.4:
                                linear_relationships.append({
                                    'var1': col1,
                                    'var2': col2,
                                    'corr': corr_val,
                                    'strength': '보통 정도의 관계'
                                })
                
                # 2. 분포 특성 분석
                distribution_insights = []
                for col in [y_col] + x_cols:
                    data = df_no_na[col]
                    skewness = data.skew()
                    kurtosis = data.kurtosis()
                    
                    # 왜도 분석
                    if abs(skewness) > 1:
                        if skewness > 0:
                            distribution_insights.append({
                                'var': col,
                                'type': '오른쪽으로 치우친 분포',
                                'value': f'왜도: {skewness:.2f}'
                            })
                        else:
                            distribution_insights.append({
                                'var': col,
                                'type': '왼쪽으로 치우친 분포',
                                'value': f'왜도: {skewness:.2f}'
                            })
                    
                    # 첨도 분석
                    if kurtosis > 3:
                        distribution_insights.append({
                            'var': col,
                            'type': '뾰족한 모양의 분포',
                            'value': f'첨도: {kurtosis:.2f}'
                        })
                    elif kurtosis < 1:
                        distribution_insights.append({
                            'var': col,
                            'type': '넓게 퍼진 분포',
                            'value': f'첨도: {kurtosis:.2f}'
                        })
                
                # 3. 비선형 관계 탐지
                nonlinear_insights = []
                for i, col1 in enumerate([y_col] + x_cols):
                    for j, col2 in enumerate([y_col] + x_cols):
                        if i < j:
                            corr_val = corr.loc[col1, col2]
                            # 상관계수가 낮지만 산점도에서 패턴이 있을 수 있는 경우
                            if abs(corr_val) < 0.3:
                                # 여기서는 간단한 휴리스틱으로 판단
                                nonlinear_insights.append({
                                    'var1': col1,
                                    'var2': col2,
                                    'description': '직선 관계는 약하지만 다른 형태의 관계가 있을 수 있습니다'
                                })
                                break
                
                # Insight를 3개 컬럼으로 나누어 표시
                col1, col2, col3 = st.columns(3)
                
                with col1:
                    if linear_relationships:
                        st.markdown("""
                        <div style="background: linear-gradient(135deg, #667eea 0%, #764ba2 100%); padding: 8px; border-radius: 10px; margin: 10px 0; color: white;">
                            <h4 style="margin: 0; text-align: center;">🔗 선형 관계</h4>
                        </div>
                        """, unsafe_allow_html=True)
                        for rel in linear_relationships[:3]:  # 상위 3개만 표시
                            st.markdown(f"""
                            <div style="background-color: rgba(255,255,255,0.15); padding: 15px; border-radius: 8px; margin: 10px 0; border: 2px solid rgba(255,255,255,0.3); box-shadow: 0 2px 8px rgba(0,0,0,0.1);">
                                <div style="display: flex; justify-content: center; align-items: center; margin-bottom: 10px;">
                                    <span style="background-color: #fff; color: #667eea; padding: 6px 12px; border-radius: 0px; font-weight: bold; font-size: 14px; border: 1px solid #ddd;">{rel['var1']}</span>
                                    <span style="margin: 0 12px; font-size: 18px; color: #333; font-weight: bold;">⟷</span>
                                    <span style="background-color: #fff; color: #667eea; padding: 6px 12px; border-radius: 0px; font-weight: bold; font-size: 14px; border: 1px solid #ddd;">{rel['var2']}</span>
                                </div>
                                <div style="font-size: 16px; font-weight: 500; text-align: center; color: #333;">{rel['strength']} (r={rel['corr']:.3f})</div>
                            </div>
                            """, unsafe_allow_html=True)
                
                with col2:
                    if distribution_insights:
                        st.markdown("""
                        <div style="background: linear-gradient(135deg, #f093fb 0%, #f5576c 100%); padding: 8px; border-radius: 10px; margin: 10px 0; color: white;">
                            <h4 style="margin: 0; text-align: center;">📊 분포 특성</h4>
                        </div>
                        """, unsafe_allow_html=True)
                        for insight in distribution_insights[:3]:  # 상위 3개만 표시
                            st.markdown(f"""
                            <div style="background-color: rgba(255,255,255,0.15); padding: 15px; border-radius: 8px; margin: 10px 0; border: 2px solid rgba(255,255,255,0.3); box-shadow: 0 2px 8px rgba(0,0,0,0.1);">
                                <div style="display: flex; justify-content: center; align-items: center; margin-bottom: 10px;">
                                    <span style="background-color: #fff; color: #f093fb; padding: 6px 12px; border-radius: 0px; font-weight: bold; font-size: 14px; border: 1px solid #ddd;">{insight['var']}</span>
                                </div>
                                <div style="font-size: 16px; font-weight: 500; text-align: center; color: #333;">{insight['type']} ({insight['value']})</div>
                            </div>
                            """, unsafe_allow_html=True)
                
                with col3:
                    if nonlinear_insights:
                        st.markdown("""
                        <div style="background: linear-gradient(135deg, #4facfe 0%, #00f2fe 100%); padding: 8px; border-radius: 10px; margin: 10px 0; color: white;">
                            <h4 style="margin: 0; text-align: center;">🔄 비선형 관계</h4>
                        </div>
                        """, unsafe_allow_html=True)
                        for insight in nonlinear_insights[:2]:  # 상위 2개만 표시
                            st.markdown(f"""
                            <div style="background-color: rgba(255,255,255,0.15); padding: 15px; border-radius: 8px; margin: 10px 0; border: 2px solid rgba(255,255,255,0.3); box-shadow: 0 2px 8px rgba(0,0,0,0.1);">
                                <div style="display: flex; justify-content: center; align-items: center; margin-bottom: 10px;">
                                    <span style="background-color: #fff; color: #4facfe; padding: 6px 12px; border-radius: 0px; font-weight: bold; font-size: 14px; border: 1px solid #ddd;">{insight['var1']}</span>
                                    <span style="margin: 0 12px; font-size: 18px; color: #333; font-weight: bold;">⟷</span>
                                    <span style="background-color: #fff; color: #4facfe; padding: 6px 12px; border-radius: 0px; font-weight: bold; font-size: 14px; border: 1px solid #ddd;">{insight['var2']}</span>
                                </div>
                                <div style="font-size: 16px; font-weight: 500; text-align: center; color: #333;">{insight['description']}</div>
                            </div>
                            """, unsafe_allow_html=True)
                    else:
                        st.markdown("""
                        <div style="background: linear-gradient(135deg, #4facfe 0%, #00f2fe 100%); padding: 15px; border-radius: 10px; margin: 10px 0; color: white;">
                            <h4 style="margin: 0 0 10px 0; text-align: center;">📈 데이터 품질</h4>
                        </div>
                        """, unsafe_allow_html=True)
                        total_vars = len([y_col] + x_cols)
                        if total_vars <= 3:
                            st.markdown("• 분석할 변수가 적어 다양한 관계를 찾기 어려울 수 있습니다")
                        elif total_vars >= 8:
                            st.markdown("• 변수가 많아 복잡한 상호작용이 있을 수 있습니다")
                        
                        if len(df_no_na) < 100:
                            st.markdown("• 데이터가 적어 결과의 신뢰도에 주의가 필요합니다")
                
                # 요약 박스
                st.markdown("""
                <div style="background: linear-gradient(135deg, #fa709a 0%, #fee140 100%); padding: 15px; border-radius: 10px; margin: 20px 0; color: white;">
                    <h4 style="margin: 0; text-align: center;">🎯 핵심 분석 결과</h4>
                    <p style="margin: 5px 0; text-align: center; font-size: 14px;">
                        위의 그래프들을 통해 변수들 간의 관계, 데이터 분포 모양, 특이한 데이터 값들을 한눈에 파악할 수 있습니다. 
                        이를 통해 공정 운영에 중요한 패턴을 발견할 수 있습니다.
                    </p>
                </div>
                """, unsafe_allow_html=True)
                
                # 4. 데이터 분석 및 결과 표출 (변수 선택이 완료된 이후에만 노출)
                st.markdown('<hr style="border:3px solid #333;">', unsafe_allow_html=True)
                st.header("4. 데이터 분석 결과")
                
                # 선형회귀분석
                st.markdown("""
                <div style="background: linear-gradient(135deg, #667eea 0%, #764ba2 100%); padding: 15px; border-radius: 15px; margin: 20px 0; color: white;">
                    <h3 style="margin: 0; font-size: 24px; text-align: center;">📊 선형회귀분석 결과</h3>
                </div>
                """, unsafe_allow_html=True)
                
                X = df_no_na[x_cols]
                y = df_no_na[y_col]
                X_train, X_test, y_train, y_test = train_test_split(X, y, test_size=0.2, random_state=42)
                lr = LinearRegression()
                lr.fit(X_train, y_train)
                y_pred = lr.predict(X_test)
                
                # 수식 생성 (Y에 대한 비중이 높은 순서대로 정렬)
                # 실제 데이터에서 각 항의 기여도 계산
                X_sample = X_test.iloc[:10]  # 테스트 데이터의 처음 10개 샘플 사용
                y_sample = y_test.iloc[:10]
                
                # 각 항의 기여도 계산
                term_contributions = {}
                for col in x_cols:
                    term_values = lr.coef_[x_cols.index(col)] * X_sample[col]
                    term_contributions[col] = term_values
                
                # 절편 기여도
                intercept_contribution = np.full(len(X_sample), lr.intercept_)
                term_contributions['절편'] = intercept_contribution
                
                # 예측값
                y_pred_sample = lr.predict(X_sample)
                
                # 각 항의 평균 기여도와 비중 계산
                term_analysis = []
                for term_name, contributions in term_contributions.items():
                    avg_contribution = np.mean(contributions)
                    contribution_ratio = (avg_contribution / np.mean(y_pred_sample)) * 100
                    
                    if term_name == '절편':
                        term_analysis.append({
                            'name': term_name,
                            'avg_contribution': avg_contribution,
                            'ratio': contribution_ratio,
                            'coef': lr.intercept_
                        })
                    else:
                        coef = lr.coef_[x_cols.index(term_name)]
                        term_analysis.append({
                            'name': term_name,
                            'avg_contribution': avg_contribution,
                            'ratio': contribution_ratio,
                            'coef': coef
                        })
                
                # Y에 대한 비중이 높은 순서대로 정렬 (절편 제외)
                non_intercept_terms = [term for term in term_analysis if term['name'] != '절편']
                non_intercept_terms.sort(key=lambda x: abs(x['ratio']), reverse=True)
                
                # 절편을 마지막에 추가
                intercept_term = [term for term in term_analysis if term['name'] == '절편'][0]
                sorted_terms = non_intercept_terms + [intercept_term]
                
                # 수식 생성
                equation_parts = []
                for i, term in enumerate(sorted_terms):
                    coef = term['coef']
                    ratio = term['ratio']
                    
                    # 계수의 유효숫자 결정
                    if abs(coef) >= 0.001:
                        coef_str = f"{coef:.3f}"
                    elif abs(coef) >= 0.0001:
                        coef_str = f"{coef:.4f}"
                    elif abs(coef) >= 0.00001:
                        coef_str = f"{coef:.5f}"
                    else:
                        coef_str = f"{coef:.2e}"
                    
                    # 첫 번째 항이 아닌 경우 부호 추가 (+ 또는 -)
                    if i > 0:
                        if coef >= 0:
                            if term['name'] == '절편':
                                equation_parts.append(f"&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;+ {coef_str} <span style='color: #6c757d;'>(절편)</span> <span style='color: #6c757d; font-size: 14px;'>({ratio:.1f}%)</span>")
                            else:
                                equation_parts.append(f"&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;+ {coef_str} × {term['name']} <span style='color: #6c757d; font-size: 14px;'>({ratio:.1f}%)</span>")
                        else:
                            # 음수 계수도 동일한 유효숫자 규칙 적용
                            abs_coef = abs(coef)
                            if abs_coef >= 0.001:
                                abs_coef_str = f"{abs_coef:.3f}"
                            elif abs_coef >= 0.0001:
                                abs_coef_str = f"{abs_coef:.4f}"
                            elif abs_coef >= 0.00001:
                                abs_coef_str = f"{abs_coef:.5f}"
                            else:
                                abs_coef_str = f"{abs_coef:.2e}"
                            
                            if term['name'] == '절편':
                                equation_parts.append(f"&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;- {abs_coef_str} <span style='color: #6c757d;'>(절편)</span> <span style='color: #6c757d; font-size: 14px;'>({ratio:.1f}%)</span>")
                            else:
                                equation_parts.append(f"&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;- {abs_coef_str} × {term['name']} <span style='color: #6c757d; font-size: 14px;'>({ratio:.1f}%)</span>")
                    else:
                        if coef >= 0:
                            if term['name'] == '절편':
                                equation_parts.append(f"&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;{coef_str} <span style='color: #6c757d;'>(절편)</span> <span style='color: #6c757d; font-size: 14px;'>({ratio:.1f}%)</span>")
                            else:
                                equation_parts.append(f"&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;{coef_str} × {term['name']} <span style='color: #6c757d; font-size: 14px;'>({ratio:.1f}%)</span>")
                        else:
                            # 음수 계수도 동일한 유효숫자 규칙 적용
                            abs_coef = abs(coef)
                            if abs_coef >= 0.001:
                                abs_coef_str = f"{abs_coef:.3f}"
                            elif abs_coef >= 0.0001:
                                abs_coef_str = f"{abs_coef:.4f}"
                            elif abs_coef >= 0.00001:
                                abs_coef_str = f"{abs_coef:.5f}"
                            else:
                                abs_coef_str = f"{abs_coef:.2e}"
                            
                            if term['name'] == '절편':
                                equation_parts.append(f"&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;- {abs_coef_str} <span style='color: #6c757d;'>(절편)</span> <span style='color: #6c757d; font-size: 14px;'>({ratio:.1f}%)</span>")
                            else:
                                equation_parts.append(f"&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;- {abs_coef_str} × {term['name']} <span style='color: #6c757d; font-size: 14px;'>({ratio:.1f}%)</span>")
                
                # Y 변수명을 괄호 안에 표시
                y_variable_name = y_col if y_col else "Target Variable"
                equation = f"Y ({y_variable_name}) = <br>" + "<br>".join(equation_parts)
                
                # 결과 표시
                col1, col2 = st.columns([1, 1])
                
                with col1:
                    st.markdown("""
                    <div style="background-color: #f8f9fa; padding: 20px; border-radius: 10px; border-left: 5px solid #dc3545;">
                        <h4 style="margin: 0 0 15px 0; color: #dc3545;">🎯 선형 상관관계 방정식</h4>
                        <div style="font-family: 'Arial', 'Helvetica', sans-serif; font-size: 16px; font-weight: bold; color: #333; line-height: 1.5;">
                            {}
                        </div>
                        <div style="margin-top: 10px; padding: 8px; background-color: #e3f2fd; border-radius: 5px; border-left: 3px solid #2196f3;">
                            <small style="color: #1976d2; font-size: 12px;">💡 <strong>설명:</strong> 괄호 안의 숫자는 각 변수가 Y값에 미치는 상대적 비중(%)을 나타냅니다.<br>&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;예를 들어 (11.6%)는 해당 변수가 Y값 변화에 11.6%의 영향을 미친다는 의미입니다.</small>
                        </div>
                    </div>
                    """.format(equation), unsafe_allow_html=True)
                    
                    # 계수 해석
                    st.markdown("""
                    <div style="background-color: #fff3cd; padding: 15px; border-radius: 10px; border-left: 5px solid #ffc107; margin-top: 15px;">
                        <h4 style="margin: 0 0 10px 0; color: #856404;">📈 계수 해석</h4>
                    """, unsafe_allow_html=True)
                    
                    for col, coef in zip(x_cols, lr.coef_):
                        if coef > 0:
                            interpretation = f"<span style='color: #28a745;'>📈 양의 영향</span>"
                        else:
                            interpretation = f"<span style='color: #dc3545;'>📉 음의 영향</span>"
                        
                        # 계수의 유효숫자 결정
                        if abs(coef) >= 0.001:
                            coef_display = f"{coef:.3f}"
                        elif abs(coef) >= 0.0001:
                            coef_display = f"{coef:.4f}"
                        elif abs(coef) >= 0.00001:
                            coef_display = f"{coef:.5f}"
                        else:
                            coef_display = f"{coef:.2e}"
                        
                        st.markdown(f"""
                        <div style="margin: 5px 0; padding: 8px; background-color: white; border-radius: 5px; border-left: 4px solid #ffc107;">
                            <strong>{col}:</strong> {coef_display} ({interpretation})
                            <br><small style="color: #6c757d;">→ {col}이 1단위 증가할 때 Y가 {coef_display}단위 변화</small>
                        </div>
                        """, unsafe_allow_html=True)
                    
                    st.markdown("</div>", unsafe_allow_html=True)
                    
                    # 항 해석 (각 항이 Y에 미치는 실제 비중 분석)
                    st.markdown("""
                    <div style="background-color: #e8f5e8; padding: 15px; border-radius: 10px; border-left: 5px solid #28a745; margin-top: 15px;">
                        <h4 style="margin: 0 0 10px 0; color: #155724;">🎯 항 해석 (실제 비중 분석)</h4>
                    """, unsafe_allow_html=True)
                    
                    # 실제 데이터에서 각 항의 기여도 계산
                    X_sample = X_test.iloc[:10]  # 테스트 데이터의 처음 10개 샘플 사용
                    y_sample = y_test.iloc[:10]
                    
                    # 각 항의 기여도 계산
                    term_contributions = {}
                    for col in x_cols:
                        term_values = lr.coef_[x_cols.index(col)] * X_sample[col]
                        term_contributions[col] = term_values
                    
                    # 절편 기여도
                    intercept_contribution = np.full(len(X_sample), lr.intercept_)
                    term_contributions['절편'] = intercept_contribution
                    
                    # 예측값
                    y_pred_sample = lr.predict(X_sample)
                    
                    # 각 항의 평균 기여도와 비중 계산
                    term_analysis = []
                    for term_name, contributions in term_contributions.items():
                        avg_contribution = np.mean(contributions)
                        contribution_ratio = (avg_contribution / np.mean(y_pred_sample)) * 100
                        
                        if term_name == '절편':
                            term_analysis.append({
                                'name': term_name,
                                'avg_contribution': avg_contribution,
                                'ratio': contribution_ratio,
                                'color': '#6c757d'
                            })
                        else:
                            term_analysis.append({
                                'name': term_name,
                                'avg_contribution': avg_contribution,
                                'ratio': contribution_ratio,
                                'color': '#28a745'
                            })
                    
                    # 비중 순으로 정렬
                    term_analysis.sort(key=lambda x: abs(x['ratio']), reverse=True)
                    
                    for term in term_analysis:
                        if term['avg_contribution'] > 0:
                            impact_icon = "📈"
                            impact_text = "양의 기여"
                        else:
                            impact_icon = "📉"
                            impact_text = "음의 기여"
                        
                        st.markdown(f"""
                        <div style="margin: 8px 0; padding: 10px; background-color: white; border-radius: 6px; border-left: 4px solid {term['color']};">
                            <div style="display: flex; justify-content: space-between; align-items: center; margin-bottom: 5px;">
                                <strong style="color: {term['color']};">{term['name']}</strong>
                                <span style="font-size: 12px; color: #666;">{impact_icon} {impact_text}</span>
                            </div>
                            <div style="font-size: 14px; color: #333;">
                                평균 기여도: <strong>{term['avg_contribution']:.2f}</strong>
                                <br>Y에 대한 비중: <strong>{term['ratio']:.1f}%</strong>
                            </div>
                        </div>
                        """, unsafe_allow_html=True)
                    
                    st.markdown("</div>", unsafe_allow_html=True)
                    
                    # 모델 성능 계산 (col2에서 사용하기 위해 미리 계산)
                    score = lr.score(X_test, y_test)
                    
                    # 분석 결과를 세션 상태에 저장
                    analysis_results = {
                        'equation': equation,
                        'model_performance': {
                            'r2': score,
                            'rmse': np.sqrt(np.mean((y_test - lr.predict(X_test))**2)),
                            'mae': np.mean(np.abs(y_test - lr.predict(X_test)))
                        },
                        'variable_importance': {},
                        'term_analysis': term_analysis
                    }
                    
                    # 변수 중요도 계산 (절대값 기준)
                    for term in term_analysis:
                        if term['name'] != '절편':
                            analysis_results['variable_importance'][term['name']] = abs(term['ratio'])
                    
                    st.session_state['analysis_results'] = analysis_results
                
                with col2:
                    # 성능 등급 결정
                    if score >= 0.8:
                        performance_grade = "🟢 우수"
                        performance_color = "#28a745"
                    elif score >= 0.6:
                        performance_grade = "🟡 양호"
                        performance_color = "#ffc107"
                    elif score >= 0.4:
                        performance_grade = "🟠 보통"
                        performance_color = "#fd7e14"
                    else:
                        performance_grade = "🔴 미흡"
                        performance_color = "#dc3545"
                    
                    st.markdown(f"""
                    <div style="background-color: #e9ecef; padding: 20px; border-radius: 10px; border-left: 5px solid #007bff;">
                        <h4 style="margin: 0 0 15px 0; color: #007bff;">📊 모델 성능</h4>
                        <div style="text-align: center;">
                            <div style="font-size: 36px; font-weight: bold; color: {performance_color}; margin: 10px 0;">
                                {score:.1%}
                            </div>
                            <div style="font-size: 18px; color: #6c757d; margin-bottom: 15px;">
                                R² (설명력)
                            </div>
                            <div style="background-color: white; padding: 10px; border-radius: 5px; font-weight: bold; color: {performance_color};">
                                {performance_grade}
                            </div>
                        </div>
                    </div>
                    """, unsafe_allow_html=True)
                    
                    # 인사이트 제공
                    st.markdown("""
                    <div style="background-color: #d1ecf1; padding: 15px; border-radius: 10px; border-left: 5px solid #17a2b8; margin-top: 15px;">
                        <h4 style="margin: 0 0 10px 0; color: #0c5460;">💡 주요 인사이트</h4>
                    """, unsafe_allow_html=True)
                    
                    if score < 0.3:
                        insight = "• 모델의 설명력이 낮습니다. 더 많은 변수나 비선형 관계를 고려해보세요."
                    elif score < 0.6:
                        insight = "• 모델의 설명력이 보통 수준입니다. 추가 변수나 특성 엔지니어링을 고려해보세요."
                    else:
                        insight = "• 모델의 설명력이 양호합니다. 선형 관계가 잘 포착되었습니다."
                    
                    # 항 해석 결과를 주요 인사이트에 추가
                    important_terms = []
                    less_important_terms = []
                    
                    for term in term_analysis:
                        if abs(term['ratio']) >= 5:  # 5% 이상인 항
                            important_terms.append(term)
                        else:  # 5% 미만인 항
                            less_important_terms.append(term)
                    
                    # 텍스트 설명 제거 - 빈 div만 유지
                    st.markdown("""
                    <div style="color: #0c5460; font-size: 16px; line-height: 1.8;">
                    </div>
                    """, unsafe_allow_html=True)
                    
                    # 항 해석 결과를 주요 인사이트의 하부 항목으로 표시
                    if important_terms:
                        st.markdown("""
                        <div style="margin: 5px 0; padding: 8px; background-color: white; border-radius: 5px; border-left: 3px solid #1976d2;">
                            <h4 style="margin: 0; color: #0d47a1;">🎯 핵심 영향 변수 (Y에 대한 비중 5% 이상)</h4>
                        </div>
                        """, unsafe_allow_html=True)
                        
                        for term in important_terms:
                            impact_icon = "📈" if term['avg_contribution'] > 0 else "📉"
                            impact_text = "양의 영향" if term['avg_contribution'] > 0 else "음의 영향"
                            
                            st.markdown(f"""
                            <div style="margin: 5px 0; padding: 8px; background-color: white; border-radius: 5px; border-left: 3px solid #1976d2;">
                                <div style="display: flex; justify-content: space-between; align-items: center;">
                                    <strong style="color: #0d47a1;">{term['name']}</strong>
                                    <span style="font-size: 12px; color: #1976d2;">{impact_icon} {impact_text}</span>
                                </div>
                                <div style="font-size: 13px; color: #0d47a1; margin-top: 3px;">
                                    평균 기여도: <strong>{term['avg_contribution']:.2f}</strong> | Y에 대한 비중: <strong>{term['ratio']:.1f}%</strong>
                                </div>
                            </div>
                            """, unsafe_allow_html=True)
                    
                    if less_important_terms:
                        st.markdown("""
                        <div style="margin: 5px 0; padding: 8px; background-color: white; border-radius: 5px; border-left: 3px solid #7b1fa2;">
                            <h4 style="margin: 0; color: #4a148c;">📊 보조 영향 변수 (Y에 대한 비중 5% 미만)</h4>
                        </div>
                        """, unsafe_allow_html=True)
                        
                        for term in less_important_terms:
                            impact_icon = "📈" if term['avg_contribution'] > 0 else "📉"
                            impact_text = "양의 영향" if term['avg_contribution'] > 0 else "음의 영향"
                            
                            st.markdown(f"""
                            <div style="margin: 5px 0; padding: 8px; background-color: white; border-radius: 5px; border-left: 3px solid #7b1fa2;">
                                <div style="display: flex; justify-content: space-between; align-items: center;">
                                    <strong style="color: #4a148c;">{term['name']}</strong>
                                    <span style="font-size: 12px; color: #7b1fa2;">{impact_icon} {impact_text}</span>
                                </div>
                                <div style="font-size: 13px; color: #4a148c; margin-top: 3px;">
                                    평균 기여도: <strong>{term['avg_contribution']:.2f}</strong> | Y에 대한 비중: <strong>{term['ratio']:.1f}%</strong>
                                </div>
                            </div>
                            """, unsafe_allow_html=True)
                    
                    st.markdown("</div>", unsafe_allow_html=True)
                    
                    # 변수 중요도 차트 생성
                    if term_analysis:
                        # 절편을 제외한 변수들만 선택
                        variables = [term['name'] for term in term_analysis if term['name'] != '절편']
                        importance_values = [abs(term['ratio']) for term in term_analysis if term['name'] != '절편']
                        
                        # 변수 중요도 차트 생성
                        fig_importance = px.bar(
                            x=variables,
                            y=importance_values,
                            title="변수 중요도 (Y에 대한 비중)",
                            labels={'x': '변수', 'y': '중요도 (%)'},
                            color=importance_values,
                            color_continuous_scale='viridis'
                        )
                        
                        fig_importance.update_layout(
                            xaxis_title="변수",
                            yaxis_title="중요도 (%)",
                            height=400,
                            showlegend=False
                        )
                        
                        # 변수 중요도 차트를 session state에 저장
                        st.session_state['importance_chart'] = fig_importance
                        
                        # 차트 표시
                        st.plotly_chart(fig_importance, use_container_width=True)
                
                # 예측 결과 그래프
                st.markdown("""
                <div style="margin: 5px 0; padding: 0px;">
                    <h4 style="text-align: left; color: #495057; margin: 0px; padding: 0px; line-height: 1.2;">📈 Target(Y값) 예측 vs 실제값 비교 시뮬레이션</h4>
                </div>
                """, unsafe_allow_html=True)
                
                # x축과 y축의 범위를 동일하게 설정
                min_val = min(y_test.min(), y_pred.min())
                max_val = max(y_test.max(), y_pred.max())
                
                # 예측값 vs 실제값 산점도
                fig = px.scatter(x=y_test, y=y_pred, 
                               labels={'x':'실제값', 'y':'예측값'}, 
                               title="",
                               color_discrete_sequence=['#007bff'])
                
                # 완벽한 예측선 추가
                fig.add_shape(type="line", x0=min_val, y0=min_val, 
                            x1=max_val, y1=max_val, 
                            line=dict(color="red", dash="dash", width=2))
                
                # 범례를 위한 더미 데이터 추가
                fig.add_scatter(x=[None], y=[None], mode='lines', 
                              line=dict(color='red', dash='dash', width=2),
                              name='Y=X 기준선 (완벽한 예측선)', showlegend=True)
                
                fig.update_layout(
                    title="",  # undefined 제거
                    plot_bgcolor='white',
                    paper_bgcolor='white',
                    font=dict(size=12),
                    xaxis=dict(
                        range=[min_val, max_val],
                        showgrid=True,
                        gridcolor='lightgray',
                        zeroline=True,
                        zerolinecolor='black',
                        zerolinewidth=1,
                        showline=True,
                        linecolor='black',
                        linewidth=1
                    ),
                    yaxis=dict(
                        range=[min_val, max_val],
                        showgrid=True,
                        gridcolor='lightgray',
                        zeroline=True,
                        zerolinecolor='black',
                        zerolinewidth=1,
                        showline=True,
                        linecolor='black',
                        linewidth=1
                    ),
                    width=600,  # 정방형을 위한 가로 크기
                    height=600,  # 정방형을 위한 세로 크기
                    showlegend=True,
                    legend=dict(
                        x=0.02,
                        y=0.98,
                        bgcolor='rgba(255,255,255,0.9)',
                        bordercolor='black',
                        borderwidth=1,
                        font=dict(size=12)
                    )
                )
                
                # 범례 텍스트 설정
                fig.data[0].name = "실제 데이터 (예측값 vs 실제값)"
                
                st.plotly_chart(fig, use_container_width=False)
                
                # 정량적인 정확성 데이터 추가
                from sklearn.metrics import r2_score, mean_squared_error, mean_absolute_error
                
                # 예측 정확도 지표 계산
                r2 = r2_score(y_test, y_pred)
                mse = mean_squared_error(y_test, y_pred)
                rmse = np.sqrt(mse)
                mae = mean_absolute_error(y_test, y_pred)
                
                # 예측값과 실제값의 차이 계산
                residuals = y_pred - y_test
                residual_std = np.std(residuals)
                residual_mean = np.mean(residuals)
                
                # 정확성 데이터를 컬럼으로 표시
                col1, col2, col3 = st.columns(3)
                
                with col1:
                    st.markdown("""
                    <div style="background-color: #e3f2fd; padding: 15px; border-radius: 10px; border-left: 5px solid #1976d2;">
                        <h4 style="margin: 0 0 10px 0; color: #0d47a1;">📊 모델 성능 지표</h4>
                        <div style="font-size: 14px; line-height: 1.6;">
                            <strong>R² 점수:</strong> {:.4f}<br>
                            <strong>RMSE:</strong> {:.4f}<br>
                            <strong>MAE:</strong> {:.4f}
                        </div>
                    </div>
                    """.format(r2, rmse, mae), unsafe_allow_html=True)
                
                with col2:
                    st.markdown("""
                    <div style="background-color: #f3e5f5; padding: 15px; border-radius: 10px; border-left: 5px solid #7b1fa2;">
                        <h4 style="margin: 0 0 10px 0; color: #4a148c;">📈 예측 오차 분석</h4>
                        <div style="font-size: 14px; line-height: 1.6;">
                            <strong>평균 오차:</strong> {:.4f}<br>
                            <strong>오차 표준편차:</strong> {:.4f}<br>
                            <strong>MSE:</strong> {:.4f}
                        </div>
                    </div>
                    """.format(residual_mean, residual_std, mse), unsafe_allow_html=True)
                
                with col3:
                    st.markdown("""
                    <div style="background-color: #e8f5e8; padding: 15px; border-radius: 10px; border-left: 5px solid #28a745;">
                        <h4 style="margin: 0 0 10px 0; color: #155724;">🎯 분포 특성</h4>
                        <div style="font-size: 14px; line-height: 1.6;">
                            <strong>데이터 포인트:</strong> {}개<br>
                            <strong>예측 범위:</strong> {:.2f} ~ {:.2f}<br>
                            <strong>실제 범위:</strong> {:.2f} ~ {:.2f}
                        </div>
                    </div>
                    """.format(len(y_test), y_pred.min(), y_pred.max(), y_test.min(), y_test.max()), unsafe_allow_html=True)

                # 여백 추가
                st.markdown("<div style='margin: 40px 0;'></div>", unsafe_allow_html=True)

                # 구분선 추가
                st.markdown('<hr style="border:2px solid #667eea; margin: 20px 0;">', unsafe_allow_html=True)

                # 머신러닝 분석
                st.markdown("""
                <div style="background: linear-gradient(135deg, #667eea 0%, #764ba2 100%); padding: 15px; border-radius: 15px; margin: 20px 0; color: white;">
                    <h3 style="margin: 0; font-size: 24px; text-align: center;">🧠 머신러닝 모델 분석 결과</h3>
                </div>
                """, unsafe_allow_html=True)
                
                from sklearn.ensemble import RandomForestRegressor
                from sklearn.svm import SVR
                from sklearn.metrics import r2_score, mean_squared_error, mean_absolute_error
                import numpy as np
                
                # 4가지 머신러닝 모델 정의
                from sklearn.ensemble import RandomForestRegressor, GradientBoostingRegressor
                from sklearn.svm import SVR
                from sklearn.neural_network import MLPRegressor
                
                models = {
                    'Random Forest': RandomForestRegressor(n_estimators=100, max_depth=10, random_state=42),
                    'Gradient Boosting': GradientBoostingRegressor(n_estimators=100, random_state=42),
                    'Support Vector Machine': SVR(kernel='rbf', C=1.0, epsilon=0.1),
                    'Neural Network': MLPRegressor(hidden_layer_sizes=(100, 50), max_iter=500, random_state=42)
                }
                
                # 모델 성능 비교
                model_results = {}
                
                with st.spinner("🧠 4가지 머신러닝 모델을 훈련하고 성능을 비교하는 중..."):
                    # 각 모델 훈련 및 평가
                    for name, model in models.items():
                        model.fit(X_train, y_train)
                        y_pred = model.predict(X_test)
                        
                        # 성능 지표 계산
                        r2 = r2_score(y_test, y_pred)
                        mse = mean_squared_error(y_test, y_pred)
                        mae = mean_absolute_error(y_test, y_pred)
                        rmse = np.sqrt(mse)
                        
                        model_results[name] = {
                            'r2': r2,
                            'rmse': rmse,
                            'mae': mae,
                            'y_pred': y_pred
                        }
                
                # 결과 표시
                col1, col2 = st.columns([1, 1])
                
                with col1:
                    # 모델 성능 비교 테이블
                    st.markdown("""
                    <div style="background-color: #f8f9fa; padding: 20px; border-radius: 10px; border-left: 5px solid #6f42c1;">
                        <h4 style="margin: 0 0 15px 0; color: #6f42c1;">🏆 모델 성능 비교</h4>
                    """, unsafe_allow_html=True)
                    
                    # 성능 순위 결정
                    sorted_models = sorted(model_results.items(), key=lambda x: x[1]['r2'], reverse=True)
                    best_model_name = sorted_models[0][0]
                    
                    for i, (name, results) in enumerate(sorted_models):
                        if name == best_model_name:
                            rank_icon = "🥇"
                            rank_color = "#ffd700"
                        elif i == 1:
                            rank_icon = "🥈"
                            rank_color = "#c0c0c0"
                        elif i == 2:
                            rank_icon = "🥉"
                            rank_color = "#cd7f32"
                        else:
                            rank_icon = "📊"
                            rank_color = "#6c757d"
                        
                        # 성능 등급 결정
                        if results['r2'] >= 0.8:
                            grade = "🟢 우수"
                            grade_color = "#28a745"
                        elif results['r2'] >= 0.6:
                            grade = "🟡 양호"
                            grade_color = "#ffc107"
                        elif results['r2'] >= 0.4:
                            grade = "🟠 보통"
                            grade_color = "#fd7e14"
                        else:
                            grade = "🔴 미흡"
                            grade_color = "#dc3545"
                        
                        # 모델별 설명
                        model_explanations = {
                            'Random Forest': '여러 의사결정 트리를 독립적으로 학습시켜 평균을 내는 앙상블 기법',
                            'Gradient Boosting': '여러 약한 예측기를 순차적으로 조합하여 정확도를 높이는 앙상블 기법',
                            'Support Vector Machine': '데이터를 고차원 공간으로 변환하여 비선형 패턴을 찾는 기법',
                            'Neural Network': '인간 뇌의 신경망을 모방하여 복잡한 패턴을 학습하는 기법'
                        }
                        
                        st.markdown(f"""
                        <div style="background-color: white; padding: 15px; border-radius: 8px; margin: 10px 0; border-left: 4px solid {rank_color};">
                            <div style="display: flex; align-items: center; margin-bottom: 8px;">
                                <span style="font-size: 20px; margin-right: 10px;">{rank_icon}</span>
                                <h5 style="margin: 0; color: #495057;">{name}</h5>
                            </div>
                            <p style="margin: 0 0 10px 0; color: #6c757d; font-size: 13px; font-style: italic;">
                                {model_explanations.get(name, '알 수 없는 모델')}
                            </p>
                            <div style="display: grid; grid-template-columns: 1fr 1fr; gap: 10px;">
                                <div style="text-align: center;">
                                    <div style="font-size: 24px; font-weight: bold; color: {grade_color};">
                                        {results['r2']:.1%}
                                    </div>
                                    <div style="font-size: 12px; color: #6c757d;">R²</div>
                                </div>
                                <div style="text-align: center;">
                                    <div style="font-size: 16px; font-weight: bold; color: #495057;">
                                        {grade}
                                    </div>
                                    <div style="font-size: 12px; color: #6c757d;">등급</div>
                                </div>
                            </div>
                        </div>
                        """, unsafe_allow_html=True)
                    
                    st.markdown("</div>", unsafe_allow_html=True)
                
                with col2:
                    # 최적 모델 정보
                    st.markdown(f"""
                    <div style="background-color: #e9ecef; padding: 20px; border-radius: 10px; border-left: 5px solid #28a745;">
                        <h4 style="margin: 0 0 15px 0; color: #495057;">🏆 최적 모델: {best_model_name}</h4>
                    """, unsafe_allow_html=True)
                    
                    best_results = model_results[best_model_name]
                    
                    # 모델별 특성 설명
                    model_descriptions = {
                        'Random Forest': '앙상블 기법으로 여러 결정 트리를 조합한 모델',
                        'Gradient Boosting': '순차적으로 약한 학습기를 강화하는 부스팅 모델',
                        'Support Vector Machine': '커널 트릭을 사용한 비선형 회귀 모델로, 비선형 관계를 학습할 수 있는 서포트 벡터 회귀(SVR)',
                        'Neural Network': '다층 퍼셉트론 기반 신경망 모델'
                    }
                    
                    st.markdown(f"""
                    <div style="background-color: white; padding: 15px; border-radius: 8px; margin-bottom: 15px;">
                        <h5 style="margin: 0 0 10px 0; color: #495057;">📋 모델 특성</h5>
                        <p style="margin: 0; color: #6c757d; font-size: 14px;">
                            {model_descriptions.get(best_model_name, '알 수 없는 모델')}
                        </p>
                    </div>
                    """, unsafe_allow_html=True)
                    
                    # 성능 지표
                    st.markdown("""
                    <div style="background-color: white; padding: 15px; border-radius: 8px;">
                        <h5 style="margin: 0 0 10px 0; color: #495057;">📊 성능 지표</h5>
                    """, unsafe_allow_html=True)
                    
                    st.markdown(f"""
                    <div style="display: grid; grid-template-columns: 1fr 1fr; gap: 10px; margin-bottom: 10px;">
                        <div style="text-align: center; padding: 8px; background-color: #f8f9fa; border-radius: 5px;">
                            <div style="font-size: 18px; font-weight: bold; color: #28a745;">
                                {best_results['r2']:.1%}
                            </div>
                            <div style="font-size: 12px; color: #6c757d;">설명력 (R²)</div>
                        </div>
                        <div style="text-align: center; padding: 8px; background-color: #f8f9fa; border-radius: 5px;">
                            <div style="font-size: 18px; font-weight: bold; color: #dc3545;">
                                {best_results['mae']:.4f}
                            </div>
                            <div style="font-size: 12px; color: #6c757d;">평균 절대 오차</div>
                        </div>
                    </div>
                    """, unsafe_allow_html=True)
                    
                    st.markdown("</div>", unsafe_allow_html=True)
                    
                    # 인사이트 제공
                    st.markdown("""
                    <div style="background-color: #d1ecf1; padding: 15px; border-radius: 10px; border-left: 5px solid #17a2b8; margin-top: 15px;">
                        <h4 style="margin: 0 0 10px 0; color: #0c5460;">💡 머신러닝 인사이트</h4>
                    """, unsafe_allow_html=True)
                    
                    # 모델 비교 인사이트
                    best_r2 = best_results['r2']
                    worst_r2 = min([results['r2'] for results in model_results.values()])
                    performance_gap = best_r2 - worst_r2
                    
                    insight = f"• {best_model_name}이 가장 좋은 성능을 보입니다 (R²: {best_r2:.4f})"
                    if performance_gap > 0.1:
                        insight += f" (성능 차이: {performance_gap:.4f})"
                    
                    
                    
                    st.markdown(f"""
                    <div style="color: #0c5460; font-size: 14px;">
                        {insight}
                        <br>• 4가지 머신러닝 기법으로 데이터 패턴을 포착했습니다.
                        <br>• 모델 복잡도를 고려할 때 {best_model_name}이 가장 적합한 모델로 선택되었음
                        <br>• 예측 정확도: 평균 절대 오차 {best_results['mae']:.4f} 단위
                    </div>
                    """, unsafe_allow_html=True)
                    
                    st.markdown("</div>", unsafe_allow_html=True)
                
                # 머신러닝 모델 성능 비교 차트
                st.markdown("""
                <div style="margin: 5px 0; padding: 0px;">
                    <h4 style="text-align: left; color: #495057; margin: 0px; padding: 0px; line-height: 1.2;">📊 4가지 머신러닝 모델 성능 비교</h4>
                </div>
                """, unsafe_allow_html=True)
                
                # 모든 모델의 예측 결과를 하나의 차트에 표시
                fig_comparison = go.Figure()
                
                # 색상 설정 - 최적 모델은 파란색, 나머지는 흐린 색
                colors = {}
                # R² 값 순으로 모델 정렬하여 최적 모델 찾기
                sorted_models = sorted(model_results.items(), key=lambda x: x[1]['r2'], reverse=True)
                best_model = sorted_models[0][0]  # R²가 가장 높은 모델
                
                # 최적 모델은 파란색, 나머지는 흐린 색으로 설정
                for model_name in model_results.keys():
                    if model_name == best_model:
                        colors[model_name] = '#007bff'  # 파란색 (최적 모델)
                    else:
                        # 나머지 모델들은 흐린 색으로 설정
                        if model_name == 'Random Forest':
                            colors[model_name] = '#87ceeb'  # 연한 하늘색
                        elif model_name == 'Neural Network':
                            colors[model_name] = '#dda0dd'  # 연한 보라색
                        elif model_name == 'Support Vector Machine':
                            colors[model_name] = '#98fb98'  # 연한 연두색
                        else:
                            colors[model_name] = '#cccccc'  # 기본 회색
                
                # 이미 위에서 sorted_models가 정의되었으므로 제거
                
                # 먼저 나머지 모델들을 추가 (뒤에 표시)
                for model_name, results in sorted_models:
                    if model_name in colors and model_name != best_model_name:
                        fig_comparison.add_trace(go.Scatter(
                            x=y_test,
                            y=results['y_pred'],
                            mode='markers',
                            name=f"{model_name} (R²: {results['r2']:.3f})",
                            marker=dict(
                                color=colors[model_name],
                                size=8,  # 점 크기 증가
                                opacity=0.6  # 반투명
                            ),
                            hovertemplate=f'<b>{model_name}</b><br>' +
                                        '실제값: %{x:.3f}<br>' +
                                        '예측값: %{y:.3f}<br>' +
                                        f'R²: {results["r2"]:.3f}<extra></extra>'
                        ))
                
                # 마지막에 최적 모델을 추가 (맨 앞에 표시)
                best_model_name = best_model  # 최적 모델명 업데이트
                best_results = model_results[best_model_name]
                fig_comparison.add_trace(go.Scatter(
                    x=y_test,
                    y=best_results['y_pred'],
                    mode='markers',
                    name=f"{best_model_name} (R²: {best_results['r2']:.3f})",
                    marker=dict(
                        color=colors[best_model_name],
                        size=10,  # 최적 모델 점 크기 더 증가
                        opacity=1.0  # 불투명
                    ),
                    hovertemplate=f'<b>{best_model_name}</b><br>' +
                                '실제값: %{x:.3f}<br>' +
                                '예측값: %{y:.3f}<br>' +
                                f'R²: {best_results["r2"]:.3f}<extra></extra>'
                ))
                
                # 완벽한 예측선 추가
                min_val = min(y_test.min(), min([results['y_pred'].min() for results in model_results.values()]))
                max_val = max(y_test.max(), max([results['y_pred'].max() for results in model_results.values()]))
                
                fig_comparison.add_trace(go.Scatter(
                    x=[min_val, max_val],
                    y=[min_val, max_val],
                    mode='lines',
                    name='완벽한 예측선 (Y=X)',
                    line=dict(color='red', dash='dash', width=2),
                    showlegend=True
                ))
                
                fig_comparison.update_layout(
                    title="",
                    plot_bgcolor='white',
                    paper_bgcolor='white',
                    font=dict(size=12),
                    margin=dict(l=50, r=50, t=50, b=50),  # 균등한 여백으로 정방형 복원
                    xaxis=dict(
                        title="실제값",
                        range=[min_val, max_val],
                        showgrid=True,
                        gridcolor='lightgray',
                        zeroline=True,
                        zerolinecolor='black',
                        zerolinewidth=1,
                        showline=True,
                        linecolor='black',
                        linewidth=1
                    ),
                    yaxis=dict(
                        title="예측값",
                        range=[min_val, max_val],
                        showgrid=True,
                        gridcolor='lightgray',
                        zeroline=True,
                        zerolinecolor='black',
                        zerolinewidth=1,
                        showline=True,
                        linecolor='black',
                        linewidth=1
                    ),
                    width=680,  # 85% 크기로 조정
                    height=680,  # 85% 크기로 조정
                    showlegend=True,
                    legend=dict(
                        x=0.02,
                        y=0.98,
                        bgcolor='rgba(255,255,255,0.9)',
                        bordercolor='black',
                        borderwidth=1,
                        font=dict(size=12)
                    )
                )
                
                st.plotly_chart(fig_comparison, use_container_width=False)
                # 회귀 분석 차트를 session state에 저장
                st.session_state['regression_chart'] = fig_comparison
                



# 다운로드 기능 함수들
def create_word_document():
    """현재 화면의 모든 내용을 캡처하여 Word 문서로 생성"""
    import streamlit as st
    import plotly.io as pio
    import base64
    import io
    from PIL import Image
    import requests
    
    doc = Document()
    
    try:
        # 제목 추가
        title = doc.add_heading('공정 데이터 상관관계 분석 보고서', 0)
        title.alignment = 1  # 가운데 정렬
        
        # 생성 날짜 추가
        doc.add_paragraph(f'생성 날짜: {datetime.now().strftime("%Y년 %m월 %d일 %H:%M")}')
        doc.add_paragraph('')
        
        # 1. 파일 업로드 섹션
        doc.add_heading('1. 파일 업로드', level=1)
        if 'data_file_path' in st.session_state and st.session_state['data_file_path']:
            doc.add_paragraph(f'업로드된 파일: {st.session_state["data_file_path"]}')
        doc.add_paragraph('')
        
        # 2. 데이터 전처리 섹션
        doc.add_heading('2. 데이터 전처리', level=1)
        
        if 'df' in st.session_state and st.session_state['df'] is not None:
            df = st.session_state['df']
            
            # 데이터 요약 정보
            doc.add_heading('2.1 데이터 요약', level=2)
            doc.add_paragraph(f'총행 수: {df.shape[0]:,}개')
            doc.add_paragraph(f'총열 수: {df.shape[1]:,}개')
            
            # 결측치 정보
            missing_count = df.isnull().sum().sum()
            total_cells = df.shape[0] * df.shape[1]
            missing_percentage = (missing_count / total_cells) * 100
            doc.add_paragraph(f'총 결측치: {missing_count:,}개 (총 {total_cells:,}개 데이터 중 {missing_percentage:.1f}%)')
            doc.add_paragraph('')
            
            # 날짜형 변환 섹션 (색상 박스 형태로)
            if 'date_column_name' in st.session_state:
                doc.add_heading('2.2 날짜형 변환', level=2)
                doc.add_paragraph('날짜/시간 데이터가 문자형으로 인식되었습니다. 날짜형으로 변환하면 더 정확한 분석이 가능합니다.')
                doc.add_paragraph(f'변환된 날짜 열: {st.session_state["date_column_name"]}')
                
                # 데이터 기간 정보
                if 'data_period' in st.session_state:
                    period = st.session_state['data_period']
                    doc.add_paragraph(f'기간: {period["start_date"]} ~ {period["end_date"]}')
                    doc.add_paragraph(f'날짜 열: {period["date_column"]}')
                doc.add_paragraph('')
        
        # 3. 상관관계 분석 섹션
        if 'correlation_matrix' in st.session_state:
            doc.add_heading('3. 상관관계 분석', level=1)
            
            # 상관관계 히트맵 이미지 추가
            if 'correlation_heatmap' in st.session_state:
                try:
                    # Plotly 차트를 이미지로 변환
                    fig = st.session_state['correlation_heatmap']
                    img_bytes = pio.to_image(fig, format='png')
                    img_stream = io.BytesIO(img_bytes)
                    doc.add_picture(img_stream, width=Inches(6))
                    doc.add_paragraph('')
                except Exception as e:
                    doc.add_paragraph('상관관계 히트맵 이미지 삽입 실패')
            
            # 높은 상관관계 변수 쌍
            if 'high_correlation_pairs' in st.session_state:
                high_corr_pairs = st.session_state['high_correlation_pairs']
                if high_corr_pairs:
                    doc.add_heading('3.1 높은 상관관계 변수 쌍', level=2)
                    for pair in high_corr_pairs:
                        doc.add_paragraph(f'• {pair["var1"]} ↔ {pair["var2"]}: r = {pair["correlation"]:.3f}')
                    doc.add_paragraph('')
        
        # 4. 회귀 분석 섹션
        if 'analysis_results' in st.session_state:
            doc.add_heading('4. 회귀 분석', level=1)
            results = st.session_state['analysis_results']
            
            # 회귀 방정식
            if 'equation' in results:
                doc.add_heading('4.1 회귀 방정식', level=2)
                equation_text = results['equation'].replace('<br>', '\n').replace('<span style="color: #6c757d; font-size: 14px;">', '').replace('</span>', '')
                doc.add_paragraph(equation_text)
                doc.add_paragraph('')
            
            # 모델 성능
            if 'model_performance' in results:
                doc.add_heading('4.2 모델 성능', level=2)
                performance = results['model_performance']
                doc.add_paragraph(f'R² (설명력): {performance.get("r2", "N/A"):.4f}')
                doc.add_paragraph(f'RMSE: {performance.get("rmse", "N/A"):.4f}')
                doc.add_paragraph(f'MAE: {performance.get("mae", "N/A"):.4f}')
                doc.add_paragraph('')
            
            # 변수 중요도 차트
            if 'term_analysis' in results:
                doc.add_heading('4.3 변수 중요도', level=2)
                
                # 변수 중요도 차트 이미지 추가
                if 'importance_chart' in st.session_state:
                    try:
                        fig = st.session_state['importance_chart']
                        img_bytes = pio.to_image(fig, format='png')
                        img_stream = io.BytesIO(img_bytes)
                        doc.add_picture(img_stream, width=Inches(6))
                        doc.add_paragraph('')
                    except Exception as e:
                        doc.add_paragraph('변수 중요도 차트 이미지 삽입 실패')
                
                # 변수 중요도 텍스트
                term_analysis = results['term_analysis']
                sorted_terms = sorted(term_analysis, key=lambda x: abs(x['ratio']), reverse=True)
                
                for term in sorted_terms:
                    if term['name'] == '절편':
                        doc.add_paragraph(f'(절편): 평균 기여도 {term["avg_contribution"]:.2f}, Y에 대한 비중 {term["ratio"]:.1f}%')
                    else:
                        impact = "양의 영향" if term['avg_contribution'] > 0 else "음의 영향"
                        doc.add_paragraph(f'{term["name"]}: 평균 기여도 {term["avg_contribution"]:.2f}, Y에 대한 비중 {term["ratio"]:.1f}% ({impact})')
                doc.add_paragraph('')
        
        # 5. 시각화 섹션
        doc.add_heading('5. 시각화', level=1)
        
        # 히스토그램 이미지들 추가
        if 'histogram_charts' in st.session_state:
            doc.add_heading('5.1 데이터 분포 히스토그램', level=2)
            for i, fig in enumerate(st.session_state['histogram_charts']):
                try:
                    img_bytes = pio.to_image(fig, format='png')
                    img_stream = io.BytesIO(img_bytes)
                    doc.add_picture(img_stream, width=Inches(6))
                    doc.add_paragraph('')
                except Exception as e:
                    doc.add_paragraph(f'히스토그램 {i+1} 이미지 삽입 실패')
        
        # 산점도 이미지들 추가
        if 'scatter_charts' in st.session_state:
            doc.add_heading('5.2 산점도', level=2)
            for i, fig in enumerate(st.session_state['scatter_charts']):
                try:
                    img_bytes = pio.to_image(fig, format='png')
                    img_stream = io.BytesIO(img_bytes)
                    doc.add_picture(img_stream, width=Inches(6))
                    doc.add_paragraph('')
                except Exception as e:
                    doc.add_paragraph(f'산점도 {i+1} 이미지 삽입 실패')
        
        # 회귀 분석 차트
        if 'regression_chart' in st.session_state:
            doc.add_heading('5.3 회귀 분석 차트', level=2)
            try:
                fig = st.session_state['regression_chart']
                img_bytes = pio.to_image(fig, format='png')
                img_stream = io.BytesIO(img_bytes)
                doc.add_picture(img_stream, width=Inches(6))
                doc.add_paragraph('')
            except Exception as e:
                doc.add_paragraph('회귀 분석 차트 이미지 삽입 실패')
        
        # 6. 결론 및 권장사항
        doc.add_heading('6. 결론 및 권장사항', level=1)
        doc.add_paragraph('이 분석을 통해 다음과 같은 인사이트를 얻을 수 있습니다:')
        
        if 'analysis_results' in st.session_state:
            results = st.session_state['analysis_results']
            if 'term_analysis' in results:
                # 양의 영향을 주는 변수들
                positive_vars = [term for term in results['term_analysis'] if term['avg_contribution'] > 0 and term['name'] != '절편']
                if positive_vars:
                    doc.add_paragraph('• 양의 영향을 주는 주요 변수들:')
                    for var in positive_vars[:3]:
                        doc.add_paragraph(f'  - {var["name"]} (기여도: {var["avg_contribution"]:.2f})')
                
                # 음의 영향을 주는 변수들
                negative_vars = [term for term in results['term_analysis'] if term['avg_contribution'] < 0 and term['name'] != '절편']
                if negative_vars:
                    doc.add_paragraph('• 음의 영향을 주는 주요 변수들:')
                    for var in negative_vars[:3]:
                        doc.add_paragraph(f'  - {var["name"]} (기여도: {var["avg_contribution"]:.2f})')
        
        doc.add_paragraph('')
        doc.add_paragraph('이 보고서는 공정 데이터의 상관관계를 분석하여 주요 영향 변수를 식별하고,')
        doc.add_paragraph('프로세스 최적화를 위한 인사이트를 제공합니다.')
        
    except Exception as e:
        st.error(f"Word 문서 생성 중 오류: {str(e)}")
        # 오류 발생 시 기본 텍스트만 추가
        doc.add_heading("공정 데이터 상관관계 분석 보고서", level=1)
        doc.add_paragraph("문서 생성에 실패했습니다. 기본 보고서를 생성합니다.")
    
    return doc

def create_pdf_report():
    """현재 페이지의 모든 내용을 PDF로 생성"""
    try:
        from reportlab.lib.pagesizes import letter, A4
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import inch
        from reportlab.lib import colors
        from reportlab.pdfgen import canvas
        import io
        
        # PDF 문서 생성
        buffer = io.BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=A4, rightMargin=72, leftMargin=72, topMargin=72, bottomMargin=18)
        story = []
        
        # 스타일 정의
        styles = getSampleStyleSheet()
        
        # 한글 폰트 설정
        try:
            from reportlab.pdfbase import pdfmetrics
            from reportlab.pdfbase.ttfonts import TTFont
            
            # 한글 폰트 등록 (시스템에 따라 다를 수 있음)
            try:
                pdfmetrics.registerFont(TTFont('NanumGothic', 'C:/Windows/Fonts/malgun.ttf'))
                font_name = 'NanumGothic'
            except:
                try:
                    pdfmetrics.registerFont(TTFont('NanumGothic', 'C:/Windows/Fonts/gulim.ttc'))
                    font_name = 'NanumGothic'
                except:
                    font_name = 'Helvetica'  # 기본 폰트 사용
        except:
            font_name = 'Helvetica'  # 기본 폰트 사용
        
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=16,
            spaceAfter=30,
            alignment=1,  # 가운데 정렬
            fontName=font_name
        )
        heading1_style = ParagraphStyle(
            'CustomHeading1',
            parent=styles['Heading1'],
            fontSize=14,
            spaceAfter=12,
            spaceBefore=12,
            fontName=font_name
        )
        heading2_style = ParagraphStyle(
            'CustomHeading2',
            parent=styles['Heading2'],
            fontSize=12,
            spaceAfter=6,
            spaceBefore=6,
            fontName=font_name
        )
        normal_style = ParagraphStyle(
            'CustomNormal',
            parent=styles['Normal'],
            fontSize=10,
            spaceAfter=6,
            fontName=font_name
        )
        
        # 제목
        story.append(Paragraph('공정 데이터 상관관계 분석 보고서', title_style))
        story.append(Spacer(1, 12))
        
        # 생성 날짜
        story.append(Paragraph(f'생성 날짜: {datetime.now().strftime("%Y년 %m월 %d일 %H:%M")}', normal_style))
        story.append(Spacer(1, 12))
        
        # 1. 파일 업로드 정보
        story.append(Paragraph('1. 파일 업로드', heading1_style))
        if 'data_file_path' in st.session_state and st.session_state['data_file_path']:
            story.append(Paragraph(f'업로드된 파일: {st.session_state["data_file_path"]}', normal_style))
        story.append(Spacer(1, 12))
        
        # 2. 데이터 전처리 정보
        story.append(Paragraph('2. 데이터 전처리', heading1_style))
        
        if 'df' in st.session_state and st.session_state['df'] is not None:
            df = st.session_state['df']
            
            # 데이터 개요
            story.append(Paragraph('2.1 데이터 개요', heading2_style))
            story.append(Paragraph(f'데이터 크기: {df.shape[0]}행 × {df.shape[1]}열', normal_style))
            story.append(Paragraph(f'변수 수: {len(df.columns)}개', normal_style))
            story.append(Spacer(1, 12))
            
            # 변수 목록
            story.append(Paragraph('변수 목록:', normal_style))
            for i, col in enumerate(df.columns, 1):
                story.append(Paragraph(f'{i}. {col}', normal_style))
            story.append(Spacer(1, 12))
            
                    # 결측치 정보
        missing_info = df.isnull().sum()
        if missing_info.sum() > 0:
            story.append(Paragraph('2.2 결측치 정보', heading2_style))
            for col, missing_count in missing_info.items():
                if missing_count > 0:
                    story.append(Paragraph(f'{col}: {missing_count}개 결측치', normal_style))
            story.append(Spacer(1, 12))
        
        # 데이터 타입 정보
        story.append(Paragraph('2.3 데이터 타입 정보', heading2_style))
        dtype_info = df.dtypes.value_counts()
        for dtype, count in dtype_info.items():
            story.append(Paragraph(f'{dtype}: {count}개 변수', normal_style))
        story.append(Spacer(1, 12))
        
        # 수치형 변수 통계
        numeric_cols = df.select_dtypes(include=[np.number]).columns
        if len(numeric_cols) > 0:
            story.append(Paragraph('2.4 수치형 변수 통계', heading2_style))
            story.append(Paragraph(f'수치형 변수 수: {len(numeric_cols)}개', normal_style))
            story.append(Paragraph(f'범주형 변수 수: {len(df.columns) - len(numeric_cols)}개', normal_style))
            story.append(Spacer(1, 12))
        
        # 3. 상관관계 분석 결과
        if 'correlation_matrix' in st.session_state:
            story.append(Paragraph('3. 상관관계 분석', heading1_style))
            
            corr_matrix = st.session_state['correlation_matrix']
            story.append(Paragraph('3.1 상관관계 매트릭스', heading2_style))
            story.append(Paragraph(f'분석된 변수 수: {len(corr_matrix.columns)}개', normal_style))
            
            # 상관관계 통계
            corr_values = corr_matrix.values[np.triu_indices_from(corr_matrix.values, k=1)]
            story.append(Paragraph(f'총 상관관계 쌍 수: {len(corr_values)}개', normal_style))
            story.append(Paragraph(f'평균 절댓값 상관계수: {np.mean(np.abs(corr_values)):.3f}', normal_style))
            story.append(Paragraph(f'최대 절댓값 상관계수: {np.max(np.abs(corr_values)):.3f}', normal_style))
            story.append(Spacer(1, 12))
            
            # 높은 상관관계 변수 쌍 정보
            if 'high_correlation_pairs' in st.session_state:
                high_corr_pairs = st.session_state['high_correlation_pairs']
                if high_corr_pairs:
                    story.append(Paragraph('3.2 높은 상관관계 변수 쌍 (|r| ≥ 0.7)', heading2_style))
                    story.append(Paragraph(f'높은 상관관계 쌍 수: {len(high_corr_pairs)}개', normal_style))
                    for pair in high_corr_pairs:
                        story.append(Paragraph(f'• {pair["var1"]} ↔ {pair["var2"]}: r = {pair["correlation"]:.3f}', normal_style))
                    story.append(Spacer(1, 12))
        
        # 4. 회귀 분석 결과
        if 'analysis_results' in st.session_state:
            results = st.session_state['analysis_results']
            story.append(Paragraph('4. 회귀 분석 결과', heading1_style))
            
            # 회귀 방정식
            if 'equation' in results:
                story.append(Paragraph('4.1 회귀 방정식', heading2_style))
                equation_text = results['equation'].replace('<br>', '\n').replace('<span style="color: #6c757d; font-size: 14px;">', '').replace('</span>', '')
                story.append(Paragraph(equation_text, normal_style))
                story.append(Spacer(1, 12))
            
            # 모델 성능
            if 'model_performance' in results:
                story.append(Paragraph('4.2 모델 성능', heading2_style))
                performance = results['model_performance']
                story.append(Paragraph(f'R² (설명력): {performance.get("r2", "N/A"):.4f}', normal_style))
                story.append(Paragraph(f'RMSE: {performance.get("rmse", "N/A"):.4f}', normal_style))
                story.append(Paragraph(f'MAE: {performance.get("mae", "N/A"):.4f}', normal_style))
                story.append(Spacer(1, 12))
            
            # 변수 중요도 및 기여도
            if 'term_analysis' in results:
                story.append(Paragraph('4.3 변수 중요도 및 기여도', heading2_style))
                term_analysis = results['term_analysis']
                
                # 중요도 순으로 정렬
                sorted_terms = sorted(term_analysis, key=lambda x: abs(x['ratio']), reverse=True)
                
                for term in sorted_terms:
                    if term['name'] == '절편':
                        story.append(Paragraph(f'(절편): 평균 기여도 {term["avg_contribution"]:.2f}, Y에 대한 비중 {term["ratio"]:.1f}%', normal_style))
                    else:
                        impact = "양의 영향" if term['avg_contribution'] > 0 else "음의 영향"
                        story.append(Paragraph(f'{term["name"]}: 평균 기여도 {term["avg_contribution"]:.2f}, Y에 대한 비중 {term["ratio"]:.1f}% ({impact})', normal_style))
                story.append(Spacer(1, 12))
            
            # 주요 인사이트
            story.append(Paragraph('5. 주요 인사이트', heading1_style))
            
            # 핵심 영향 변수 (5% 이상)
            important_vars = [term for term in results.get('term_analysis', []) if abs(term['ratio']) >= 5]
            if important_vars:
                story.append(Paragraph('5.1 핵심 영향 변수 (Y에 대한 비중 5% 이상)', heading2_style))
                for term in important_vars:
                    impact = "양의 영향" if term['avg_contribution'] > 0 else "음의 영향"
                    story.append(Paragraph(f'• {term["name"]}: {term["ratio"]:.1f}% ({impact})', normal_style))
                story.append(Spacer(1, 12))
            
            # 보조 영향 변수 (5% 미만)
            less_important_vars = [term for term in results.get('term_analysis', []) if abs(term['ratio']) < 5]
            if less_important_vars:
                story.append(Paragraph('5.2 보조 영향 변수 (Y에 대한 비중 5% 미만)', heading2_style))
                for term in less_important_vars:
                    impact = "양의 영향" if term['avg_contribution'] > 0 else "음의 영향"
                    story.append(Paragraph(f'• {term["name"]}: {term["ratio"]:.1f}% ({impact})', normal_style))
                story.append(Spacer(1, 12))
            
                    # 모델 해석
        story.append(Paragraph('5.3 모델 해석', heading2_style))
        if 'r2' in performance:
            r2 = performance['r2']
            if r2 >= 0.8:
                story.append(Paragraph('• 모델의 설명력이 매우 높습니다 (R² ≥ 0.8)', normal_style))
                story.append(Paragraph('  - 이 모델은 종속변수의 변화를 매우 잘 설명합니다', normal_style))
            elif r2 >= 0.6:
                story.append(Paragraph('• 모델의 설명력이 높습니다 (R² ≥ 0.6)', normal_style))
                story.append(Paragraph('  - 이 모델은 종속변수의 변화를 잘 설명합니다', normal_style))
            elif r2 >= 0.4:
                story.append(Paragraph('• 모델의 설명력이 보통입니다 (R² ≥ 0.4)', normal_style))
                story.append(Paragraph('  - 이 모델은 종속변수의 변화를 어느 정도 설명합니다', normal_style))
            else:
                story.append(Paragraph('• 모델의 설명력이 낮습니다 (R² < 0.4)', normal_style))
                story.append(Paragraph('  - 이 모델은 종속변수의 변화를 충분히 설명하지 못합니다', normal_style))
        
        # 가장 중요한 변수
        if sorted_terms:
            most_important = sorted_terms[0]
            if most_important['name'] != '절편':
                story.append(Paragraph(f'• 가장 중요한 변수: {most_important["name"]} (비중: {most_important["ratio"]:.1f}%)', normal_style))
                story.append(Paragraph(f'  - 이 변수가 종속변수에 가장 큰 영향을 미칩니다', normal_style))
        
        # 모델 성능 평가
        story.append(Paragraph('5.4 모델 성능 평가', heading2_style))
        if 'rmse' in performance and 'mae' in performance:
            rmse = performance['rmse']
            mae = performance['mae']
            story.append(Paragraph(f'• RMSE (Root Mean Square Error): {rmse:.4f}', normal_style))
            story.append(Paragraph(f'• MAE (Mean Absolute Error): {mae:.4f}', normal_style))
            if rmse < mae:
                story.append(Paragraph('  - RMSE가 MAE보다 작아 예측 오차가 상대적으로 균등하게 분포합니다', normal_style))
            else:
                story.append(Paragraph('  - MAE가 RMSE보다 작아 큰 오차가 적습니다', normal_style))
        
        story.append(Spacer(1, 12))
        
        # 6. 시각화 정보
        story.append(Paragraph('6. 시각화', heading1_style))
        story.append(Paragraph('이 보고서에는 다음과 같은 시각화가 포함됩니다:', normal_style))
        story.append(Paragraph('• 상관관계 히트맵', normal_style))
        story.append(Paragraph('• 산점도 (선택된 변수 쌍)', normal_style))
        story.append(Paragraph('• 회귀 분석 결과 차트', normal_style))
        story.append(Paragraph('• 변수 중요도 차트', normal_style))
        story.append(Paragraph('• 데이터 분포 히스토그램', normal_style))
        story.append(Paragraph('• 상관관계 네트워크', normal_style))
        story.append(Spacer(1, 12))
        
        # 시각화 상세 정보
        story.append(Paragraph('6.1 시각화 상세 정보', heading2_style))
        story.append(Paragraph('• 상관관계 히트맵: 모든 변수 간의 상관계수를 색상으로 표현', normal_style))
        story.append(Paragraph('• 산점도: 선택된 변수 쌍의 관계를 점으로 표현', normal_style))
        story.append(Paragraph('• 회귀 분석 차트: 실제값 vs 예측값 비교', normal_style))
        story.append(Paragraph('• 변수 중요도: 각 변수의 기여도를 막대그래프로 표현', normal_style))
        story.append(Paragraph('• 데이터 분포: 각 변수의 분포를 히스토그램으로 표현', normal_style))
        story.append(Paragraph('• 상관관계 네트워크: 높은 상관관계를 가진 변수들을 연결', normal_style))
        story.append(Spacer(1, 12))
        
        # 7. 결론 및 권장사항
        story.append(Paragraph('7. 결론 및 권장사항', heading1_style))
        story.append(Paragraph('이 분석을 통해 다음과 같은 인사이트를 얻을 수 있습니다:', normal_style))
        
        if 'analysis_results' in st.session_state:
            results = st.session_state['analysis_results']
            if 'term_analysis' in results:
                # 양의 영향을 주는 변수들
                positive_vars = [term for term in results['term_analysis'] if term['avg_contribution'] > 0 and term['name'] != '절편']
                if positive_vars:
                    story.append(Paragraph('• 양의 영향을 주는 주요 변수들:', normal_style))
                    for var in positive_vars[:3]:  # 상위 3개만
                        story.append(Paragraph(f'  - {var["name"]} (기여도: {var["avg_contribution"]:.2f})', normal_style))
                
                # 음의 영향을 주는 변수들
                negative_vars = [term for term in results['term_analysis'] if term['avg_contribution'] < 0 and term['name'] != '절편']
                if negative_vars:
                    story.append(Paragraph('• 음의 영향을 주는 주요 변수들:', normal_style))
                    for var in negative_vars[:3]:  # 상위 3개만
                        story.append(Paragraph(f'  - {var["name"]} (기여도: {var["avg_contribution"]:.2f})', normal_style))
        
        story.append(Spacer(1, 12))
        story.append(Paragraph('이 보고서는 공정 데이터의 상관관계를 분석하여 주요 영향 변수를 식별하고,', normal_style))
        story.append(Paragraph('프로세스 최적화를 위한 인사이트를 제공합니다.', normal_style))
        story.append(Spacer(1, 12))
        
        # 추가 권장사항
        story.append(Paragraph('8. 추가 권장사항', heading1_style))
        story.append(Paragraph('이 분석 결과를 바탕으로 다음과 같은 추가 작업을 권장합니다:', normal_style))
        story.append(Paragraph('• 핵심 변수에 대한 더 깊은 분석 수행', normal_style))
        story.append(Paragraph('• 시계열 분석을 통한 트렌드 파악', normal_style))
        story.append(Paragraph('• 예측 모델의 정기적인 업데이트', normal_style))
        story.append(Paragraph('• 새로운 변수 추가 시 모델 재검증', normal_style))
        story.append(Paragraph('• 실시간 모니터링 시스템 구축 검토', normal_style))
        story.append(Spacer(1, 12))
        
        # 보고서 정보
        story.append(Paragraph('9. 보고서 정보', heading1_style))
        story.append(Paragraph('• 생성 도구: 공정 데이터 상관관계 분석 도우미', normal_style))
        story.append(Paragraph('• 분석 방법: 다중 선형 회귀 분석', normal_style))
        story.append(Paragraph('• 시각화 도구: Plotly, Matplotlib', normal_style))
        story.append(Paragraph('• 데이터 전처리: 자동 결측치 처리, 정규화', normal_style))
        story.append(Spacer(1, 12))
        
        # PDF 생성
        doc.build(story)
        buffer.seek(0)
        return buffer.getvalue()
        
    except ImportError as e:
        st.error("PDF 생성을 위해 reportlab 라이브러리가 필요합니다. 'pip install reportlab'을 실행해주세요.")
        st.error(f"ImportError 상세: {str(e)}")
        return None
    except Exception as e:
        st.error(f"PDF 보고서 생성 중 오류: {str(e)}")
        st.error("오류 유형: " + type(e).__name__)
        return None

def download_report(format_type):
    """보고서 다운로드 함수"""
    try:
        if format_type == "Word":
            doc = create_word_document()
            
            # 임시 파일로 저장
            tmp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
            tmp_file_path = tmp_file.name
            tmp_file.close()
            
            try:
                doc.save(tmp_file_path)
                
                # 파일 읽기
                with open(tmp_file_path, 'rb') as f:
                    file_data = f.read()
                
                # 다운로드 버튼 생성
                st.download_button(
                    label="📄 Word 문서 다운로드",
                    data=file_data,
                    file_name=f"상관관계_분석_보고서_{datetime.now().strftime('%Y%m%d_%H%M')}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
            finally:
                # 임시 파일 삭제
                try:
                    if os.path.exists(tmp_file_path):
                        os.unlink(tmp_file_path)
                except Exception as e:
                    st.warning(f"임시 파일 삭제 중 오류: {str(e)}")
        
        elif format_type == "PDF":
            # PDF 생성
            pdf_data = create_pdf_report()
            
            if pdf_data:
                # 다운로드 버튼 생성
                st.download_button(
                    label="📄 PDF 보고서 다운로드",
                    data=pdf_data,
                    file_name=f"상관관계_분석_보고서_{datetime.now().strftime('%Y%m%d_%H%M')}.pdf",
                    mime="application/pdf"
                )
            else:
                st.error("PDF 생성에 실패했습니다.")
                st.info("💡 대안: Word 문서를 다운로드한 후 PDF로 변환해주세요.")
                st.info("또는 브라우저에서 Ctrl+P를 눌러 페이지를 PDF로 저장할 수 있습니다.")
                
    except Exception as e:
        st.error(f"보고서 생성 중 오류가 발생했습니다: {str(e)}")

 